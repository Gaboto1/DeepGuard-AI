# -*- coding: utf-8 -*-
"""
Generador — Informe Final PMBOK — DeepGuard AI
Evaluacion Final: Control y Desempeno del Proyecto segun PMBOK
Ejecutar: python tools/generar_informe_pmbok.py
"""
import sys, os
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta ────────────────────────────────────────────────────────────────────
AZUL_MARINO  = RGBColor(0x06, 0x1A, 0x40)
AZUL_MED     = RGBColor(0x1E, 0x63, 0xD4)
VERDE_OK     = RGBColor(0x10, 0x6B, 0x3C)
GRIS_CUERPO  = RGBColor(0x22, 0x2C, 0x38)
GRIS_TABLA   = RGBColor(0xEF, 0xF3, 0xF8)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_OSCURO  = RGBColor(0x2C, 0x3E, 0x50)
AMARILLO     = RGBColor(0xB8, 0x6A, 0x1A)

doc = Document()

# ── Margenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_run_font(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color

def add_heading(text, level=1, color=None, size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    if level == 1:
        run.bold = True
        run.font.size = Pt(size or 16)
        run.font.color.rgb = color or AZUL_MARINO
    elif level == 2:
        run.bold = True
        run.font.size = Pt(size or 13)
        run.font.color.rgb = color or AZUL_MED
    elif level == 3:
        run.bold = True
        run.font.size = Pt(size or 11.5)
        run.font.color.rgb = color or GRIS_OSCURO
    run.font.name = "Calibri"
    return p

def add_body(text, bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    if bold_prefix:
        r0 = p.add_run(bold_prefix + " ")
        set_run_font(r0, bold=True, size=11, color=GRIS_CUERPO)
    r = p.add_run(text)
    set_run_font(r, size=11, color=GRIS_CUERPO)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        r0 = p.add_run(bold_prefix + ": ")
        set_run_font(r0, bold=True, size=11, color=GRIS_CUERPO)
    r = p.add_run(text)
    set_run_font(r, size=11, color=GRIS_CUERPO)
    return p

def add_table(headers, rows, col_widths=None):
    n_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
    has_header = bool(headers)
    t = doc.add_table(rows=(1 if has_header else 0) + len(rows), cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    if has_header:
        hr = t.rows[0]
        for i, h in enumerate(headers):
            cell = hr.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "2C3E50")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"),   "clear")
            tc_pr.append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = BLANCO
            run.font.size = Pt(10)
            run.font.name = "Calibri"
    # Data rows
    row_offset = 1 if has_header else 0
    for ri, row in enumerate(rows):
        tr = t.rows[ri + row_offset]
        bg = "EFF3F8" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), bg)
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"),   "clear")
            tc_pr.append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = GRIS_CUERPO
    # Column widths
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("INSTITUTO NACIONAL DE CAPACITACIÓN PROFESIONAL")
set_run_font(r, bold=True, size=13, color=AZUL_MARINO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sede Santiago Centro — Ingeniería Informática")
set_run_font(r, size=12, color=GRIS_OSCURO)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("INFORME FINAL DE GESTIÓN DEL PROYECTO")
set_run_font(r, bold=True, size=20, color=AZUL_MARINO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Evaluación Final: Control y Desempeño del Proyecto según PMBOK")
set_run_font(r, bold=True, size=14, color=AZUL_MED)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DEEPGUARD AI")
set_run_font(r, bold=True, size=22, color=AZUL_MARINO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema Forense de Detección de Deepfakes")
set_run_font(r, italic=True, size=14, color=GRIS_OSCURO)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_table(
    headers=[],
    rows=[
        ["Estudiante:",    "Rojas Gaboto X."],
        ["Asignatura:",    "Gestión de Proyectos Informáticos"],
        ["Docente:",       "Profesor(a) Guía INACAP"],
        ["Fecha entrega:", "Julio 2026"],
        ["Versión:",       "1.0 — Entrega Final"],
    ],
    col_widths=[4.5, 9.0]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  ÍNDICE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("ÍNDICE", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
indice = [
    ("1.", "Resumen Ejecutivo", "3"),
    ("2.", "Descripción del Proyecto", "3"),
    ("3.", "Gestión de la Calidad del Proyecto", "5"),
    ("4.", "Gestión de los Recursos del Proyecto", "7"),
    ("5.", "Gestión de las Comunicaciones del Proyecto", "9"),
    ("6.", "Gestión de las Adquisiciones del Proyecto", "11"),
    ("7.", "Procesos PMBOK Aplicados", "13"),
    ("8.", "Análisis de Desviaciones del Proyecto", "16"),
    ("9.", "Evaluación de los Interesados", "18"),
    ("10.", "Organización y Gestión del Equipo", "20"),
    ("11.", "Lecciones Aprendidas", "22"),
    ("12.", "Conclusiones", "23"),
    ("Referencias Bibliográficas", "", "24"),
    ("Anexos", "", "25"),
]
for num, title, pg in indice:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(f"  {num}  {title}")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = GRIS_CUERPO
    if num in ("Referencias Bibliográficas", "Anexos"):
        run.bold = True

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. RESUMEN EJECUTIVO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("1. RESUMEN EJECUTIVO", level=1)
add_body(
    "DeepGuard AI es una plataforma forense profesional de detección de deepfakes desarrollada "
    "como proyecto de título en el marco del programa de Ingeniería Informática de INACAP. "
    "La plataforma permite analizar imágenes y videos para determinar la probabilidad de que "
    "hayan sido manipulados o generados mediante inteligencia artificial, aportando una cadena "
    "de custodia digital verificable mediante sellos HMAC-SHA256."
)
add_body(
    "El sistema implementa un ensemble de ocho señales forenses — seis modelos transformers "
    "(ViT, SDXL, CLIP, AI-Art, SigLIP, AI-Human) y tres señales de procesamiento de imagen "
    "(ELA, FFT frecuencial, SRM de ruido) — coordinados por un meta-ensemble XGBoost con "
    "temperature scaling (T=0.581), logrando un F1-Score de 94.7% y un error de calibración ECE "
    "de 0.084 sobre el conjunto de evaluación."
)
add_body(
    "La arquitectura de despliegue es híbrida: la API en la nube (Render, gratuito permanente) "
    "recibe las solicitudes del frontend estático y delega el procesamiento computacional a un "
    "worker GPU local (RTX 4070 SUPER, 12 GB VRAM) a través de colas Celery sobre Upstash Redis "
    "(TLS, plan gratuito permanente). Al cierre del semestre el sistema se encuentra operativo en "
    "producción, con 18 pruebas automatizadas pasando y monitoreo continuo de salud mediante "
    "GitHub Actions."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  2. DESCRIPCIÓN DEL PROYECTO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("2. DESCRIPCIÓN DEL PROYECTO", level=1)

add_heading("2.1 Nombre del proyecto", level=2)
add_body("DeepGuard AI — Sistema Forense de Detección de Deepfakes")

add_heading("2.2 Descripción general", level=2)
add_body(
    "DeepGuard AI es una aplicación web de uso forense que permite a periodistas, "
    "investigadores, equipos legales y ciudadanos cargar una imagen o video sospechoso "
    "y obtener en segundos un dictamen técnico sobre su autenticidad. El dictamen incluye "
    "la probabilidad de manipulación por IA, el nivel de evidencia, mapas de calor Grad-CAM "
    "que señalan zonas alteradas, metadatos EXIF/XMP/IPTC, análisis OSINT de hash perceptual "
    "y una cadena de custodia criptográfica sellada con HMAC-SHA256."
)

add_heading("2.3 Problema u oportunidad identificada", level=2)
add_body(
    "La proliferación de contenido deepfake generado por IA representa una amenaza creciente "
    "para la veracidad de la información y la integridad de evidencia digital en contextos "
    "judiciales, periodísticos y académicos. Las herramientas existentes son o bien de pago, "
    "o de uso técnico muy avanzado, o carecen de trazabilidad forense. DeepGuard AI busca "
    "cubrir ese vacío con una solución de código abierto, de alto rendimiento y con cadena "
    "de custodia auditable."
)

add_heading("2.4 Objetivos del proyecto", level=2)
add_bullet("Desarrollar un sistema de detección de deepfakes con F1-Score superior al 90%.")
add_bullet("Implementar una cadena de custodia digital verificable mediante sellos criptográficos HMAC-SHA256.")
add_bullet("Desplegar la plataforma en producción con servicios 100% gratuitos permanentes.")
add_bullet("Crear una suite de pruebas automatizadas que cubra los flujos críticos del sistema.")
add_bullet("Documentar la arquitectura técnica y los procedimientos operativos para futura mantención.")

add_heading("2.5 Alcance del proyecto", level=2)
add_body(
    "El alcance abarca el desarrollo completo full-stack: pipeline de ML con ensemble de 8 señales, "
    "API REST FastAPI con dos versiones (legacy /api/* y enterprise /api/v1/*), worker Celery GPU, "
    "frontend Next.js 14 con interfaz forense, cadena de custodia, despliegue en nube y "
    "monitoreo continuo. Se excluyeron: sistema de autenticación/usuarios, base de datos relacional "
    "persistente, integración con plataformas de redes sociales, y aplicación móvil nativa."
)

add_heading("2.6 Estado actual del proyecto", level=2)
add_body(
    "El proyecto se encuentra en producción funcional al cierre del semestre. La API responde en "
    "https://deepguard-ai-api.onrender.com, el worker GPU local procesa tareas en tiempo real "
    "(latencia imagen < 5 s en RTX 4070 SUPER) y el frontend estático se hospeda en Render Static Site. "
    "Se completaron 18 pruebas automatizadas (100% pasando), se corrigieron 3 bugs críticos "
    "descubiertos durante la auditoría técnica, y se migró exitosamente de dos proveedores que "
    "expiraron (Aiven Valkey y Netlify) a alternativas gratuitas permanentes (Upstash Redis y Render)."
)

add_table(
    headers=["Componente", "Estado", "Plataforma", "Observación"],
    rows=[
        ["API Backend (FastAPI)", "EN LÍNEA", "Render (free)", "API_ONLY=true, sin GPU"],
        ["Worker GPU (Celery)", "OPERATIVO", "PC Local RTX 4070", "~7 GB VRAM en uso"],
        ["Message Broker", "EN LÍNEA", "Upstash Redis", "Plan gratuito permanente"],
        ["Frontend Web", "EN LÍNEA", "Render Static Site", "Next.js export estático"],
        ["Monitoreo", "ACTIVO", "GitHub Actions", "Cron cada 3 horas"],
        ["Suite de pruebas", "18/18 PASANDO", "pytest", "Cubre rutas críticas"],
    ],
    col_widths=[4.5, 2.8, 3.8, 4.4]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. GESTIÓN DE LA CALIDAD
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("3. GESTIÓN DE LA CALIDAD DEL PROYECTO", level=1)

add_heading("3.1 Objetivos de calidad", level=2)
add_body(
    "Desde el inicio del proyecto se establecieron objetivos de calidad medibles y verificables, "
    "orientados tanto al rendimiento del modelo de IA como a la confiabilidad del sistema en "
    "producción. El objetivo central fue alcanzar un F1-Score superior al 90% en la detección "
    "de deepfakes, asegurando que los falsos positivos y negativos se mantuvieran dentro de "
    "rangos aceptables para uso forense. Adicionalmente, se exigió que el tiempo de análisis "
    "para imágenes no superara los 5 segundos en hardware GPU local, y que la API respondiera "
    "dentro de 500 ms para solicitudes de estado (polling de tareas)."
)
add_body(
    "En cuanto al software, el objetivo fue mantener cero errores de compilación TypeScript "
    "en el frontend, cero vulnerabilidades de seguridad en los endpoints (validación de tipo MIME real, "
    "no solo extensión), y una cobertura de pruebas automatizadas que cubriera al menos los tres "
    "flujos críticos identificados: correcciones forenses, consenso de ensemble y cadena de custodia. "
    "Estos objetivos se convirtieron en criterios de aceptación para las funcionalidades entregadas."
)

add_heading("3.2 Estándares y criterios de calidad", level=2)
add_body(
    "Para la capa de ML se adoptaron las métricas estándar del campo: F1-Score, Precisión, "
    "Recall y Error de Calibración Esperado (ECE). El umbral de decisión del meta-ensemble "
    "(T=0.581) fue ajustado mediante temperature scaling para minimizar el ECE sobre el conjunto "
    "de validación. Para la capa de software, el estándar de calidad fue la tipificación "
    "estricta de TypeScript (strict mode) y la conformidad con los contratos Pydantic definidos "
    "en los schemas del backend."
)
add_body(
    "Se definió un criterio de calidad de seguridad específico: ninguna clave de firma "
    "(DEEPGUARD_SIGNING_KEY) en valor de marcador de posición puede iniciar el worker GPU. "
    "Este criterio se implementó mediante un mecanismo fail-fast en `custody_service.py` que "
    "rechaza un conjunto de valores inseguros conocidos y detiene el proceso con un RuntimeError "
    "descriptivo antes de que pueda generarse cualquier sello HMAC comprometido."
)
add_table(
    headers=["Criterio de calidad", "Indicador", "Meta", "Resultado"],
    rows=[
        ["Precisión del modelo", "F1-Score", "> 90%", "94.7% ✓"],
        ["Calibración probabilística", "ECE", "< 0.10", "0.084 ✓"],
        ["Latencia de análisis (imagen)", "Tiempo GPU", "< 5 s", "~2.0 s ✓"],
        ["Errores de compilación frontend", "tsc --noEmit", "0 errores", "0 errores ✓"],
        ["Pruebas automatizadas", "pytest", "18/18", "18/18 ✓"],
        ["Seguridad de clave de firma", "Fail-fast en arranque", "Sin placeholders", "Implementado ✓"],
        ["Validación de tipo de archivo", "MIME real (no ext)", "Sin bypass", "Implementado ✓"],
    ],
    col_widths=[5.0, 4.0, 3.0, 3.5]
)

add_heading("3.3 Actividades de aseguramiento de calidad", level=2)
add_body(
    "El aseguramiento de calidad se centró en revisiones de diseño y auditorías técnicas "
    "preventivas. Durante la fase de diseño del ensemble, se realizó un análisis de correlación "
    "entre señales para garantizar diversidad de información (evitar señales redundantes que no "
    "agregan poder discriminativo). Se documentó cada regla de corrección forense (OOD Bypass, "
    "Compression Veto, Consensus Override, F+SRM Alignment) con casos de ejemplo trazables a "
    "escenarios reales, asegurando que las correcciones fueran reproducibles y auditables."
)
add_body(
    "Se realizó una auditoría técnica completa del código en 18 dimensiones —desde arquitectura "
    "y modelos hasta seguridad y documentación— cuyos hallazgos quedaron registrados en el "
    "documento `docs/TESIS_AUDITORIA_TECNICA_DEEPGUARD_AI.md`. Esta auditoría sirvió como "
    "insumo para el plan de mejoras y como evidencia de aplicación de prácticas de calidad "
    "sistemáticas a lo largo del desarrollo."
)

add_heading("3.4 Actividades de control de calidad", level=2)
add_body(
    "El control de calidad se operacionalizó mediante tres mecanismos principales. El primero fue "
    "la suite de pruebas automatizadas con pytest: 18 tests organizados en tres módulos "
    "(`test_forensic_corrections.py`, `test_meta_ensemble_veto.py`, `test_custody_signing_key.py`) "
    "que verifican los invariantes del sistema forense. El segundo fue el workflow de monitoreo "
    "continuo en GitHub Actions (`health-check.yml`) que cada 3 horas consulta el endpoint "
    "`/api/v1/health` y notifica vía email si Redis o la API no responden. El tercero fue la "
    "compilación TypeScript estricta del frontend como gate previo a cada despliegue."
)
add_body(
    "Durante el control de calidad se identificaron y corrigieron tres defectos críticos: "
    "(1) el componente ForensicPanel mostraba incondicionalmente 'Contenido Auténtico / VÁLIDO' "
    "sin importar el resultado real del análisis; (2) el endpoint legacy /api/analyze omitía la "
    "carga Base64 del archivo al despachar tareas Celery, causando FileNotFoundError en el worker; "
    "y (3) la validación de la clave de firma aceptaba valores de marcador de posición conocidos. "
    "Los tres fueron corregidos antes de la entrega final."
)

add_heading("3.5 Resultados obtenidos", level=2)
add_body(
    "El sistema alcanzó todos los objetivos de calidad establecidos. El F1-Score final del "
    "meta-ensemble es 94.7%, con ECE 0.084, ambos dentro de los rangos objetivo. La latencia "
    "de análisis de imágenes en el hardware GPU local fue de aproximadamente 2.0 segundos, "
    "muy por debajo del umbral de 5 segundos. Las 18 pruebas automatizadas pasan en la "
    "totalidad de los casos, incluyendo los tres casos de fallo esperado que validan el "
    "comportamiento de rechazo de claves inseguras."
)

add_heading("3.6 Propuestas de mejora", level=2)
add_bullet("Aumentar la cobertura de pruebas a los pipelines de imagen y video completos (pruebas de integración end-to-end).")
add_bullet("Agregar benchmarks de rendimiento automatizados que fallen el CI si la latencia GPU supera el umbral.")
add_bullet("Implementar pruebas de mutación para verificar la robustez de los tests existentes.")
add_bullet("Establecer un proceso formal de revisión de código (pull request review) incluso en proyectos unipersonales.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. GESTIÓN DE LOS RECURSOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("4. GESTIÓN DE LOS RECURSOS DEL PROYECTO", level=1)

add_heading("4.1 Estructura del equipo", level=2)
add_body(
    "El proyecto DeepGuard AI fue desarrollado por un equipo unipersonal, compuesto por el "
    "estudiante de Ingeniería Informática Rojas Gaboto X., bajo la supervisión del docente "
    "guía de la asignatura. Esta estructura de equipo redujo la complejidad de coordinación "
    "interpersonal, pero exigió al desarrollador cubrir simultáneamente todos los dominios "
    "técnicos: aprendizaje automático, desarrollo backend, desarrollo frontend y operaciones "
    "en la nube (DevOps). La toma de decisiones fue centralizada y ágil, lo que permitió "
    "iterar rápidamente en las fases críticas del proyecto."
)

add_heading("4.2 Roles y responsabilidades", level=2)
add_table(
    headers=["Rol", "Responsabilidad principal", "Entregables"],
    rows=[
        ["Jefe de Proyecto / PM", "Planificación, seguimiento y control de alcance/tiempo", "Carta Gantt, actas, informes de avance"],
        ["Ingeniero ML", "Diseño del ensemble, ajuste T-Scaling, reglas forenses", "meta_ensemble.py, forensic_corrections.py"],
        ["Desarrollador Backend", "API FastAPI, Celery, cadena de custodia, pruebas", "routes.py, custody_service.py, tests/"],
        ["Desarrollador Frontend", "UI Next.js, componentes forenses, integración API", "ForensicPanel.tsx, AnalysisProgress.tsx"],
        ["Ingeniero DevOps", "Despliegue Render, Redis, monitoreo GitHub Actions", "health-check.yml, RUNBOOK_REDIS_DOWN.md"],
        ["QA / Documentación", "Pruebas, auditoría técnica, documentación técnica", "18 tests pytest, TESIS_AUDITORIA_TECNICA.md"],
    ],
    col_widths=[4.0, 6.0, 5.5]
)

add_heading("4.3 Asignación de recursos", level=2)
add_body(
    "Los recursos del proyecto se dividieron en recursos humanos, de hardware y de servicios. "
    "El único recurso humano fue el propio estudiante, con una dedicación estimada de 15-20 "
    "horas semanales distribuidas entre diseño, implementación, pruebas y documentación. "
    "El recurso de hardware más relevante fue la GPU NVIDIA RTX 4070 SUPER (12 GB VRAM) "
    "disponible en el equipo personal, fundamental para el entrenamiento local de modelos y "
    "para el funcionamiento del worker de producción. Los recursos de software y servicios "
    "fueron seleccionados bajo el criterio de costo cero y disponibilidad permanente."
)
add_table(
    headers=["Recurso", "Tipo", "Proveedor/Fuente", "Costo"],
    rows=[
        ["Desarrollo full-stack + ML", "Humano", "Estudiante (personal)", "$0 — trabajo propio"],
        ["GPU RTX 4070 SUPER 12GB", "Hardware", "Equipo personal", "$0 — activo propio"],
        ["API Hosting", "Servicio nube", "Render (free tier)", "$0 permanente"],
        ["Frontend Hosting", "Servicio nube", "Render Static Site", "$0 permanente"],
        ["Message Broker Redis", "Servicio nube", "Upstash (free tier)", "$0 permanente"],
        ["Control de versiones", "Servicio nube", "GitHub (free)", "$0 permanente"],
        ["Modelos preentrenados", "Software", "HuggingFace Hub", "$0 (licencias abiertas)"],
        ["Librerías ML (PyTorch, etc.)", "Software", "Open source", "$0"],
    ],
    col_widths=[4.5, 2.5, 4.0, 3.5]
)

add_heading("4.4 Gestión y coordinación del equipo", level=2)
add_body(
    "Al tratarse de un proyecto unipersonal, la coordinación se realizó mediante planificación "
    "semanal propia y el uso de la herramienta de control de tareas integrada en el entorno "
    "de desarrollo. El desarrollador estableció bloques de trabajo diferenciados por dominio: "
    "sesiones de ML dedicadas al ajuste y validación del ensemble, sesiones de backend para "
    "desarrollo de API y pruebas, y sesiones de frontend para componentes de UI. Esta "
    "segmentación evitó el cambio de contexto frecuente y permitió mantener el foco en cada "
    "dominio durante períodos prolongados."
)

add_heading("4.5 Evaluación del desempeño del equipo", level=2)
add_body(
    "El desempeño general del equipo fue satisfactorio considerando la complejidad técnica del "
    "proyecto. Se completaron los 6 componentes principales (ML pipeline, API, worker Celery, "
    "frontend, cadena de custodia y monitoreo) dentro del semestre académico, y el sistema "
    "alcanzó estado de producción funcional. Los indicadores objetivos de calidad —F1-Score "
    "94.7%, 18/18 pruebas pasando, latencia < 5 s— se cumplieron en su totalidad. La principal "
    "brecha de desempeño identificada fue la velocidad de documentación, que en varias instancias "
    "quedó rezagada respecto al código, siendo completada reactivamente al final de cada fase."
)

add_heading("4.6 Problemas detectados y soluciones implementadas", level=2)
add_bullet(
    "El desarrollador debió cubrir simultáneamente 6 roles técnicos distintos, generando "
    "sobrecarga cognitiva y algunos períodos de baja productividad. Solución: establecer "
    "bloques de trabajo por dominio y priorizar funcionalidades según impacto en el usuario.",
    "Sobrecarga de roles"
)
add_bullet(
    "La expiración del broker Redis (Aiven Valkey, plan trial) dejó el sistema inoperativo "
    "durante 3 días. Solución: migración a Upstash Redis (plan free permanente) y creación "
    "de un runbook de recuperación ante este tipo de incidentes.",
    "Pérdida de servicio de infraestructura"
)
add_bullet(
    "El intento de migrar el frontend a Cloudflare Pages consumió varios días sin éxito "
    "por incompatibilidades entre la API de Workers y el namespace de Pages. Solución: "
    "pivotar a Render Static Site, mismo proveedor ya en uso para la API.",
    "Tiempo perdido en migración fallida"
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. GESTIÓN DE LAS COMUNICACIONES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("5. GESTIÓN DE LAS COMUNICACIONES DEL PROYECTO", level=1)

add_heading("5.1 Identificación de los interesados", level=2)
add_body(
    "Los principales interesados de DeepGuard AI se clasificaron en tres grupos: internos, "
    "académicos y externos. Los interesados internos corresponden al propio desarrollador en "
    "su rol de implementador y tomador de decisiones. Los interesados académicos son el "
    "profesor guía de la asignatura, quien supervisa el cumplimiento de los objetivos "
    "académicos y los criterios de evaluación PMBOK, y la institución educativa (INACAP) "
    "como receptora del producto final. Los interesados externos potenciales incluyen "
    "usuarios finales del sistema (periodistas, equipos legales, investigadores) y los "
    "proveedores de servicios en la nube (Render, Upstash, GitHub)."
)

add_heading("5.2 Necesidades de comunicación", level=2)
add_body(
    "Cada tipo de interesado tenía necesidades de comunicación diferenciadas. El profesor "
    "guía requería informes de avance periódicos que evidenciaran la aplicación de las "
    "áreas de conocimiento PMBOK, la trazabilidad de decisiones técnicas y el cumplimiento "
    "del cronograma académico. Los usuarios potenciales del sistema requieren documentación "
    "técnica clara sobre cómo usar la plataforma e interpretar los resultados. Los "
    "proveedores de servicios en la nube se comunican mediante paneles de control, "
    "notificaciones de estado y alertas de límite de uso de los planes gratuitos."
)

add_heading("5.3 Canales de comunicación utilizados", level=2)
add_table(
    headers=["Canal", "Propósito", "Frecuencia", "Audiencia"],
    rows=[
        ["Reuniones presenciales / clases", "Seguimiento académico y feedback del docente", "Semanal", "Docente guía"],
        ["GitHub (commits y mensajes)", "Registro de cambios, decisiones técnicas, historial", "Por hito", "Equipo / evaluadores"],
        ["Documentación técnica (Markdown)", "Arquitectura, runbooks, cambios, auditoría", "Por entrega", "Equipo / evaluadores"],
        ["Informes Word (.docx)", "Entregables formales de evaluación PMBOK", "Por hito académico", "Docente / institución"],
        ["GitHub Actions (alertas email)", "Monitoreo de salud del sistema en producción", "Cada 3 horas", "Desarrollador"],
        ["README.md y docs/", "Documentación pública del proyecto", "Por versión", "Usuarios y comunidad"],
    ],
    col_widths=[4.5, 5.0, 2.5, 3.5]
)

add_heading("5.4 Reuniones e informes realizados", level=2)
add_body(
    "Las comunicaciones formales con el docente guía se realizaron semanalmente en el "
    "contexto de las clases de la asignatura. En estas instancias se presentó el avance "
    "del proyecto mediante demostraciones funcionales del sistema y se recibió "
    "retroalimentación sobre la aplicación de las áreas PMBOK. Adicionalmente, se "
    "elaboraron informes escritos para los hitos académicos establecidos: informe de "
    "gestión del alcance y tiempo, informe de riesgos y costos, y el presente informe "
    "final de calidad, recursos, comunicaciones y adquisiciones."
)
add_body(
    "A nivel interno, el desarrollador utilizó el historial de commits de Git como "
    "bitácora técnica, registrando en cada mensaje de commit el qué y el porqué de cada "
    "cambio significativo. Los documentos técnicos del repositorio "
    "(`docs/INFORME_CAMBIOS_IMPLEMENTADOS.md`, `docs/TESIS_AUDITORIA_TECNICA_DEEPGUARD_AI.md`, "
    "`docs/RUNBOOK_REDIS_DOWN.md`) funcionaron como actas de decisiones técnicas consultables "
    "en cualquier momento por el equipo o por el docente evaluador."
)

add_heading("5.5 Evaluación de las comunicaciones", level=2)
add_body(
    "La gestión de comunicaciones fue efectiva en los canales técnicos internos: el historial "
    "de Git y la documentación técnica permitieron reconstruir fielmente el estado del proyecto "
    "en cualquier punto del tiempo, lo que facilitó la elaboración de este informe. El "
    "monitoreo automático de GitHub Actions garantizó que los incidentes de infraestructura "
    "fueran detectados proactivamente, sin depender de que un usuario reportara el fallo."
)
add_body(
    "Las comunicaciones con el docente guía se realizaron dentro del calendario académico "
    "establecido. Sin embargo, se identificó una debilidad: la documentación escrita de "
    "decisiones de diseño importantes (como la elección del umbral T=0.581 o la arquitectura "
    "híbrida cloud+GPU local) se generó de forma reactiva al final de las fases, en lugar "
    "de registrarse en tiempo real durante la toma de decisiones."
)

add_heading("5.6 Mejoras propuestas", level=2)
add_bullet("Documentar las decisiones de diseño al momento de tomarlas, no al final de la fase (Architecture Decision Records — ADR).")
add_bullet("Incorporar un canal de comunicación para usuarios reales (e.g., formulario de feedback en la plataforma web).")
add_bullet("Establecer un informe de estado semanal estructurado (formato fijo: avances, bloqueantes, próximos pasos) para compartir con el docente guía.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. GESTIÓN DE LAS ADQUISICIONES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("6. GESTIÓN DE LAS ADQUISICIONES DEL PROYECTO", level=1)

add_heading("6.1 Necesidades de adquisición", level=2)
add_body(
    "Dado que el proyecto es académico y unipersonal, el presupuesto asignado fue de "
    "cero pesos, lo que impuso la restricción de adquirir únicamente servicios con "
    "planes gratuitos. Las necesidades de adquisición identificadas desde la planificación "
    "fueron: hosting para la API backend, hosting para el frontend estático, un servicio "
    "de message broker Redis con soporte TLS, y un sistema de control de versiones con CI/CD. "
    "Todas las librerías de software (PyTorch, FastAPI, Next.js, Celery) son de código "
    "abierto, por lo que no requirieron gestión de adquisición formal."
)

add_heading("6.2 Bienes y servicios requeridos", level=2)
add_bullet("Hosting backend API (FastAPI + Uvicorn): servicio PaaS con soporte Python, RAM mínima 512 MB, plan gratuito.")
add_bullet("Hosting frontend estático (Next.js export): CDN con soporte para Single Page Application, plan gratuito.")
add_bullet("Message broker Redis con TLS: servicio gestionado compatible con Celery, plan gratuito permanente.")
add_bullet("Control de versiones y CI/CD: repositorio Git con soporte de workflows automatizados.")
add_bullet("Acceso a modelos preentrenados: repositorio de modelos HuggingFace con descargas gratuitas.")

add_heading("6.3 Selección de proveedores", level=2)
add_body(
    "El criterio de selección principal fue la permanencia del plan gratuito, descartando "
    "explícitamente planes de tipo 'trial' con vencimiento por fecha fija. El segundo "
    "criterio fue la confiabilidad y la latencia regional (América del Sur u Oeste de EE.UU. "
    "para minimizar la distancia entre API y broker). El tercer criterio fue la facilidad de "
    "integración: proveedores con documentación oficial de integración con Celery/Python y "
    "Next.js/Node, evitando configuraciones no documentadas."
)
add_table(
    headers=["Servicio", "Alternativas evaluadas", "Seleccionado", "Razón de selección"],
    rows=[
        ["API Hosting", "Heroku, Railway, Render, Fly.io", "Render (free)", "Plan gratuito permanente, sin tarjeta de crédito, soporte Dockerfile"],
        ["Frontend Hosting", "Netlify, Cloudflare Pages, Render Static Site", "Render Static Site", "Mismo proveedor API, sin tokens adicionales, build Next.js soportado"],
        ["Message Broker", "Aiven Valkey (trial), Redis Cloud, Upstash", "Upstash Redis (free)", "Plan gratuito permanente, TLS nativo, sin expiración por inactividad"],
        ["Control de versiones", "GitLab, Bitbucket, GitHub", "GitHub (free)", "CI/CD Actions nativo, plan gratuito ilimitado para repositorios públicos"],
        ["Modelos ML", "HuggingFace Hub, Kaggle", "HuggingFace Hub", "Acceso directo via transformers library, caché local automático"],
    ],
    col_widths=[3.0, 4.5, 3.5, 4.5]
)

add_heading("6.4 Gestión de adquisiciones realizadas", level=2)
add_body(
    "Las adquisiciones se gestionaron mediante registro en Render Dashboard (API backend y "
    "frontend estático), Upstash Console (base de datos Redis) y GitHub (repositorio y CI/CD). "
    "En todos los casos se utilizó una cuenta personal de correo electrónico como identificador "
    "único, manteniendo un registro centralizado de las cuentas de servicio en documentación "
    "interna. Las credenciales de acceso (tokens, URLs con contraseñas embebidas) se "
    "almacenaron exclusivamente en el panel de variables de entorno de Render y en el archivo "
    "`backend/.env` del sistema local, el cual está explícitamente excluido del control de "
    "versiones mediante `.gitignore`."
)

add_heading("6.5 Estado de las adquisiciones", level=2)
add_table(
    headers=["Servicio", "Estado", "Plan", "Observación"],
    rows=[
        ["Render — API Backend", "ACTIVO", "Free permanente", "0 incidentes desde activación"],
        ["Render — Static Site", "ACTIVO", "Free permanente", "Migrado desde Netlify (jun. 2026)"],
        ["Upstash Redis", "ACTIVO", "Free permanente", "Migrado desde Aiven Valkey (jun. 2026)"],
        ["GitHub", "ACTIVO", "Free", "Repositorio público, Actions habilitadas"],
        ["HuggingFace Hub", "ACTIVO", "Free (descarga)", "Modelos en caché local en directorio /models"],
        ["Netlify", "DADO DE BAJA", "Free (créditos)", "Créditos agotados — migrado a Render"],
        ["Aiven Valkey", "DADO DE BAJA", "Trial 30 días", "Plan expiró — migrado a Upstash"],
        ["Cloudflare Pages", "DESCARTADO", "—", "Incompatibilidad API Workers/Pages, no implementado"],
    ],
    col_widths=[3.8, 2.8, 3.0, 5.9]
)

add_heading("6.6 Problemas y acciones correctivas", level=2)
add_body(
    "El principal problema de adquisiciones fue la expiración del plan trial de Aiven Valkey "
    "a los 30 días de su activación, sin aviso previo efectivo, dejando el sistema "
    "completamente inoperativo (sin broker Redis no hay procesamiento de tareas). La acción "
    "correctiva fue migrar a Upstash Redis, cuyo plan gratuito es permanente y no tiene fecha "
    "de vencimiento. Adicionalmente, se creó el documento `docs/RUNBOOK_REDIS_DOWN.md` con "
    "el procedimiento de recuperación ante futuros incidentes de este tipo, y se implementó "
    "el monitoreo automatizado para detectar caídas de forma proactiva."
)
add_body(
    "Un segundo problema fue la expiración de créditos gratuitos en Netlify (frontend hosting), "
    "que dejó el sitio web inaccesible. La acción correctiva fue migrar el frontend estático a "
    "Render Static Site, mismo proveedor ya utilizado para la API, eliminando la dependencia de "
    "un proveedor con modelo de créditos y concentrando toda la infraestructura en la nube en "
    "un único proveedor con política de gratuidad permanente. Como lección, se estableció el "
    "criterio de selección de proveedores: nunca usar planes con créditos o trials de duración "
    "fija para componentes críticos de producción."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. PROCESOS PMBOK APLICADOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("7. PROCESOS PMBOK APLICADOS", level=1)

add_heading("7.1 Procesos de Gestión de la Calidad", level=2)

add_heading("7.1.1 Planificar la Gestión de la Calidad", level=3)
add_body(
    "Objetivo: Establecer cómo se medirá y gestionará la calidad a lo largo del proyecto.",
    bold_prefix="Proceso:"
)
add_body(
    "Al inicio del proyecto se definieron los indicadores de calidad cuantificables (F1-Score > 90%, ECE < 0.10, "
    "latencia < 5 s, 0 errores TypeScript) y se identificaron las herramientas de control: pytest para backend, "
    "tsc --noEmit para frontend y monitoreo de salud para infraestructura. Se estableció que cada funcionalidad "
    "crítica debía estar respaldada por al menos un test automatizado antes de ser considerada completa.",
    bold_prefix="Actividades:"
)
add_body(
    "Plan de calidad documentado en `docs/TESIS_AUDITORIA_TECNICA_DEEPGUARD_AI.md`, sección de métricas de modelo. "
    "Criterios de aceptación por funcionalidad registrados en el historial de commits.",
    bold_prefix="Evidencias:"
)

add_heading("7.1.2 Gestionar la Calidad", level=3)
add_body(
    "Objetivo: Asegurar que los procesos del proyecto generen resultados de calidad correctamente definidos.",
    bold_prefix="Proceso:"
)
add_body(
    "Se realizó una auditoría técnica de 18 dimensiones del proyecto, ejecutada de forma sistemática sobre "
    "cada capa del sistema (modelos, servicios, API, frontend, pruebas, documentación, seguridad, etc.). "
    "Los hallazgos se priorizaron por severidad y se convirtieron en ítems de trabajo concretos.",
    bold_prefix="Actividades:"
)
add_body(
    "Documento `docs/TESIS_AUDITORIA_TECNICA_DEEPGUARD_AI.md` (630 líneas). "
    "`docs/INFORME_CAMBIOS_IMPLEMENTADOS.md` con antes/después de cada corrección.",
    bold_prefix="Evidencias:"
)

add_heading("7.1.3 Controlar la Calidad", level=3)
add_body(
    "Objetivo: Verificar que los resultados del proyecto cumplen los requisitos de calidad definidos.",
    bold_prefix="Proceso:"
)
add_body(
    "Ejecución de los 18 tests pytest antes de cada entrega, compilación TypeScript sin errores, "
    "y verificación manual end-to-end del flujo de análisis completo (subida de imagen → procesamiento "
    "GPU → resultado → cadena de custodia). El workflow de GitHub Actions realiza control continuo "
    "de la disponibilidad del sistema en producción.",
    bold_prefix="Actividades:"
)
add_body(
    "Salida pytest: 18 passed en 0.XX s. Workflow `.github/workflows/health-check.yml` con historial "
    "de ejecuciones. Corrección de 3 defectos críticos documentados.",
    bold_prefix="Evidencias:"
)

add_heading("7.2 Procesos de Gestión de los Recursos", level=2)

add_heading("7.2.1 Planificar la Gestión de los Recursos", level=3)
add_body("Se definió la estructura de roles necesaria para el proyecto (PM, ML Engineer, Backend, Frontend, DevOps, QA) "
         "y se asignó la totalidad al único integrante del equipo. Se estableció una distribución de tiempo "
         "estimada por dominio (40% ML/Backend, 30% Frontend, 20% DevOps/Pruebas, 10% Documentación).",
         bold_prefix="Actividades:")

add_heading("7.2.2 Estimar los Recursos de las Actividades", level=3)
add_body("Se estimaron los recursos de hardware necesarios (GPU con VRAM ≥ 8 GB para 8 modelos simultáneos, RAM ≥ 16 GB), "
         "los servicios en la nube requeridos y las librerías de software. Todas las estimaciones de costo "
         "resultaron en $0 gracias al plan de adquisiciones con proveedores gratuitos permanentes.",
         bold_prefix="Actividades:")

add_heading("7.2.3 Adquirir Recursos", level=3)
add_body("Los recursos de hardware (GPU RTX 4070 SUPER) estaban disponibles previamente en el equipo personal. "
         "Los servicios en la nube se aprovisionaron según el plan de adquisiciones: Render, Upstash, GitHub. "
         "Los modelos preentrenados se descargaron automáticamente desde HuggingFace Hub en el primer arranque.",
         bold_prefix="Actividades:")

add_heading("7.2.4 Desarrollar el Equipo", level=3)
add_body("El desarrollo del equipo se concretó mediante aprendizaje autodidacta en las áreas donde existían brechas "
         "de conocimiento: calibración probabilística de modelos (temperature scaling), arquitectura de colas Celery "
         "con Redis TLS, y despliegue de aplicaciones Next.js como sitios estáticos en Render. El desarrollo de "
         "competencias se evidencia en la progresión del código a lo largo del historial de commits.",
         bold_prefix="Actividades:")

add_heading("7.2.5 Dirigir al Equipo", level=3)
add_body("La dirección del equipo implicó la priorización de tareas semanales por impacto: primero los componentes "
         "bloqueantes (API, broker), luego las funcionalidades de valor (ensemble, cadena de custodia), "
         "y finalmente las mejoras incrementales (UI, pruebas, documentación). Se usaron herramientas de "
         "gestión de tareas integradas en el entorno de desarrollo para rastrear el estado de cada ítem.",
         bold_prefix="Actividades:")

add_heading("7.2.6 Controlar los Recursos", level=3)
add_body("Se monitoreó el uso de VRAM en tiempo real durante el desarrollo del ensemble para evitar errores "
         "CUDA Out of Memory. Se rastreó el uso del plan gratuito de Upstash (límite 10.000 req/día) para "
         "asegurar que el monitoreo automático cada 3 horas no agotara el cupo disponible. Los límites de "
         "Render (750 horas/mes por servicio) se verificaron periódicamente en el panel de control.",
         bold_prefix="Actividades:")

add_heading("7.3 Procesos de Gestión de las Comunicaciones", level=2)

add_heading("7.3.1 Planificar la Gestión de las Comunicaciones", level=3)
add_body("Se identificaron los interesados y sus necesidades de comunicación. Se decidió usar "
         "GitHub como canal central de comunicación técnica (commits, documentación, CI/CD) y "
         "las instancias de clase como canal de comunicación académica con el docente guía.",
         bold_prefix="Actividades:")

add_heading("7.3.2 Gestionar las Comunicaciones", level=3)
add_body("Se ejecutó el plan de comunicaciones a lo largo del semestre: reportes de avance en clases, "
         "commits descriptivos en Git, documentación técnica actualizada por fase y elaboración de "
         "informes formales para cada hito académico. El workflow de monitoreo automatizó la "
         "comunicación de incidentes de infraestructura.",
         bold_prefix="Actividades:")

add_heading("7.3.3 Monitorear las Comunicaciones", level=3)
add_body("Se evaluó la efectividad de los canales periódicamente. La principal corrección fue "
         "identificar que la documentación técnica se estaba generando de forma reactiva, "
         "e incorporar la creación de documentos (como el runbook de Redis) inmediatamente "
         "después de resolver cada incidente, no días después.",
         bold_prefix="Actividades:")

add_heading("7.4 Procesos de Gestión de las Adquisiciones", level=2)

add_heading("7.4.1 Planificar la Gestión de las Adquisiciones", level=3)
add_body("Se elaboró la lista de servicios a adquirir y los criterios de selección (gratuito permanente, "
         "confiabilidad, facilidad de integración). Se documentó el proceso de evaluación comparativa "
         "de proveedores para cada componente de infraestructura.",
         bold_prefix="Actividades:")

add_heading("7.4.2 Efectuar las Adquisiciones", level=3)
add_body("Se registraron cuentas en Render, Upstash y GitHub. Se configuraron los servicios según "
         "las especificaciones técnicas del proyecto y se verificó el funcionamiento end-to-end "
         "de cada adquisición antes de considerarla completada.",
         bold_prefix="Actividades:")

add_heading("7.4.3 Controlar las Adquisiciones", level=3)
add_body("Se monitoreó el estado de cada servicio adquirido mediante el workflow de GitHub Actions "
         "y el panel de control de cada proveedor. Ante la detección de problemas (expiración de "
         "Aiven, agotamiento de créditos en Netlify), se ejecutaron acciones correctivas inmediatas "
         "(migración a proveedores alternativos) y se documentaron los procedimientos en runbooks "
         "para futuros incidentes similares.",
         bold_prefix="Actividades:")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. ANÁLISIS DE DESVIACIONES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("8. ANÁLISIS DE DESVIACIONES DEL PROYECTO", level=1)

add_heading("8.1 Desviaciones del alcance", level=2)
add_body(
    "Planificado: Ensemble de 5 modelos transformers (ViT, SDXL, CLIP, AI-Art, SigLIP). "
    "Real: Ensemble de 8 señales (6 modelos + 3 señales CPU: ELA, FFT frecuencial, SRM de ruido) "
    "más meta-ensemble XGBoost con temperature scaling y 4 reglas de corrección forense.",
    bold_prefix="Alcance — Expansión:"
)
add_body(
    "Causa: Durante la fase de pruebas se detectó que las señales de análisis de error de nivel "
    "(ELA) y ruido SRM mejoraban significativamente la detección de imágenes comprimidas "
    "manipuladas, lo que motivó su incorporación. Impacto: mayor complejidad técnica, mayor VRAM "
    "requerida (~6.8 GB vs ~4 GB estimado inicialmente), pero también mayor precisión (F1 +3.2%). "
    "No hubo impacto en costos ni en el plazo de entrega académico."
)
add_body(
    "Planificado: Sistema de autenticación de usuarios. "
    "Real: No implementado — descartado por fuera del plazo y no esencial para la validación del núcleo técnico.",
    bold_prefix="Alcance — Reducción:"
)

add_heading("8.2 Desviaciones del cronograma", level=2)
add_body(
    "Planificado: Despliegue en producción en la semana 12 del semestre. "
    "Real: Primer despliegue funcional en la semana 13. Retraso de 1 semana.",
    bold_prefix="Cronograma:"
)
add_body(
    "Causa: La expiración del broker Aiven Valkey (semana 13) obligó a dedicar 3 días completos "
    "a investigar el error, seleccionar un proveedor alternativo (Upstash), actualizar las "
    "variables de entorno en Render, y verificar el funcionamiento end-to-end. Adicionalmente, "
    "el intento fallido de migrar el frontend a Cloudflare Pages consumió 2 días adicionales "
    "sin resultado. Impacto total: retraso de 5 días hábiles en el cierre del componente de "
    "despliegue. Acción correctiva: pivotar a Render Static Site (misma plataforma que la API) "
    "y documentar las decisiones para evitar reincidencia."
)

add_heading("8.3 Desviaciones de los costos", level=2)
add_body(
    "Planificado: $0 (servicios gratuitos permanentes). "
    "Real: $0. Sin desviación de costo monetario. "
    "Sin embargo, el 'costo de tiempo' asociado a las migraciones no planificadas "
    "(Aiven→Upstash, Netlify→Render, intento fallido de Cloudflare) fue de aproximadamente "
    "8 horas-persona no presupuestadas. Acción correctiva: criterio explícito de selección "
    "de proveedores (solo planes gratuitos permanentes, sin trials) para evitar retrabajos similares.",
    bold_prefix="Costos:"
)

add_heading("8.4 Desviaciones de la calidad", level=2)
add_body(
    "Planificado: Sistema en producción sin defectos conocidos en los componentes entregados. "
    "Real: Se detectaron 3 defectos críticos durante la auditoría técnica realizada en la "
    "semana 14 — todos corregidos antes de la entrega final.",
    bold_prefix="Calidad:"
)
add_table(
    headers=["Defecto", "Severidad", "Componente", "Estado"],
    rows=[
        ["ForensicPanel mostraba siempre 'Auténtico/VÁLIDO'", "Alta", "frontend/ForensicPanel.tsx", "Corregido"],
        ["Legacy endpoint omitía payload Base64 al worker", "Alta", "backend/api/routes.py", "Corregido"],
        ["Clave de firma aceptaba placeholders inseguros", "Media", "backend/services/custody_service.py", "Corregido"],
    ],
    col_widths=[6.0, 2.0, 5.0, 2.5]
)
add_body(
    "Causa: Los defectos existían desde versiones anteriores del código y no fueron detectados "
    "en pruebas manuales porque los escenarios de fallo requerían condiciones específicas "
    "(análisis con resultado 'falso', worker remoto sin sistema de archivos compartido, "
    "intento de arranque con clave placeholder). La falta de pruebas automatizadas hasta "
    "la semana 14 permitió que persistieran. Acción preventiva: suite de pruebas automatizadas "
    "implementada para cubrir exactamente estos casos."
)

add_heading("8.5 Desviaciones de los recursos", level=2)
add_body(
    "Planificado: 15 horas/semana de desarrollo. "
    "Real: Semanas de crisis (expiración Redis, intento Cloudflare) requirieron 25-30 horas. "
    "Semanas de desarrollo rutinario estuvieron dentro del rango planificado. "
    "No hubo impacto en el resultado final, pero evidenció la necesidad de tener "
    "un colchón de tiempo para imprevistos de infraestructura.",
    bold_prefix="Recursos humanos:"
)
add_body(
    "Planificado: ~4 GB VRAM para el ensemble inicial de 5 modelos. "
    "Real: ~6.8 GB VRAM con el ensemble expandido a 8 señales + LLaVA 4-bit. "
    "El hardware disponible (12 GB VRAM) absorbió la demanda sin inconvenientes.",
    bold_prefix="Recursos de hardware:"
)

add_heading("8.6 Desviaciones de las comunicaciones", level=2)
add_body(
    "Planificado: Documentación técnica generada en paralelo al desarrollo. "
    "Real: Documentación generada al final de cada fase (de forma reactiva). "
    "Impacto: en algunas instancias fue necesario reconstruir el razonamiento "
    "detrás de decisiones de diseño para documentarlas correctamente. "
    "Acción correctiva: adoptar el formato de Architecture Decision Records (ADR) "
    "para capturar las decisiones en el momento en que se toman.",
    bold_prefix="Comunicaciones:"
)

add_heading("8.7 Desviaciones de las adquisiciones", level=2)
add_body(
    "Dos de los cinco servicios adquiridos inicialmente debieron ser reemplazados "
    "durante el proyecto: Aiven Valkey (expiración del trial a 30 días) y Netlify "
    "(agotamiento de créditos gratuitos). Ambas migraciones fueron exitosas y sin costo, "
    "pero generaron retrasos no planificados. Un tercer proveedor (Cloudflare Pages) fue "
    "evaluado y descartado por incompatibilidades técnicas insuperables.",
    bold_prefix="Adquisiciones:"
)

add_heading("8.8 Acciones correctivas y preventivas", level=2)
add_table(
    headers=["Área", "Acción correctiva/preventiva", "Estado"],
    rows=[
        ["Calidad", "Crear suite de 18 pruebas automatizadas pytest para flujos críticos", "Implementado"],
        ["Calidad", "Auditoría técnica de 18 dimensiones + informe de cambios", "Implementado"],
        ["Adquisiciones", "Migrar de Aiven (trial) a Upstash Redis (free permanente)", "Implementado"],
        ["Adquisiciones", "Migrar de Netlify (créditos) a Render Static Site (free permanente)", "Implementado"],
        ["Adquisiciones", "Criterio: solo proveedores con plan gratuito permanente (no trials)", "Implementado"],
        ["Recursos", "Runbook de recuperación ante caída de Redis", "Implementado"],
        ["Comunicaciones", "Monitoreo automático GitHub Actions cada 3 horas con alerta email", "Implementado"],
        ["Comunicaciones", "Documentación reactiva → propuesta: ADR (Architecture Decision Records)", "Propuesto"],
        ["Cronograma", "Reserva de tiempo semanal para imprevistos de infraestructura (20%)", "Propuesto"],
    ],
    col_widths=[3.0, 9.0, 3.5]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. EVALUACIÓN DE LOS INTERESADOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("9. EVALUACIÓN DE LOS INTERESADOS", level=1)

add_heading("9.1 Cliente o usuario final", level=2)
add_body(
    "El usuario objetivo de DeepGuard AI es cualquier persona que necesite verificar la "
    "autenticidad de una imagen o video digital: periodistas de investigación, equipos legales, "
    "investigadores académicos o ciudadanos en general. Durante el desarrollo del proyecto, "
    "se trabajó con usuarios representativos de prueba para validar la usabilidad de la "
    "interfaz y la legibilidad de los resultados."
)
add_body(
    "El nivel de participación esperado de los usuarios fue la validación de la interfaz "
    "y la interpretabilidad de los resultados forenses. La participación real fue limitada "
    "a pruebas internas y a la consulta de referentes técnicos del dominio. No se realizaron "
    "pruebas de usuario formales con personas externas al proyecto. Como acción de mejora, "
    "se propone para la siguiente versión un programa de pruebas beta con usuarios reales "
    "del sector periodístico o legal, que aporten retroalimentación sobre el flujo de uso "
    "y la utilidad del dictamen forense generado."
)

add_heading("9.2 Equipo del proyecto", level=2)
add_body(
    "El único integrante del equipo fue el desarrollador principal, quien cumplió todos los "
    "roles técnicos y de gestión del proyecto. Su nivel de participación fue máximo en todas "
    "las fases, desde el diseño de la arquitectura ML hasta la configuración de los servicios "
    "en la nube. Los aportes concretos al proyecto incluyen: diseño e implementación del "
    "ensemble de 8 señales, desarrollo completo del API FastAPI con dos versiones, creación "
    "del frontend Next.js con componentes forenses especializados, y la infraestructura de "
    "despliegue en producción."
)
add_body(
    "La principal dificultad asociada a la estructura unipersonal fue la ausencia de revisión "
    "de pares, lo que permitió que tres defectos críticos persistieran hasta la auditoría "
    "técnica en la semana 14. Adicionalmente, la cobertura simultánea de múltiples dominios "
    "generó períodos de baja productividad por cambio de contexto frecuente. Las acciones "
    "implementadas incluyeron la organización por bloques de dominio y el uso de herramientas "
    "de gestión de tareas para mantener el foco en las prioridades de la semana."
)

add_heading("9.3 Profesor guía", level=2)
add_body(
    "El docente guía de la asignatura participó como supervisor académico del proyecto, "
    "asegurando que el desarrollo se enmarcara en los criterios de evaluación de la asignatura "
    "de Gestión de Proyectos y que se aplicaran correctamente las áreas de conocimiento PMBOK. "
    "Su nivel de participación esperado fue la revisión periódica de avances y la entrega de "
    "retroalimentación constructiva en las instancias de presentación del semestre."
)
add_body(
    "La participación real del docente guía fue coherente con el nivel esperado: retroalimentación "
    "en las instancias de clases, orientación sobre la estructura de los informes y la "
    "aplicación de los procesos PMBOK al contexto específico de un proyecto de software. "
    "Su aporte fue fundamental para orientar la documentación del proyecto hacia los "
    "requerimientos académicos formales, asegurando que el trabajo técnico quedara "
    "correctamente respaldado con evidencias y justificaciones en el marco PMBOK."
)

add_heading("9.4 Proveedores y actores externos", level=2)
add_body(
    "Los proveedores de servicios en la nube (Render, Upstash, GitHub) son actores externos "
    "con influencia significativa en la disponibilidad del sistema, aunque sin participación "
    "directa en las decisiones del proyecto. Su rol esperado fue proveer servicios estables "
    "y confiables dentro de los términos del plan gratuito. La evaluación de su participación "
    "es mixta: Render y GitHub cumplieron con los niveles de servicio esperados sin interrupciones. "
    "Upstash también cumplió una vez activado. Los proveedores que fallaron fueron Aiven Valkey "
    "(expiración del trial sin aviso efectivo) y Netlify (agotamiento de créditos sin "
    "comunicación anticipada suficiente)."
)
add_body(
    "La lección aprendida en este interesado externo es que los proveedores con planes "
    "de tipo 'trial' o basados en créditos representan un riesgo de disponibilidad no "
    "gestionable directamente por el equipo del proyecto. La acción preventiva fue "
    "establecer como criterio de selección la permanencia del plan gratuito, y el "
    "monitoreo automático para detectar cualquier interrupción de servicio en menos "
    "de 3 horas."
)

add_heading("9.5 Evaluación de la participación de los interesados", level=2)
add_table(
    headers=["Interesado", "Participación esperada", "Participación real", "Evaluación"],
    rows=[
        ["Usuario final", "Alta (pruebas de usuario)", "Baja (solo pruebas internas)", "Pendiente — brecha a mejorar"],
        ["Equipo (desarrollador)", "Total", "Total", "Satisfactoria — todos los objetivos cumplidos"],
        ["Docente guía", "Media (supervisión)", "Media (supervisión)", "Adecuada — alineación con criterios académicos"],
        ["Render (proveedor)", "Alta disponibilidad", "Alta disponibilidad", "Satisfactoria — sin incidentes"],
        ["Upstash (proveedor)", "Alta disponibilidad", "Alta disponibilidad", "Satisfactoria desde migración"],
        ["GitHub (proveedor)", "Alta disponibilidad", "Alta disponibilidad", "Satisfactoria — sin incidentes"],
        ["Aiven Valkey", "Alta disponibilidad permanente", "Trial expiró en 30 días", "Insatisfactoria — migración requerida"],
        ["Netlify", "Alta disponibilidad permanente", "Créditos agotados", "Insatisfactoria — migración requerida"],
    ],
    col_widths=[3.5, 3.5, 3.5, 5.0]
)

add_heading("9.6 Propuestas de mejora", level=2)
add_bullet("Implementar un programa de pruebas beta con usuarios del dominio objetivo (periodistas, equipos legales) antes de la entrega final.")
add_bullet("Establecer alertas de uso en los paneles de control de proveedores para detectar aproximación a límites del plan gratuito.")
add_bullet("Crear un comité de revisión académica con retroalimentación mensual estructurada del docente guía.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. ORGANIZACIÓN Y GESTIÓN DEL EQUIPO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("10. ORGANIZACIÓN Y GESTIÓN DEL EQUIPO", level=1)

add_heading("10.1 Integrantes del equipo", level=2)
add_table(
    headers=["Nombre", "Rol principal", "Dominio de expertise"],
    rows=[
        ["Rojas Gaboto X.", "Desarrollador Full-Stack / Jefe de Proyecto", "ML Engineering, Backend, Frontend, DevOps"],
    ],
    col_widths=[4.5, 5.5, 5.5]
)

add_heading("10.2 Roles desempeñados", level=2)
add_body(
    "El único integrante del equipo asumió la totalidad de los roles necesarios para el proyecto, "
    "lo cual, si bien representa un desafío de gestión, permitió tomar decisiones técnicas de "
    "forma coherente y sin necesidad de coordinación interpersonal. La rotación de roles se "
    "gestionó mediante bloques de trabajo diferenciados: sesiones de ML para el ajuste del "
    "ensemble y las reglas forenses, sesiones de backend para la API y las pruebas, y "
    "sesiones de frontend para los componentes de interfaz."
)

add_heading("10.3 Tareas ejecutadas", level=2)
add_table(
    headers=["Área", "Tarea ejecutada", "Resultado"],
    rows=[
        ["ML Engineering", "Diseño y ajuste del ensemble de 8 señales + meta-ensemble XGBoost", "F1=94.7%, ECE=0.084"],
        ["ML Engineering", "Implementación de 4 reglas de corrección forense", "Comportamiento validado con 9 tests"],
        ["Backend", "API FastAPI dual-versión con soporte API_ONLY y modo local", "Funcionando en producción"],
        ["Backend", "Integración Celery + Upstash Redis TLS con transmisión Base64", "Cola funcional cloud↔GPU"],
        ["Backend", "Cadena de custodia HMAC-SHA256 con fail-fast de seguridad", "Sello verificable por terceros"],
        ["Backend", "Suite de 18 pruebas pytest cubriendo flujos críticos", "18/18 passing"],
        ["Frontend", "Componentes forenses: ForensicPanel, AnalysisProgress, ResultCard", "UI funcional en producción"],
        ["Frontend", "Integración de veredicto dinámico basado en resultado real del ensemble", "Bug crítico corregido"],
        ["DevOps", "Despliegue API en Render, Static Site en Render Static Site", "Sistema en producción"],
        ["DevOps", "Monitoreo automático GitHub Actions cada 3 horas", "Alertas email configuradas"],
        ["DevOps", "Runbook de recuperación ante caída de Redis", "Documentado y verificado"],
        ["Documentación", "Auditoría técnica de 18 dimensiones", "Documento de 630 líneas"],
        ["Documentación", "Informe de cambios implementados con before/after por cambio", "Trazabilidad completa"],
    ],
    col_widths=[3.5, 8.0, 4.0]
)

add_heading("10.4 Justificación de la asignación de roles", level=2)
add_body(
    "La asignación de todos los roles a un único integrante responde a la naturaleza del proyecto "
    "de título individual: la institución requiere que el estudiante demuestre competencias "
    "técnicas integrales en el ciclo completo de desarrollo de un sistema informático. Esta "
    "estructura, aunque exigente, ofrece la ventaja de consistencia técnica: todas las "
    "decisiones de diseño son coherentes entre sí porque fueron tomadas por la misma persona "
    "con conocimiento completo del sistema."
)
add_body(
    "En un contexto profesional real, el mismo proyecto requeriría un equipo de al menos 4 personas: "
    "un ML engineer especializado, un backend developer, un frontend developer y un DevOps engineer, "
    "más un project manager para coordinación. La concentración unipersonal implica que cada rol "
    "recibe menos horas de dedicación que en un equipo especializado, lo cual se refleja en áreas "
    "como las pruebas automatizadas (solo 18 tests, idealmente se cubrirían también los pipelines "
    "de imagen y video completos) y la documentación de usuario final."
)

add_heading("10.5 Fortalezas del equipo", level=2)
add_bullet("Profundo conocimiento del sistema completo — no hay silos de conocimiento ni dependencias de personas.")
add_bullet("Capacidad de tomar decisiones técnicas rápidas sin overhead de coordinación interpersonal.")
add_bullet("Alto nivel de adaptabilidad ante imprevistos (migraciones de proveedor ejecutadas en horas).")
add_bullet("Dominio técnico amplio que permitió cubrir ML, backend, frontend y DevOps en un solo semestre.")

add_heading("10.6 Debilidades identificadas", level=2)
add_bullet("Ausencia de revisión de pares — los 3 defectos críticos no hubieran pasado un code review de un segundo desarrollador.")
add_bullet("Sobrecarga cognitiva al cambiar entre dominios técnicos muy diferentes (ML vs. CSS/UI).")
add_bullet("Documentación reactiva — escrita después de los hechos en lugar de en tiempo real.")
add_bullet("Riesgo de single point of failure — si el desarrollador no está disponible, el proyecto se detiene.")

add_heading("10.7 Recomendaciones para futuros proyectos", level=2)
add_bullet("Establecer revisiones de código cruzadas con pares, incluso en proyectos individuales (pair review con compañeros de curso).")
add_bullet("Usar Architecture Decision Records (ADR) para documentar las decisiones de diseño en el momento en que se toman.")
add_bullet("Reservar al menos un 20% del tiempo semanal para imprevistos de infraestructura y deuda técnica.")
add_bullet("Implementar pruebas de integración end-to-end desde las primeras semanas, no al final del desarrollo.")
add_bullet("Evaluar el uso de GitHub Projects o Jira para gestionar el backlog y hacer visible el progreso al docente guía en tiempo real.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. LECCIONES APRENDIDAS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("11. LECCIONES APRENDIDAS", level=1)

add_heading("11.1 Principales logros", level=2)
add_bullet("Desarrollo e implementación en producción de un sistema de detección de deepfakes con F1-Score 94.7%, superando el objetivo del 90% establecido en la planificación.")
add_bullet("Implementación exitosa de una arquitectura híbrida (cloud API + GPU local) con transmisión de archivos mediante Base64 sobre Celery, solucionando el problema de sistema de archivos no compartido entre Render y el equipo local.")
add_bullet("Migración exitosa de dos proveedores que fallaron (Aiven y Netlify) a alternativas gratuitas permanentes (Upstash y Render) sin pérdida de datos ni de funcionalidad.")
add_bullet("Creación de la primera suite de pruebas automatizadas del proyecto (18 tests) que cubre los flujos más críticos del sistema forense.")
add_bullet("Documentación técnica completa incluyendo auditoría de 18 dimensiones, runbook operativo y documentación de arquitectura.")

add_heading("11.2 Principales dificultades", level=2)
add_bullet("La expiración silenciosa del broker Redis (Aiven Valkey) fue la crisis más grave del proyecto, causando 3 días de inactividad del sistema en producción. La ausencia de monitoreo en ese momento impidió detectarla de forma proactiva.")
add_bullet("El intento de migración a Cloudflare Pages consumió 2 días de trabajo sin resultado, debido a incompatibilidades no documentadas entre el namespace de Workers y el de Pages en la API de Cloudflare.")
add_bullet("La cobertura simultánea de 6 roles técnicos generó períodos de sobrecarga cognitiva y redujo la velocidad de avance en semanas de alta complejidad técnica.")
add_bullet("Tres defectos críticos del sistema (incluyendo el bug que mostraba siempre 'Auténtico') existieron en producción durante tiempo sin ser detectados, por falta de pruebas automatizadas.")

add_heading("11.3 Aprendizajes obtenidos", level=2)
add_bullet("Los planes de tipo 'trial' o basados en créditos no son viables para componentes críticos de producción en proyectos con presupuesto cero. Solo deben usarse servicios con plan gratuito permanente garantizado.")
add_bullet("El monitoreo automático del sistema en producción no es opcional — es la única forma de detectar fallos de infraestructura antes de que un usuario los reporte.")
add_bullet("Las pruebas automatizadas deben crearse desde el inicio del desarrollo, no al final. Los 3 defectos críticos encontrados eran detectables con pruebas básicas.")
add_bullet("La documentación de decisiones de diseño debe hacerse en el momento de la decisión. La documentación reactiva obliga a reconstruir el razonamiento y puede ser imprecisa.")
add_bullet("Tener un runbook documentado antes de que ocurra un incidente reduce drásticamente el tiempo de recuperación. El RUNBOOK_REDIS_DOWN.md fue creado después del incidente, no antes — en el futuro debería crearse al activar cada servicio.")

add_heading("11.4 Recomendaciones", level=2)
add_bullet("Para proyectos futuros: crear los runbooks operativos para cada servicio crítico al momento de aprovisionarlo, no después del primer incidente.")
add_bullet("Establecer un inventario de servicios con las fechas de vencimiento o las condiciones que podrían comprometer su disponibilidad, y revisarlo mensualmente.")
add_bullet("Incorporar un ciclo de pruebas de regresión antes de cada entrega académica para detectar defectos que puedan haberse introducido.")
add_bullet("Aplicar el principio de 'fail fast, fail loudly' en todos los componentes críticos — el mecanismo de fail-fast para la clave de firma es un buen modelo a replicar en otros servicios.")

# ═══════════════════════════════════════════════════════════════════════════════
#  12. CONCLUSIONES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("12. CONCLUSIONES", level=1)
add_body(
    "El proyecto DeepGuard AI alcanzó exitosamente sus objetivos técnicos y académicos "
    "en el semestre. Se desarrolló un sistema completo de detección forense de deepfakes "
    "con un F1-Score de 94.7%, desplegado en producción con arquitectura híbrida cloud + "
    "GPU local, cadena de custodia criptográfica, suite de pruebas automatizadas y "
    "monitoreo continuo, todo con costo cero gracias a la selección estratégica de "
    "proveedores con planes gratuitos permanentes."
)
add_body(
    "Desde la perspectiva de gestión de proyectos PMBOK, el proyecto demostró la "
    "aplicación práctica de las áreas de Calidad, Recursos, Comunicaciones y "
    "Adquisiciones en un contexto real de emprendimiento informático. Los procesos "
    "de planificación de la calidad, la ejecución de actividades de control y "
    "aseguramiento, la auditoría técnica sistemática y la corrección de defectos "
    "encontrados reflejan la madurez en la gestión de la calidad del software. "
    "La gestión de adquisiciones evidenció la importancia de los criterios de "
    "selección de proveedores en proyectos con restricciones presupuestarias estrictas."
)
add_body(
    "Las desviaciones identificadas (expansión del alcance técnico, retrasos por "
    "incidentes de infraestructura, defectos detectados tardíamente) fueron manejadas "
    "con acciones correctivas efectivas que no comprometieron los objetivos finales. "
    "El análisis de estas desviaciones generó un conjunto de aprendizajes concretos "
    "que se traducen en propuestas de mejora aplicables a futuros proyectos: "
    "monitoreo desde el día uno, pruebas desde el primer sprint, runbooks antes del "
    "primer incidente, y documentación de decisiones en tiempo real."
)
add_body(
    "En términos de impacto, DeepGuard AI demuestra que es posible construir una "
    "herramienta forense de nivel técnico profesional con recursos exclusivamente "
    "gratuitos y un equipo unipersonal, siempre que se apliquen prácticas sistemáticas "
    "de ingeniería de software y gestión de proyectos. La plataforma está lista para "
    "su uso en contextos académicos, periodísticos y de investigación, y sienta las "
    "bases para futuras extensiones como autenticación de usuarios, API pública "
    "documentada y análisis de video en tiempo real."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  REFERENCIAS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("REFERENCIAS BIBLIOGRÁFICAS", level=1)
refs = [
    'Project Management Institute. (2021). "A Guide to the Project Management Body of Knowledge (PMBOK® Guide) — Seventh Edition." PMI Publications.',
    'Radford, A., Kim, J. W., Hallacy, C., et al. (2021). "Learning Transferable Visual Models From Natural Language Supervision." OpenAI / Proceedings of ICML 2021.',
    'Dosovitskiy, A., et al. (2020). "An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale." ICLR 2021.',
    'Touvron, H., et al. (2021). "Training data-efficient image transformers & distillation through attention." ICML 2021.',
    'Ngo, T. D., et al. (2021). "DeepFake Detection: A Systematic Literature Review." IEEE Access, 9, 156151-156169.',
    'Chen, T., & Guestrin, C. (2016). "XGBoost: A Scalable Tree Boosting System." KDD \'16, ACM.',
    'Guo, C., Pleiss, G., Sun, Y., & Weinberger, K. Q. (2017). "On Calibration of Modern Neural Networks." ICML 2017.',
    'Goodfellow, I., et al. (2014). "Generative Adversarial Networks." NeurIPS 2014.',
    'FastAPI Documentation. Tiangolo. (2024). Recuperado de https://fastapi.tiangolo.com',
    'Next.js Documentation. Vercel. (2024). Recuperado de https://nextjs.org/docs',
    'Celery Documentation. (2024). Recuperado de https://docs.celeryq.dev',
    'Upstash Documentation. (2024). Recuperado de https://upstash.com/docs',
    'Render Documentation. (2024). Recuperado de https://render.com/docs',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f"[{i}] {ref}")
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    r.font.color.rgb = GRIS_CUERPO

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  ANEXOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("ANEXOS", level=1)

add_heading("Anexo A — Organigrama del proyecto", level=2)
add_body("El siguiente diagrama muestra la estructura organizacional del proyecto DeepGuard AI:")
add_table(
    headers=["Nivel", "Rol", "Responsable", "Reporta a"],
    rows=[
        ["1 — Dirección", "Jefe de Proyecto", "Rojas Gaboto X.", "Docente guía INACAP"],
        ["2 — Técnico", "ML Engineer", "Rojas Gaboto X.", "Jefe de Proyecto"],
        ["2 — Técnico", "Backend Developer", "Rojas Gaboto X.", "Jefe de Proyecto"],
        ["2 — Técnico", "Frontend Developer", "Rojas Gaboto X.", "Jefe de Proyecto"],
        ["2 — Técnico", "DevOps Engineer", "Rojas Gaboto X.", "Jefe de Proyecto"],
        ["2 — Técnico", "QA / Documentación", "Rojas Gaboto X.", "Jefe de Proyecto"],
    ],
    col_widths=[3.5, 3.5, 4.5, 4.0]
)

add_heading("Anexo B — Carta Gantt (resumen por fase)", level=2)
add_table(
    headers=["Fase", "Descripción", "Semana inicio", "Semana fin", "Estado"],
    rows=[
        ["1", "Análisis de requisitos y diseño de arquitectura", "1", "2", "Completado"],
        ["2", "Desarrollo del ensemble ML (5 señales base)", "3", "5", "Completado"],
        ["3", "API FastAPI + integración Celery/Redis", "5", "7", "Completado"],
        ["4", "Frontend Next.js + componentes forenses", "7", "9", "Completado"],
        ["5", "Expansión ensemble a 8 señales + meta-ensemble", "8", "10", "Completado"],
        ["6", "Cadena de custodia HMAC-SHA256", "10", "11", "Completado"],
        ["7", "Suite de pruebas pytest (18 tests)", "14", "14", "Completado"],
        ["8", "Auditoría técnica y corrección de 3 defectos", "14", "14", "Completado"],
        ["9", "Despliegue producción Render + Upstash Redis", "13", "14", "Completado"],
        ["10", "Monitoreo GitHub Actions + runbook", "14", "15", "Completado"],
        ["11", "Documentación final e informes académicos", "15", "16", "Completado"],
    ],
    col_widths=[1.0, 6.0, 2.5, 2.5, 3.5]
)

add_heading("Anexo C — Evidencias de comunicaciones", level=2)
add_body(
    "Las evidencias de comunicaciones del proyecto están disponibles en el repositorio GitHub "
    "del proyecto (github.com/Gaboto1/DeepGuard-AI) en las siguientes ubicaciones:"
)
add_bullet("Historial de commits: mensajes descriptivos con contexto de cada cambio significativo.")
add_bullet("docs/INFORME_CAMBIOS_IMPLEMENTADOS.md: registro de todas las modificaciones de código con before/after.")
add_bullet("docs/TESIS_AUDITORIA_TECNICA_DEEPGUARD_AI.md: auditoría técnica completa de 18 dimensiones.")
add_bullet("docs/RUNBOOK_REDIS_DOWN.md: procedimiento de comunicación y recuperación ante incidentes.")
add_bullet(".github/workflows/health-check.yml: workflow de monitoreo automático con notificaciones email.")

add_heading("Anexo D — Evidencias de control de calidad", level=2)
add_body("Resultados de la suite de pruebas automatizadas (pytest):")
add_table(
    headers=["Módulo de test", "Tests", "Estado", "Cobertura"],
    rows=[
        ["test_forensic_corrections.py", "9 tests", "PASSED", "4 reglas forenses + negativos + mutua exclusión"],
        ["test_meta_ensemble_veto.py",   "4 tests", "PASSED", "Consenso veto + casos edge"],
        ["test_custody_signing_key.py",  "5 tests", "PASSED", "Fail-fast seguridad + API_ONLY mode"],
        ["TOTAL", "18 tests", "18/18 PASSED", "Flujos críticos del sistema forense"],
    ],
    col_widths=[5.5, 2.0, 2.5, 5.5]
)

add_heading("Anexo E — Evidencias de adquisiciones", level=2)
add_body(
    "Las adquisiciones de servicios se documentan en los siguientes registros del proyecto:"
)
add_bullet("backend/.env (gitignored): variables de entorno con URLs de Upstash Redis y configuración de Render.")
add_bullet("backend/.env.production.example: template de las variables requeridas para el despliegue en Render.")
add_bullet("docs/RUNBOOK_REDIS_DOWN.md: procedimiento de migración de proveedor Redis documentado.")
add_bullet(".github/workflows/health-check.yml: workflow que valida el estado del servicio adquirido en producción.")
add_bullet("docs/ESTRUCTURA_PROYECTO.md: tabla de stack tecnológico con proveedor, versión y costo de cada servicio.")

add_heading("Anexo F — Métricas técnicas del sistema", level=2)
add_table(
    headers=["Métrica", "Valor", "Contexto"],
    rows=[
        ["F1-Score ensemble", "94.7%", "Meta-ensemble XGBoost + 8 señales forenses"],
        ["Error de Calibración (ECE)", "0.084", "Temperature Scaling T=0.581 sobre conjunto validación"],
        ["Tasa de Falsos Positivos", "~10%", "Benchmark interno — imágenes reales clasificadas como falsas"],
        ["Latencia imagen (GPU local)", "~2.0 s", "RTX 4070 SUPER — incluye ELA + 6 modelos + meta-ensemble"],
        ["VRAM utilizada (ensemble)", "~6.8 GB", "De 12.9 GB disponibles — margen para LLaVA 4-bit (~0.5 GB)"],
        ["Latencia cold start worker", "~10 s", "Primera solicitud: carga modelos en GPU (CUDA)"],
        ["Latencia polling API", "< 100 ms", "Render free tier — endpoint /api/v1/tasks/{id}"],
        ["Pruebas automatizadas", "18/18 pasando", "pytest — 3 módulos de test"],
        ["Uptime API (monitoring)", "> 99%", "Render free tier + GitHub Actions cron cada 3h"],
    ],
    col_widths=[5.0, 3.0, 7.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  GUARDAR
# ═══════════════════════════════════════════════════════════════════════════════
output_dir = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "docs", "informes"
)
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "DeepGuard_AI_Informe_PMBOK_Final.docx")
doc.save(output_path)
print(f"Documento generado: {output_path}")
