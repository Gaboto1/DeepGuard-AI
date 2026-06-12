# -*- coding: utf-8 -*-
"""
Generador del Informe Técnico Profesional — DeepGuard AI
Word (.docx) para defensa de título INACAP
Ejecutar: python generar_informe_word.py
"""
import sys
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta corporativa forense ────────────────────────────────────────────────
AZUL_MARINO   = RGBColor(0x06, 0x1A, 0x30)   # Títulos principales
AZUL_MED      = RGBColor(0x1E, 0x63, 0xD4)   # Acentos / sub-títulos
VERDE_OK      = RGBColor(0x1D, 0x7A, 0x45)   # Éxitos / positivos
ROJO_ERR      = RGBColor(0xC4, 0x2B, 0x2B)   # Errores críticos
GRIS_CUERPO   = RGBColor(0x1F, 0x2D, 0x3D)   # Texto cuerpo
GRIS_TABLA    = RGBColor(0xF0, 0xF4, 0xF8)   # Fondo tabla par
BLANCO        = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_OSCURO   = RGBColor(0x2C, 0x3E, 0x50)   # Encabezados columna


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_col_width(table, col_widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_cm):
                cell.width = Cm(col_widths_cm[i])


def cell_para(cell, text, bold=False, size=9.5, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, mono=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Courier New" if mono else "Calibri"
    if color:
        run.font.color.rgb = color


def heading(doc, text, level=1, color=None, sz=None, before=14, after=4, underline=False):
    sizes  = {1: 18, 2: 14, 3: 12, 4: 11}
    colors = {1: AZUL_MARINO, 2: AZUL_MED, 3: AZUL_MED, 4: GRIS_OSCURO}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    r.bold      = True
    r.font.size = Pt(sz or sizes.get(level, 11))
    r.font.color.rgb = color or colors.get(level, AZUL_MARINO)
    r.font.name = "Calibri"
    if underline:
        r.underline = True
    if level == 1:
        pPr   = p._p.get_or_add_pPr()
        pBdr  = OxmlElement("w:pBdr")
        bot   = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), "1E63D4")
        pBdr.append(bot)
        pPr.append(pBdr)
    return p


def body(doc, text, size=10.5, italic=False, color=None, before=0, after=5, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.italic     = italic
    r.bold       = bold
    r.font.name  = "Calibri"
    r.font.color.rgb = color or GRIS_CUERPO


def code_block(doc, lines, title=None):
    if title:
        t = doc.add_paragraph()
        t.paragraph_format.space_before = Pt(6)
        t.paragraph_format.space_after  = Pt(2)
        rt = t.add_run(title)
        rt.bold = True
        rt.font.size = Pt(9)
        rt.font.color.rgb = GRIS_OSCURO
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run(lines if isinstance(lines, str) else "\n".join(lines))
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1E, 0x63, 0xD4)
    # Fondo gris para bloque de código
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EBF0F7")
    pPr.append(shd)


def bullet(doc, text, bold_prefix=None, size=10.5, color=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    if bold_prefix:
        rb = p.add_run(bold_prefix + ": ")
        rb.bold      = True
        rb.font.size = Pt(size)
        rb.font.color.rgb = AZUL_MED
        rb.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.font.color.rgb = color or GRIS_CUERPO
    r.font.name = "Calibri"


def simple_table(doc, headers, rows, col_widths=None,
                 header_color="1E63D4", row_colors=("F0F4F8", "FFFFFF")):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabecera
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell_para(hrow.cells[i], h, bold=True, size=9, color=BLANCO,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(hrow.cells[i], header_color)

    # Filas
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = row_colors[ri % 2]
        for ci, val in enumerate(row_data):
            bold = (ci == 0 and ri < 99)
            col  = GRIS_CUERPO
            if str(val).startswith("$0") or "Gratuito" in str(val):
                col = VERDE_OK
            cell_para(row.cells[ci], str(val), bold=bold, size=9.5,
                      color=col, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_bg(row.cells[ci], bg)

    if col_widths:
        set_col_width(table, col_widths)
    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("─" * 90)
    r.font.size  = Pt(7)
    r.font.color.rgb = RGBColor(0xCC, 0xD6, 0xE0)
    r.font.name  = "Courier New"


# ═══════════════════════════════════════════════════════════════════════════════
#  CONSTRUCCIÓN DEL DOCUMENTO
# ═══════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()

    # Márgenes
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    # ── PIE DE PÁGINA ─────────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_f = fp.add_run("DeepGuard AI v6.0 — Informe Técnico de Arquitectura e Implementación  ·  INACAP  ·  2026")
        r_f.font.size = Pt(8)
        r_f.font.color.rgb = RGBColor(0x7A, 0x88, 0x99)
        r_f.font.name = "Calibri"

    # ══════════════════════════════════════════════════════════════════════════
    #  PORTADA
    # ══════════════════════════════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("DeepGuard AI")
    r.font.size  = Pt(40)
    r.bold       = True
    r.font.color.rgb = AZUL_MARINO
    r.font.name  = "Calibri"

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Sistema Global de Detección de Deepfakes")
    r2.font.size  = Pt(18)
    r2.font.color.rgb = AZUL_MED
    r2.font.name  = "Calibri"

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("INFORME TÉCNICO DE ARQUITECTURA E IMPLEMENTACIÓN")
    rs.font.size  = Pt(13)
    rs.bold       = True
    rs.font.color.rgb = GRIS_OSCURO
    rs.font.name  = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    meta_data = [
        ("Versión del sistema",   "6.0.0  —  Release de Producción"),
        ("Fecha de emisión",      "31 de mayo de 2026"),
        ("Autor",                 "Gabriel Toro Rojas"),
        ("Institución",           "Instituto Nacional de Capacitación Profesional — INACAP"),
        ("Carrera",               "Ingeniería en Informática"),
        ("Clasificación",         "Documento Técnico — Defensa de Título"),
    ]
    for label, val in meta_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(label + ":  ")
        r1.bold      = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = AZUL_MARINO
        r1.font.name = "Calibri"
        r2 = p.add_run(val)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = GRIS_CUERPO
        r2.font.name = "Calibri"

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  TABLA DE CONTENIDOS (manual)
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "Tabla de Contenidos", 1)
    toc_items = [
        ("Capítulo 1", "Resumen Ejecutivo y Justificación de Arquitectura", "3"),
        ("  1.1", "El Problema: Deepfakes en la Era de la IA Generativa", "3"),
        ("  1.2", "Justificación Teórica: Arquitectura Híbrida Desacoplada", "3"),
        ("  1.3", "Modelo de IA: Ensemble de 5 Modelos Especializados", "5"),
        ("Capítulo 2", "Bitácora Forense de Optimización — Errores Críticos", "6"),
        ("  Error A", "404 Page Not Found en Netlify", "6"),
        ("  Error B", "Network Error y Bloqueos CORS", "7"),
        ("  Error C", "Congelamiento en 2% con Errores 404 en el Polling", "9"),
        ("  Error D", "FileNotFoundError en el Worker GPU Local", "11"),
        ("  Error E", "Out of Memory en Render — 512 MB RAM", "12"),
        ("Capítulo 3", "Flujo Operativo Sectorizado y Cadena de Custodia", "14"),
        ("Capítulo 4", "Topología del Repositorio Purgado v6.0", "17"),
        ("Capítulo 5", "Métricas de Rendimiento y Validación", "19"),
        ("Capítulo 6", "Conclusiones de Ingeniería", "20"),
    ]
    for num, title, pag in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        is_chapter = not num.startswith(" ")
        r1 = p.add_run(f"{num}  ")
        r1.bold = is_chapter
        r1.font.size = Pt(10 if is_chapter else 9.5)
        r1.font.color.rgb = AZUL_MARINO if is_chapter else GRIS_CUERPO
        r1.font.name = "Calibri"
        r2 = p.add_run(title)
        r2.bold = is_chapter
        r2.font.size = Pt(10 if is_chapter else 9.5)
        r2.font.color.rgb = AZUL_MARINO if is_chapter else GRIS_CUERPO
        r2.font.name = "Calibri"
        r3 = p.add_run(f"  {'.' * (70 - len(num) - len(title))}  {pag}")
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(0x9A, 0xAA, 0xBB)
        r3.font.name = "Courier New"

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  CAPÍTULO 1 — RESUMEN EJECUTIVO
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "CAPÍTULO 1: RESUMEN EJECUTIVO Y JUSTIFICACIÓN DE ARQUITECTURA", 1)

    heading(doc, "1.1 El Problema: La Amenaza de los Deepfakes en la Era de la IA Generativa", 2)
    body(doc, (
        "El término deepfake designa contenido audiovisual sintético generado o manipulado "
        "mediante redes neuronales profundas —principalmente Generative Adversarial Networks "
        "(GANs) y, más recientemente, modelos de difusión latente como Stable Diffusion, "
        "DALL-E y Midjourney. La proliferación de estas técnicas ha generado una crisis de "
        "autenticidad digital con impactos documentados en tres dominios críticos:"
    ))
    bullet(doc, "Desinformación política: manipulación de declaraciones de figuras públicas", bold_prefix="Político")
    bullet(doc, "Fraude financiero: suplantación de identidad en videollamadas corporativas", bold_prefix="Económico")
    bullet(doc, "Abuso de imagen: generación no consentida de contenido íntimo", bold_prefix="Social")
    body(doc, (
        "Según el World Economic Forum (2024), el volumen de contenido deepfake detectado en "
        "redes sociales se duplica cada seis meses. Los métodos de verificación humana tienen "
        "una tasa de acierto inferior al 50% frente a deepfakes de última generación, "
        "estadísticamente equivalente a un lanzamiento de moneda."
    ))
    body(doc, (
        "DeepGuard AI nace como respuesta a esta problemática, implementando un pipeline de "
        "análisis forense probabilístico capaz de clasificar imágenes y videos con base en "
        "evidencia multi-modal: frecuencia espacial, coherencia semántica, metadatos forenses "
        "y análisis temporal de consistencia entre fotogramas."
    ))

    heading(doc, "1.2 Justificación Teórica: Arquitectura Híbrida Desacoplada Nube-Local", 2)
    heading(doc, "1.2.1 El Dilema del Cómputo GPU en la Nube", 3)
    body(doc, (
        "La inferencia de modelos de Deep Learning a escala media-grande presenta un cuello "
        "de botella computacional determinado por el ancho de banda de memoria de la GPU. "
        "Los proveedores de nube ofrecen capacidad GPU con los siguientes costos aproximados:"
    ))
    simple_table(doc,
        ["Proveedor", "Instancia GPU", "VRAM", "Costo Mensual (USD)"],
        [
            ["AWS", "g4dn.xlarge — Tesla T4", "16 GB", "~$380"],
            ["Google Cloud", "n1-standard-4 + T4", "16 GB", "~$350"],
            ["Azure", "NC6s_v3 — Tesla V100", "16 GB", "~$730"],
            ["Hardware propio RTX 4070 SUPER", "RTX 4070 SUPER", "12 GB", "$0 (amortizado)"],
        ],
        [5, 5, 2.5, 4]
    )
    body(doc, (
        "La naturaleza académica del proyecto impone una restricción presupuestaria de $0 "
        "en infraestructura de cómputo. Esta restricción convierte el uso de GPU en nube "
        "en técnicamente inviable para un ciclo de vida de 12 meses. La solución adoptada "
        "sigue el principio de Separación de Responsabilidades (Separation of Concerns) "
        "a nivel de infraestructura, distribuyendo las capas del sistema según su perfil "
        "de recursos y costo:"
    ))

    heading(doc, "1.2.2 Mapa de Infraestructura Híbrida", 3)
    simple_table(doc,
        ["Capa", "Tecnología", "Servicio", "RAM / Cómputo", "Costo"],
        [
            ["Presentación", "Next.js 14 + React + TypeScript", "Netlify CDN global", "~0 MB (sin servidor)", "Gratuito"],
            ["Orquestación", "FastAPI 0.104 — API_ONLY=true", "Render (free tier)", "< 80 MB / CPU mínima", "Gratuito"],
            ["Transporte", "Aiven Valkey 7.2.4 (TLS 1.3)", "Aiven (free tier)", "Broker + Result Backend", "Gratuito"],
            ["Cómputo GPU", "Celery 5.3 + PyTorch 2.x + CUDA", "Hardware propio", "6.8 GB VRAM / RTX 4070S", "$0 (amortizado)"],
            ["TOTAL MENSUAL", "", "", "", "$0"],
        ],
        [3, 5, 3.5, 4.5, 3]
    )

    body(doc, "Ventajas arquitectónicas cuantificables de la decisión adoptada:", bold=True)
    bullet(doc, "Costo operativo de $0/mes frente a los $350-730/mes de alternativas GPU cloud", bold_prefix="Costo")
    bullet(doc, "Latencia de inferencia reducida: la GPU local elimina el round-trip de red en el pipeline de inferencia", bold_prefix="Latencia")
    bullet(doc, "Sin restricciones de memoria de contenedor cloud; posibilidad de cargar LLaVA-1.5-7b (7B parámetros, ~5 GB en 4-bit)", bold_prefix="Flexibilidad")
    bullet(doc, "La caída de la API en Render no interrumpe tareas ya encoladas; el worker local opera de forma autónoma", bold_prefix="Resiliencia")

    heading(doc, "1.3 Modelo de IA: Ensemble de 5 Modelos Especializados", 2)
    body(doc, (
        "El corazón del sistema es un meta-ensemble XGBoost que combina las predicciones "
        "de cinco modelos de Deep Learning, cada uno especializado en un subtipo de manipulación:"
    ))
    simple_table(doc,
        ["ID", "Modelo HuggingFace", "Arquitectura", "Especialización"],
        [
            ["A", "prithivMLmods/Deep-Fake-Detector-v2-Model", "ViT (Vision Transformer)", "Face-swap, face-reenactment"],
            ["B", "Organika/sdxl-detector", "EfficientNet fine-tuned", "Imágenes SDXL fotorrealistas"],
            ["C", "CLIP ViT-L/14 + Linear Probe", "Contrastive Learning", "Generalización multi-dominio (FF++)"],
            ["D", "haywoodsloan/ai-image-detector-deploy", "Swin Transformer v2", "Arte generado por IA"],
            ["E", "prithivMLmods/Deepfake-Detect-Siglip2", "SigLIP", "Deepfakes audiovisuales"],
        ],
        [0.8, 6.5, 4, 5]
    )
    body(doc, (
        "El meta-ensemble XGBoost (profundidad=2, regularización L2) opera con pesos "
        "adaptativos según presencia de rostro humano. Se aplica Temperature Scaling "
        "con T=0.581 para calibración probabilística, alcanzando F1=94.7%, ECE=0.084 "
        "y FPR=10% sobre un golden set de 512 imágenes con split independiente."
    ))
    body(doc, (
        "Complementariamente, LLaVA-1.5-7b-hf (cuantizado en 4-bit NF4 mediante "
        "bitsandbytes, ~6.8 GB VRAM) realiza análisis semántico forense evaluando "
        "coherencia anatómica, física y textual, produciendo un risk_score (0-100) "
        "que corrige el output del ensemble en casos límite."
    ))

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  CAPÍTULO 2 — BITÁCORA DE ERRORES
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "CAPÍTULO 2: BITÁCORA FORENSE DE OPTIMIZACIÓN — RESOLUCIÓN DE ERRORES CRÍTICOS", 1)
    body(doc, (
        "Esta sección documenta los cinco errores más críticos encontrados durante el ciclo "
        "de integración del sistema distribuido, con análisis de causa raíz teórica y "
        "solución aplicada a nivel de código. Cada error representa un desafío inherente "
        "a la arquitectura cloud-local desacoplada adoptada."
    ))

    # ── ERROR A ────────────────────────────────────────────────────────────────
    heading(doc, "ERROR A — 404 Page Not Found en Netlify", 2, color=ROJO_ERR)

    heading(doc, "Contexto de manifestación", 3)
    body(doc, (
        "Al desplegar el frontend en Netlify, cualquier acceso a la URL de producción "
        "deepguard-ai-inacap.netlify.app devolvía el error \"Page not found\" de Netlify, "
        "no de nuestra aplicación. La página en blanco impedía cualquier interacción con "
        "el sistema desde la nube."
    ))

    heading(doc, "Causa Raíz Técnica", 3)
    body(doc, (
        "Next.js ofrece tres modos de output: default (SSR con Node.js), standalone "
        "(servidor Node.js autocontenido) y export (HTML/CSS/JS estático puro). "
        "La configuración inicial usaba output: 'standalone', que genera un servidor "
        "Node.js que debe ejecutarse con node server.js. Netlify es un servicio de "
        "hosting estático y no ejecuta procesos Node.js persistentes."
    ))
    body(doc, (
        "Adicionalmente, el archivo _redirects (necesario para que todas las rutas "
        "de la SPA apunten a index.html) estaba ubicado en la raíz del proyecto "
        "frontend, no dentro de la carpeta public/, por lo que Next.js nunca lo "
        "copiaba al directorio de output."
    ))

    heading(doc, "Solución Aplicada", 3)
    code_block(doc, [
        "# next.config.js — ANTES (incompatible con Netlify)",
        "const nextConfig = {",
        "  output: 'standalone',   // Genera servidor Node.js",
        "  async rewrites() { ... }",
        "};",
        "",
        "# next.config.js — DESPUÉS (export estático)",
        "const nextConfig = {",
        "  output: 'export',       // Genera carpeta /out — HTML/CSS/JS puro",
        "  trailingSlash: true,    // /about → /about/index.html",
        "  images: { unoptimized: true },",
        "};",
    ], "Cambio de modo de compilación Next.js")

    code_block(doc, [
        "# netlify.toml — Creado en raíz del frontend",
        "[build]",
        "  command   = \"npm run build\"",
        "  publish   = \"out\"              # Netlify publica esta carpeta",
        "",
        "[build.environment]",
        "  NODE_VERSION        = \"20\"",
        "  NEXT_PUBLIC_API_URL = \"https://deepguard-ai-api.onrender.com\"",
        "",
        "[[redirects]]",
        "  from   = \"/*\"",
        "  to     = \"/index.html\"",
        "  status = 200                   # SPA fallback: todas las rutas → index.html",
    ], "Configuración netlify.toml")

    body(doc, (
        "El archivo _redirects fue reubicado a frontend/public/_redirects para que "
        "Next.js lo copie automáticamente a la carpeta out/ durante el build. "
        "Resultado: Netlify ejecuta npm run build, obtiene la carpeta out/ con "
        "index.html y sirve la SPA correctamente."
    ))
    divider(doc)

    # ── ERROR B ────────────────────────────────────────────────────────────────
    heading(doc, "ERROR B — Network Error y Bloqueos de Seguridad CORS", 2, color=ROJO_ERR)

    heading(doc, "Contexto de manifestación", 3)
    body(doc, (
        "Al intentar analizar una imagen desde deepguard-ai-inacap.netlify.app, "
        "el navegador mostraba en consola: \"Access to fetch at "
        "https://deepguard-ai-api.onrender.com/api/v1/analyze from origin "
        "https://deepguard-ai-inacap.netlify.app has been blocked by CORS policy: "
        "No Access-Control-Allow-Origin header is present on the requested resource.\""
    ))

    heading(doc, "Causa Raíz Técnica", 3)
    body(doc, (
        "CORS (Cross-Origin Resource Sharing) es un mecanismo de seguridad que los "
        "navegadores implementan para restringir peticiones HTTP entre orígenes distintos. "
        "El navegador envía primero una petición preflight con método OPTIONS, esperando "
        "que el servidor responda con la cabecera Access-Control-Allow-Origin. "
        "Dos causas independientes y simultáneas generaban el error:"
    ))
    bullet(doc, "Render no tenía el dominio de Netlify en ALLOWED_ORIGINS_CSV. El middleware CORSMiddleware de FastAPI solo añade cabeceras CORS para orígenes listados explícitamente.", bold_prefix="Causa 1")
    bullet(doc, "NEXT_PUBLIC_API_URL no estaba definida en Netlify. Next.js embebe las variables NEXT_PUBLIC_* en el bundle JavaScript en tiempo de compilación. Sin la variable, el cliente usaba el fallback http://localhost:8000, unreachable desde internet.", bold_prefix="Causa 2")

    heading(doc, "Solución Aplicada", 3)
    code_block(doc, [
        "# backend/app/config.py — CORS ampliado",
        "class Settings(BaseSettings):",
        "    ALLOWED_ORIGINS_CSV: str = (",
        "        \"http://localhost:3000,\"",
        "        \"http://localhost:3001,\"",
        "        \"http://127.0.0.1:3000,\"",
        "        \"https://deepguard-ai-inacap.netlify.app\"  # Producción Netlify",
        "    )",
    ], "Backend — CORS ampliado")

    code_block(doc, [
        "# frontend/Navbar.tsx — mode:'no-cors' para el health check",
        "# Con no-cors: servidor UP → promesa resuelve con respuesta opaca → 'EN LÍNEA'",
        "# Con no-cors: servidor DOWN → promesa rechaza → 'FUERA DE LÍNEA'",
        "const res = await fetch(`${BASE_URL}/api/v1/health`, {",
        "  method: 'GET',",
        "  mode:   'no-cors',",
        "  signal: AbortSignal.timeout(35_000),",
        "  cache:  'no-store',",
        "});",
        "setStatus(res.type === 'opaque' || res.ok ? 'online' : 'offline');",
    ], "Frontend — Badge de estado sin error CORS")

    body(doc, (
        "Error derivado resuelto simultáneamente: el endpoint /api/v1/health importaba "
        "torch incondicionalmente, causando ModuleNotFoundError en Render (API_ONLY=true, "
        "sin PyTorch instalado) → HTTP 500. La importación se hizo condicional:"
    ))
    code_block(doc, [
        "# ANTES: import torch incondicional → ModuleNotFoundError en Render",
        "# DESPUÉS: import condicional según entorno",
        "if not settings.API_ONLY:",
        "    try:",
        "        import torch",
        "        cuda_ok  = torch.cuda.is_available()",
        "        gpu_name = torch.cuda.get_device_name(0) if cuda_ok else 'N/A'",
        "    except ImportError:",
        "        pass",
    ], "Backend — import torch condicional (fix HTTP 500)")
    divider(doc)

    # ── ERROR C ────────────────────────────────────────────────────────────────
    heading(doc, "ERROR C — Congelamiento en 2% con Errores 404 en el Polling", 2, color=ROJO_ERR)

    heading(doc, "Contexto de manifestación", 3)
    body(doc, (
        "El usuario subía un archivo. La barra de progreso avanzaba al 2% y se detenía "
        "indefinidamente. La consola del navegador mostraba GET /api/v1/tasks/{id} 404 "
        "Not Found en bucle indefinido hasta que el usuario cerraba la página."
    ))

    heading(doc, "Causa Raíz Técnica — Cadena de fallos silenciosos", 3)
    body(doc, (
        "El frontend implementa polling asíncrono: consulta GET /api/v1/tasks/{task_id} "
        "cada 1500ms. La función _get_task_result() en el servidor fallaba silenciosamente:"
    ))
    bullet(doc, "Render intentaba leer el estado de la tarea desde Aiven (AsyncResult)", bold_prefix="Paso 1")
    bullet(doc, "Si la conexión a Aiven fallaba (latencia TLS, cold start), el except silenciaba el error", bold_prefix="Paso 2")
    bullet(doc, "Render buscaba el JSON del resultado en su disco local → no existía (lo guarda el worker local)", bold_prefix="Paso 3")
    bullet(doc, "La función retornaba None → el endpoint devolvía HTTP 404", bold_prefix="Paso 4")
    bullet(doc, "Adicionalmente, el timeout de verificación Redis era de 1 segundo, insuficiente para TLS con Aiven (80-200ms de latencia)", bold_prefix="Causa raíz 2")

    heading(doc, "Solución Aplicada", 3)
    code_block(doc, [
        "# backend/app/api/v1/routes.py — Fallback PENDING en lugar de None",
        "def _get_task_result(task_id: str):",
        "    try:",
        "        ar = AsyncResult(task_id, app=celery_app)",
        "        if ar.ready():",
        "            return _enriquecer_estado({'status': 'SUCCESS', 'result': ar.result})",
        "        elif ar.state in ('PROCESSING', 'STARTED'):",
        "            return _enriquecer_estado({'status': ar.state, ...})",
        "        else:",
        "            return _enriquecer_estado({'status': 'PENDING'})",
        "    except Exception as e:",
        "        logger.debug(f'AsyncResult error: {e}')",
        "    # Búsqueda en disco del worker local",
        "    if disk_path.exists():",
        "        return _enriquecer_estado(json.load(open(disk_path)))",
        "    # FALLBACK CRÍTICO: PENDING en lugar de None → HTTP 202 (no 404)",
        "    return _enriquecer_estado({'status': 'PENDING', 'task_id': task_id})",
    ], "Fallback seguro — evita HTTP 404 durante procesamiento")

    code_block(doc, [
        "# Timeout corregido: 1s → 8s para handshake TLS con Aiven",
        "_r = redis.from_url(url, socket_connect_timeout=8, socket_timeout=8)",
    ], "Timeout Redis corregido")
    divider(doc)

    # ── ERROR D ────────────────────────────────────────────────────────────────
    heading(doc, "ERROR D — FileNotFoundError en el Worker GPU Local", 2, color=ROJO_ERR)

    heading(doc, "Contexto de manifestación", 3)
    body(doc, (
        "El worker Celery mostraba: FileNotFoundError: Archivo no encontrado: "
        "uploads/a1beeb50-4660-442f-aff0-275495df0216.webp. El análisis nunca "
        "se completaba y la tarea agotaba los reintentos (max_retries=2) antes "
        "de marcarse como fallida."
    ))

    heading(doc, "Causa Raíz Técnica", 3)
    body(doc, (
        "La arquitectura híbrida genera una disjunción de sistemas de archivos: "
        "el servidor Render (California, USA) y el worker GPU local (Santiago, Chile) "
        "son máquinas físicamente distintas que no comparten ningún volumen de "
        "almacenamiento. El archivo subido por el usuario se guardaba en el disco "
        "efímero de Render (uploads/{task_id}.ext), pero el worker buscaba ese "
        "mismo path en su propio disco local, donde naturalmente no existía."
    ))

    heading(doc, "Solución Aplicada — Serialización Base64 por Canal de Mensajes", 3)
    body(doc, (
        "La solución serializa el contenido binario del archivo como cadena Base64 "
        "dentro del payload del mensaje Celery. Esto transforma el archivo en texto "
        "ASCII transportable por Redis, eliminando la dependencia de sistema de "
        "archivos compartido. El archivo 'viaja' dentro del mensaje de la cola:"
    ))
    code_block(doc, [
        "# backend/app/api/v1/routes.py — API Gateway (Render)",
        "content = await file.read()          # bytes del archivo",
        "import base64",
        "file_content_b64 = base64.b64encode(content).decode('ascii')  # +33% tamaño",
        "",
        "analyze_image_task.apply_async(",
        "    args=[task_id, str(dest), file.filename],",
        "    kwargs={'file_content_b64': file_content_b64},  # viaja en el mensaje",
        "    task_id=task_id,",
        ")",
    ], "API Gateway — Codificación Base64 del archivo")

    code_block(doc, [
        "# backend/app/tasks/analysis_tasks.py — Worker GPU Local",
        "def analyze_image_task(",
        "    self, task_id, file_path, filename,",
        "    *, include_heatmap=True, file_content_b64='',",
        "):",
        "    path = Path(file_path)",
        "    if file_content_b64:",
        "        path.parent.mkdir(parents=True, exist_ok=True)",
        "        path.write_bytes(base64.b64decode(file_content_b64))  # Reconstruir",
        "        logger.info(f'Archivo reconstruido: {len(file_content_b64)//1024} KB')",
        "    if not path.exists():  # Solo falla si Base64 tampoco llegó",
        "        raise FileNotFoundError(f'Archivo no encontrado: {file_path}')",
    ], "Worker GPU — Decodificación y reconstrucción local")
    divider(doc)

    # ── ERROR E ────────────────────────────────────────────────────────────────
    heading(doc, "ERROR E — Out of Memory (OOM Killer) en Render — 512 MB RAM", 2, color=ROJO_ERR)

    heading(doc, "Contexto de manifestación", 3)
    body(doc, (
        "Los despliegues en Render fallaban durante el startup con la señal del "
        "sistema operativo SIGKILL (OOM Killer), seguida de reinicios automáticos "
        "en bucle. Los logs mostraban consumos superiores a 512 MB antes de que "
        "el servidor pudiera responder su primera petición HTTP."
    ))

    heading(doc, "Causa Raíz Técnica", 3)
    body(doc, (
        "El plan gratuito de Render asigna 512 MB de RAM por servicio. La versión "
        "original de main.py cargaba todos los modelos de IA durante el lifespan "
        "de arranque: PyTorch (~500 MB RAM base), 5 modelos ensemble (~6.8 GB VRAM) "
        "y el detector MTCNN. El consumo excedía los 512 MB antes de que FastAPI "
        "pudiera atender peticiones. Adicionalmente, requirements.txt incluía "
        "torch (>2.5 GB descargado), haciendo que el build time excediera los "
        "límites del plan gratuito de Render."
    ))

    heading(doc, "Solución Aplicada — Purga de Librerías e Importación Condicional", 3)
    simple_table(doc,
        ["Métrica", "Antes", "Después"],
        [
            ["RAM en startup", "> 512 MB (OOM → crash)", "< 80 MB"],
            ["Tiempo de build Docker", "> 20 min (timeout)", "~3 min"],
            ["Tiempo de arranque API", "N/A (crash infinito)", "< 10 segundos"],
            ["Tamaño imagen Docker", "~9 GB (con PyTorch+CUDA)", "~185 MB (python:3.13-slim)"],
            ["Dependencias instaladas", "~8 GB (requirements.txt)", "~180 MB (requirements-api.txt)"],
        ],
        [5, 5, 5]
    )
    code_block(doc, [
        "# backend/app/main.py — Lifespan completamente condicional",
        "@asynccontextmanager",
        "async def lifespan(app: FastAPI):",
        "    if settings.API_ONLY:          # Render: API_ONLY=true",
        "        logger.info('Modo API-only: modelos omitidos (workers GPU remotos)')",
        "        # RAM < 80 MB — solo FastAPI corriendo",
        "    else:                          # PC local: carga completa",
        "        import torch",
        "        from app.models.deepfake_detector import DeepfakeDetector",
        "        DeepfakeDetector.get_instance().load()",
        "    yield",
    ], "main.py — Carga condicional de modelos")

    code_block(doc, [
        "# Dockerfile ultra-ligero para Render",
        "FROM python:3.13-slim          # 125 MB (vs pytorch/pytorch:2.1.2-cuda: 8 GB)",
        "ENV API_ONLY=true              # Garantiza modo sin GPU en el contenedor",
        "COPY requirements-api.txt .",
        "RUN pip install --no-cache-dir -r requirements-api.txt",
        "# requirements-api.txt: solo FastAPI + Celery + Redis (sin torch)",
    ], "Dockerfile ultra-ligero — sin PyTorch")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  CAPÍTULO 3 — FLUJO OPERATIVO
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "CAPÍTULO 3: FLUJO OPERATIVO SECTORIZADO Y CADENA DE CUSTODIA", 1)

    body(doc, (
        "A continuación se describe el viaje completo de un archivo, desde su selección "
        "en el navegador hasta la presentación del resultado forense, pasando por todas "
        "las capas de la arquitectura híbrida."
    ))

    fases = [
        ("FASE 0 — CLIENTE (Navegador del usuario)", [
            ("Validación local", "Tipo MIME, tamaño máximo 500 MB, formato soportado"),
            ("Object URL", "Generación de preview local sin subida al servidor"),
            ("Multipart/form-data", "Codificación del archivo para transmisión HTTP"),
        ]),
        ("FASE 1 — API GATEWAY en Render (RAM < 80 MB)", [
            ("Validación 1", "Nombre de archivo requerido (no vacío)"),
            ("Validación 2", "Tipo MIME real verificado con python-magic (no solo extensión)"),
            ("Validación 3", "Tamaño vs MAX_FILE_SIZE_MB — rechaza antes de leer bytes"),
            ("task_id", "Generación UUID v4 único por análisis"),
            ("Base64", "file_content_b64 = base64.b64encode(content)"),
            ("HTTP 202", "Respuesta inmediata con { task_id, poll_url, estimated_seconds }"),
        ]),
        ("FASE 2 — BROKER Aiven Valkey (TLS 1.3)", [
            ("Mensaje Celery", "Serializado en JSON, encriptado TLS sobre puerto 13419"),
            ("Cola destino", "\"images\" con routing key image.analyze"),
            ("Persistencia", "AOF (Append-Only File) — no se pierde si Redis reinicia"),
            ("TTL resultado", "86400 segundos (24 horas) en el result backend"),
        ]),
        ("FASE 3 — WORKER GPU Local (RTX 4070 SUPER)", [
            ("[05%] Reconstrucción", "base64.b64decode → archivo.write_bytes() en disco local"),
            ("[10%] SHA-256", "hashlib.sha256(bytes).hexdigest() — identidad criptográfica"),
            ("[15%] EXIF/XMP/IPTC", "Detección de firmas IA: Midjourney, DALL-E, SDXL, FLUX..."),
            ("[25%] OOD Detection", "5 señales heurísticas — umbral 0.43 + hard rule extreme_pixels"),
            ("[45%] Ensemble GPU", "ViT + SDXL + CLIP + AiArt + SigLIP con pesos adaptativos"),
            ("[65%] XGBoost + T-Scaling", "Meta-ensemble con calibración T=0.581"),
            ("[80%] LLaVA 4-bit GPU", "Análisis semántico: anatomía, física, artefactos, texto"),
            ("[90%] HMAC-SHA256 v2", "Sello criptográfico — cadena de custodia forense"),
            ("[100%] Persistencia", "JSON → disco local + Aiven Redis (result backend)"),
        ]),
        ("FASE 4 — POLLING del Frontend (cada 1500ms)", [
            ("PENDING 202", "Tarea en cola → frontend muestra etapa actual del terminal"),
            ("PROCESSING 202", "Tarea en progreso → barra avanza según meta.progress"),
            ("SUCCESS 200", "Resultado completo JSON → ResultCard renderiza el análisis"),
        ]),
    ]

    for fase_title, pasos in fases:
        heading(doc, fase_title, 3)
        for label, desc in pasos:
            bullet(doc, desc, bold_prefix=label)

    heading(doc, "3.1 Cadena de Custodia Criptográfica", 2)
    body(doc, (
        "El sello de cadena de custodia (versión 2) vincula criptográficamente el "
        "archivo original con el resultado forense, haciendo imposible alterar "
        "retroactivamente cualquier resultado sin invalidar el token. El string "
        "canónico que se firma con HMAC-SHA256 incluye:"
    ))
    code_block(doc, [
        "# Versión 2 — incluye score semántico de LLaVA",
        "seal_version  = 'DG-CUSTODY-v2'",
        "sem_score_val = int(semantic_score) if semantic_score is not None else -1",
        "",
        "canonical = (",
        "    f'{seal_version}:{task_id}:{file_sha256}:'",
        "    f'{final_score:.8f}:{timestamp_unix:.3f}:sem={sem_score_val}'",
        ")",
        "custody_token = hmac.new(",
        "    SIGNING_KEY,            # Variable de entorno — no está en el código",
        "    canonical.encode('utf-8'),",
        "    hashlib.sha256,",
        ").hexdigest()               # 64 caracteres hexadecimales",
    ], "Generación del sello HMAC-SHA256 v2")

    body(doc, (
        "Verificación independiente: cualquier auditor con acceso a SIGNING_KEY puede "
        "recomputar el HMAC y compararlo contra custody_token sin necesidad de acceder "
        "al sistema DeepGuard AI, lo que dota al resultado de validez forense autónoma."
    ))

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  CAPÍTULO 4 — TOPOLOGÍA
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "CAPÍTULO 4: TOPOLOGÍA DEL REPOSITORIO PURGADO v6.0", 1)
    body(doc, (
        "Tras la refactorización de producción v6.0, el repositorio quedó limpio de "
        "caches, archivos temporales y dependencias no versionables. La estructura "
        "refleja una separación estricta por dominio de responsabilidad:"
    ))

    dirs = [
        ("frontend/", "Capa de Presentación — Next.js 14, TypeScript, React"),
        ("  src/app/", "Root layout y página única (SPA)"),
        ("  src/components/", "8 componentes: Navbar, UploadZone, AnalysisProgress, ResultCard, ForensicPanel, MetadataPanel, OsintPanel, HistorySection"),
        ("  src/lib/api.ts", "Cliente Axios con normalización de estados Celery y polling automático"),
        ("  src/types/index.ts", "Interfaces TypeScript: AnalysisResult, SemanticAnalysis, CustodySeal"),
        ("  public/_redirects", "SPA fallback Netlify: /* /index.html 200"),
        ("  netlify.toml", "Build config: publish='out', NEXT_PUBLIC_API_URL embebida"),
        ("  next.config.js", "output:'export', trailingSlash:true, images.unoptimized"),
        ("backend/", "Capa de Orquestación + Worker GPU — Python 3.13"),
        ("  app/main.py", "FastAPI entry point con lifespan condicional (API_ONLY flag)"),
        ("  app/config.py", "Settings Pydantic: API_ONLY, REDIS_URL, CORS, UPLOAD_DIR"),
        ("  app/celery_app.py", "Broker Aiven, 3 colas (images/videos/default), serialización JSON"),
        ("  app/api/v1/routes.py", "Endpoints producción cloud: analyze, tasks, health, history"),
        ("  app/models/", "5 modelos IA + MTCNN + XGBoost + OOD detector (singletons GPU)"),
        ("  app/services/", "Pipelines imagen/video, LLaVA semántico, HMAC custodia, EXIF"),
        ("  app/tasks/analysis_tasks.py", "Celery tasks con base64 payload y retry logic (max_retries=2)"),
        ("  requirements.txt", "Deps completas GPU: torch + transformers + bitsandbytes + opencv"),
        ("  requirements-api.txt", "Deps sin GPU para Render: FastAPI + Celery + Redis (~180 MB)"),
        ("  Dockerfile", "python:3.13-slim, API_ONLY=true, < 185 MB imagen final"),
        ("reports/", "Reportes de análisis (vacío en Git — auto-generado en runtime)"),
        ("  tasks/", "JSON resultados por task_id (chmod 444 — solo lectura)"),
        ("  custody/", "Sellos HMAC-SHA256 para auditoría forense"),
        ("models/", "Pesos de modelos HuggingFace (excluido de Git — varios GB)"),
        (".gitignore", "Excluye: venv/, .env, models/, uploads/, reports/, .next/, out/"),
        ("ESTRUCTURA_PROYECTO.md", "Árbol visual del proyecto para la comisión evaluadora"),
        ("START CELERY WORKER.bat", "Script de arranque del worker GPU con verificación Redis previa"),
    ]

    simple_table(doc,
        ["Ruta", "Descripción"],
        dirs,
        [5, 11.5]
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  CAPÍTULO 5 — MÉTRICAS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "CAPÍTULO 5: MÉTRICAS DE RENDIMIENTO Y VALIDACIÓN", 1)

    heading(doc, "5.1 Rendimiento del Pipeline en Producción", 2)
    simple_table(doc,
        ["Etapa del pipeline", "Tiempo (imagen ~2 MB)", "Hardware"],
        [
            ["Upload HTTP + codificación Base64", "~0.3 s", "Red internet"],
            ["Dispatch Celery (Aiven TLS)", "~0.1 s", "Red internet"],
            ["SHA-256 + EXIF/XMP/IPTC", "~0.05 s", "CPU worker local"],
            ["OOD Detection (5 señales)", "~0.10 s", "CPU worker local"],
            ["Ensemble 5 modelos (GPU)", "~1.20 s", "RTX 4070 SUPER"],
            ["XGBoost Meta-Ensemble", "~0.01 s", "CPU worker local"],
            ["LLaVA Semántico (GPU 4-bit NF4)", "~2.80 s", "RTX 4070 SUPER"],
            ["HMAC-SHA256 Cadena de Custodia", "~0.001 s", "CPU worker local"],
            ["TOTAL PIPELINE", "~4.5 s", "GPU local"],
        ],
        [6.5, 4, 6]
    )

    heading(doc, "5.2 Validación del Modelo — Golden Set 512 Imágenes", 2)
    simple_table(doc,
        ["Métrica", "Valor", "Condición"],
        [
            ["F1-Score", "94.7%", "Split independiente sobre golden set 512 imágenes, 16 categorías"],
            ["ECE (Error de Calibración)", "0.084", "Post Temperature Scaling T=0.581"],
            ["FPR (Falsos Positivos) @ 50%", "10%", "Umbral de clasificación 50%"],
            ["Temperatura T", "0.581", "Parámetro de calibración post-entrenamiento"],
            ["VRAM utilizada total", "6.8 GB / 12.9 GB", "5 modelos + LLaVA 4-bit en RTX 4070 SUPER"],
        ],
        [4.5, 3, 9]
    )

    heading(doc, "5.3 Cobertura de Tipos de Manipulación", 2)
    simple_table(doc,
        ["Categoría de deepfake", "Modelos responsables", "Tasa de detección"],
        [
            ["Face-swap deepfake", "ViT (A) + CLIP (C)", "~96%"],
            ["Face-reenactment", "ViT (A) + SigLIP (E)", "~91%"],
            ["Imágenes SDXL/Stable Diffusion", "SDXL Detector (B)", "~94%"],
            ["Arte IA (Midjourney, Firefly, FLUX)", "AI Art (D) + CLIP (C)", "~89%"],
            ["Fotos reales comprimidas (FP → corregido)", "Compression Veto + LLaVA", "FPR reducido: 68% → 10%"],
            ["Afiches híbridos IA (FN → corregido)", "OOD Bypass + LLaVA", "Detección restaurada por fusión semántica"],
        ],
        [5.5, 4.5, 4.5]
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  CAPÍTULO 6 — CONCLUSIONES
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "CAPÍTULO 6: CONCLUSIONES DE INGENIERÍA", 1)

    heading(doc, "6.1 Logros Técnicos Principales", 2)
    body(doc, (
        "El desarrollo de DeepGuard AI v6.0 demuestra que es factible construir un "
        "sistema de detección de deepfakes de grado profesional con costo operativo "
        "nulo en infraestructura, resolviendo el problema de acceso a cómputo GPU "
        "mediante una arquitectura híbrida desacoplada que separa la capa de "
        "presentación y orquestación (cloud gratuito) de la capa de inferencia "
        "(hardware propio). Los cinco errores documentados en el Capítulo 2 "
        "representan los desafíos inherentes a sistemas distribuidos multi-cloud "
        "con workers locales heterogéneos, y su resolución requirió comprensión "
        "profunda de múltiples dominios de ingeniería simultáneos."
    ))

    heading(doc, "6.2 Principios de Ingeniería Aplicados", 2)
    simple_table(doc,
        ["Principio", "Aplicación en DeepGuard AI"],
        [
            ["Separation of Concerns", "4 capas independientes: presentación, orquestación, transporte, cómputo"],
            ["Defense in Depth", "SHA-256 + HMAC + CORS + Rate Limiting + MIME validation en capas sucesivas"],
            ["Fail Safe Defaults", "Fallback PENDING en polling → frontend no aborta, sistema resiliente"],
            ["Single Responsibility", "API solo valida y despacha; Worker solo infiere — sin mezcla de roles"],
            ["Environment Parity", ".env.production.example y .env.worker.example documentados y versionados"],
            ["Immutable Infrastructure", "Docker imagen reproducible; .gitignore estricto; builds deterministas"],
            ["Graceful Degradation", "Sin LLaVA (VRAM insuficiente): sistema funciona en modo reducido sin crash"],
        ],
        [5, 11.5]
    )

    heading(doc, "6.3 Líneas de Trabajo Futuro", 2)
    bullet(doc, "Certificado TLS de Aiven en Celery: reemplazar ssl_cert_reqs=CERT_NONE por el certificado CA descargado, eliminando el warning de seguridad", bold_prefix="Seguridad")
    bullet(doc, "Object Storage (S3/R2/Cloudflare) para archivos >50 MB: reemplaza la serialización Base64 que incrementa el tamaño en 33%", bold_prefix="Escalabilidad")
    bullet(doc, "Fine-tuning del meta-ensemble: expandir el golden set a 2048+ imágenes con manipulaciones de 2025-2026 (FLUX.1, Ideogram v3, GPT-4o Image)", bold_prefix="Precisión")
    bullet(doc, "Autoscaling del worker: múltiples instancias Celery para análisis paralelo, con soporte de prioridad de cola por tipo y tamaño de archivo", bold_prefix="Rendimiento")
    bullet(doc, "API Key Management: implementar sistema de autenticación JWT para acceso controlado desde clientes externos", bold_prefix="Producción")

    # Firma final
    doc.add_paragraph()
    doc.add_paragraph()
    divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DeepGuard AI v6.0.0  ·  Proyecto de Título — Ingeniería en Informática  ·  INACAP  ·  31 de mayo de 2026")
    r.font.size  = Pt(9)
    r.font.color.rgb = RGBColor(0x7A, 0x88, 0x99)
    r.font.name  = "Calibri"
    r.italic     = True

    # Guardar
    out = r"c:\Users\gabot\OneDrive\Desktop\PROYECTO TITULO FINAL\DeepGuard_AI_Informe_Tecnico.docx"
    doc.save(out)
    sz  = __import__("os").path.getsize(out) / 1024
    print(f"Informe Word generado: {out}")
    print(f"Tamaño: {sz:.1f} KB")


if __name__ == "__main__":
    build()
