# -*- coding: utf-8 -*-
"""
Convierte docs/TESIS_AUDITORIA_TECNICA_DEEPGUARD_AI.md a un documento Word
con portada, indice automatico (campo TOC de Word) y estilos de tesis.
Ejecutar: venv/Scripts/python.exe tools/generar_tesis_word.py
"""
import sys
import re
from pathlib import Path

if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT       = Path(__file__).resolve().parent.parent
MD_PATH    = ROOT / "docs" / "TESIS_AUDITORIA_TECNICA_DEEPGUARD_AI.md"
OUT_PATH   = ROOT / "docs" / "informes" / "DeepGuard_AI_Tesis_Auditoria_Tecnica.docx"

# ── Paleta corporativa forense (misma de generar_informe_word.py) ─────────────
AZUL_MARINO = RGBColor(0x06, 0x1A, 0x30)
AZUL_MED    = RGBColor(0x1E, 0x63, 0xD4)
GRIS_OSCURO = RGBColor(0x2C, 0x3E, 0x50)
GRIS_CUERPO = RGBColor(0x1F, 0x2D, 0x3D)
GRIS_NOTA   = RGBColor(0x5A, 0x6B, 0x7D)
BLANCO      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Regex de parsing markdown ─────────────────────────────────────────────────
HEADER_RE   = re.compile(r"^(#{1,4})\s+(.*)$")
BULLET_RE   = re.compile(r"^-\s+(.*)$")
NUM_RE      = re.compile(r"^(\d+)\.\s+(.*)$")
QUOTE_RE    = re.compile(r"^>\s?(.*)$")
SEP_CELL_RE = re.compile(r"^:?-{1,}:?$")
INLINE_RE   = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")
FENCE       = "```"
LEVEL_MAP   = {2: 1, 3: 2, 4: 3}


# ── Helpers de formato ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def strip_md_inline(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text


def render_inline(paragraph, text, size=10.5, color=None, italic_base=False):
    color = color or GRIS_CUERPO
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            plain = text[pos:m.start()]
            if plain:
                r = paragraph.add_run(plain)
                r.font.size, r.font.name, r.font.color.rgb = Pt(size), "Calibri", color
                r.italic = italic_base
        if m.group(1) is not None:
            r = paragraph.add_run(m.group(1))
            r.bold = True
            r.italic = italic_base
            r.font.size, r.font.name, r.font.color.rgb = Pt(size), "Calibri", AZUL_MARINO
        elif m.group(2) is not None:
            r = paragraph.add_run(m.group(2))
            r.italic = True
            r.font.size, r.font.name, r.font.color.rgb = Pt(size), "Calibri", color
        elif m.group(3) is not None:
            r = paragraph.add_run(m.group(3))
            r.font.name = "Courier New"
            r.font.size = Pt(max(size - 0.5, 7.5))
            r.font.color.rgb = AZUL_MED
        pos = m.end()
    if pos < len(text):
        plain = text[pos:]
        if plain:
            r = paragraph.add_run(plain)
            r.font.size, r.font.name, r.font.color.rgb = Pt(size), "Calibri", color
            r.italic = italic_base


def add_heading(doc, text, md_level, page_break=False):
    if page_break:
        doc.add_page_break()
    level = LEVEL_MAP.get(md_level, 3)
    h = doc.add_heading(strip_md_inline(text), level=level)
    sizes  = {1: 17, 2: 13, 3: 11.5}
    colors = {1: AZUL_MARINO, 2: AZUL_MED, 3: GRIS_OSCURO}
    for r in h.runs:
        r.font.size = Pt(sizes.get(level, 11))
        r.font.color.rgb = colors.get(level, GRIS_OSCURO)
        r.font.name = "Calibri"
        r.bold = True
    h.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    h.paragraph_format.keep_with_next = True
    if level == 1:
        pPr = h._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), "1E63D4")
        pBdr.append(bot)
        pPr.append(pBdr)


def add_paragraph_rich(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    render_inline(p, text, size=10.5)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    render_inline(p, text, size=10.5)


def add_numbered(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    rb = p.add_run(f"{num}.  ")
    rb.bold, rb.font.size, rb.font.color.rgb, rb.font.name = True, Pt(10.5), AZUL_MED, "Calibri"
    render_inline(p, text, size=10.5)


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    render_inline(p, text, size=9.5, color=GRIS_NOTA, italic_base=True)


def add_code_block(doc, lines):
    if not any(l.strip() for l in lines):
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EBF0F7")
    pPr.append(shd)
    for idx, line in enumerate(lines):
        r = p.add_run(line if line.strip() else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x12, 0x3A, 0x6B)
        if idx < len(lines) - 1:
            r.add_break(WD_BREAK.LINE)


def split_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def is_sep_row(cells):
    return bool(cells) and all(SEP_CELL_RE.match(c) for c in cells)


def add_table_block(doc, lines):
    rows = [split_row(l) for l in lines]
    if len(rows) >= 2 and is_sep_row(rows[1]):
        header, data = rows[0], rows[2:]
    else:
        header, data = rows[0], rows[1:]
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(data), cols=n_cols)
    table.style = "Table Grid"

    hrow = table.rows[0]
    for i, htext in enumerate(header):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(strip_md_inline(htext))
        r.bold, r.font.size, r.font.color.rgb, r.font.name = True, Pt(9), BLANCO, "Calibri"
        set_cell_bg(cell, "1E63D4")

    for ri, row_vals in enumerate(data):
        row = table.rows[ri + 1]
        bg = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        for ci in range(n_cols):
            val = row_vals[ci] if ci < len(row_vals) else ""
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            render_inline(p, val, size=9)
            set_cell_bg(cell, bg)
    doc.add_paragraph()


def add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    r_el = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Haga clic derecho sobre esta linea y seleccione \"Actualizar campo\" para generar el indice."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_el.append(fld_begin)
    r_el.append(instr)
    r_el.append(fld_sep)
    r_el.append(placeholder)
    r_el.append(fld_end)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    r_el = run._r
    for kind, text in (("begin", None), ("instr", "PAGE"), ("end", None)):
        if kind == "instr":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        r_el.append(el)


# ── Portada manual ─────────────────────────────────────────────────────────────
def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("DeepGuard AI")
    r.font.size, r.bold, r.font.color.rgb, r.font.name = Pt(38), True, AZUL_MARINO, "Calibri"

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Auditoría Técnica y Memoria de Título")
    r2.font.size, r2.font.color.rgb, r2.font.name = Pt(18), AZUL_MED, "Calibri"

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("DOCUMENTACIÓN TÉCNICA EXHAUSTIVA GENERADA MEDIANTE\nANÁLISIS ESTÁTICO COMPLETO DEL REPOSITORIO")
    rs.bold, rs.font.size, rs.font.color.rgb, rs.font.name = True, Pt(12), GRIS_OSCURO, "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    meta = [
        ("Tipo de documento", "Auditoría técnica integral / Memoria de título"),
        ("Fecha de generación", "24 de junio de 2026"),
        ("Último commit analizado", "45576f1 (2026-06-12)"),
        ("Versión del sistema", "v6.0.0 — Ensemble de 8 señales + XGBoost + LLaVA semántico"),
        ("Líneas de código backend", "~6.091 líneas Python (app/)"),
        ("Archivos fuente analizados", "49 archivos .py / .tsx / .ts"),
        ("Repositorio", "Gaboto1/DeepGuard-AI"),
    ]
    for label, val in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(label + ":  ")
        r1.bold, r1.font.size, r1.font.color.rgb, r1.font.name = True, Pt(10.5), AZUL_MARINO, "Calibri"
        r2 = p.add_run(val)
        r2.font.size, r2.font.color.rgb, r2.font.name = Pt(10.5), GRIS_CUERPO, "Calibri"


# ── Construcción del documento ─────────────────────────────────────────────────
def build():
    md_text = MD_PATH.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    start = 0
    for idx, line in enumerate(lines):
        if line.strip() == "## 1. Resumen Ejecutivo":
            start = idx
            break
    body_lines = lines[start:]

    doc = Document()
    for sec in doc.sections:
        sec.top_margin, sec.bottom_margin = Cm(2.5), Cm(2.5)
        sec.left_margin, sec.right_margin = Cm(3.0), Cm(2.5)

    for section in doc.sections:
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_f = fp.add_run("DeepGuard AI — Auditoría Técnica y Memoria de Título  ·  2026  ·  Página ")
        r_f.font.size, r_f.font.color.rgb, r_f.font.name = Pt(8), RGBColor(0x7A, 0x88, 0x99), "Calibri"
        add_page_number_field(fp)

    build_cover(doc)
    doc.add_page_break()

    idx_title = doc.add_paragraph()
    idx_run = idx_title.add_run("Índice")
    idx_run.bold, idx_run.font.size, idx_run.font.color.rgb, idx_run.font.name = True, Pt(17), AZUL_MARINO, "Calibri"
    idx_title.paragraph_format.space_after = Pt(8)
    add_toc_field(doc)
    doc.add_page_break()

    i = 0
    n = len(body_lines)
    while i < n:
        raw = body_lines[i]
        stripped = raw.rstrip()

        if stripped.strip() == "":
            i += 1
            continue

        if stripped.startswith(FENCE):
            i += 1
            code_lines = []
            while i < n and not body_lines[i].startswith(FENCE):
                code_lines.append(body_lines[i].rstrip())
                i += 1
            i += 1
            add_code_block(doc, code_lines)
            continue

        m = HEADER_RE.match(stripped)
        if m:
            md_level = len(m.group(1))
            title = m.group(2).strip()
            add_heading(doc, title, md_level, page_break=(md_level == 2))
            i += 1
            continue

        if stripped.strip() == "---":
            i += 1
            continue

        if stripped.lstrip().startswith("|"):
            table_lines = []
            while i < n and body_lines[i].rstrip().lstrip().startswith("|"):
                table_lines.append(body_lines[i].rstrip())
                i += 1
            add_table_block(doc, table_lines)
            continue

        qm = QUOTE_RE.match(stripped)
        if qm:
            add_note(doc, qm.group(1))
            i += 1
            continue

        bm = BULLET_RE.match(stripped)
        if bm:
            add_bullet(doc, bm.group(1))
            i += 1
            continue

        nm = NUM_RE.match(stripped)
        if nm:
            add_numbered(doc, nm.group(1), nm.group(2))
            i += 1
            continue

        add_paragraph_rich(doc, stripped)
        i += 1

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    size_kb = OUT_PATH.stat().st_size / 1024
    print(f"Documento Word generado: {OUT_PATH}")
    print(f"Tamano: {size_kb:.1f} KB")


if __name__ == "__main__":
    build()
