# -*- coding: utf-8 -*-
"""
Generador de Informe Técnico Profesional — DeepGuard AI
Ejecutar: python generar_informe.py
"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.oxml.ns as nsmap

TODAY = date.today().strftime("%d de %B de %Y").replace(
    "January","enero").replace("February","febrero").replace("March","marzo"
    ).replace("April","abril").replace("May","mayo").replace("June","junio"
    ).replace("July","julio").replace("August","agosto").replace("September","septiembre"
    ).replace("October","octubre").replace("November","noviembre").replace("December","diciembre")

# ─── Paleta de colores corporativa ───────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x08, 0x2A, 0x4A)   # Encabezados principales
AZUL_MEDIO   = RGBColor(0x1E, 0x63, 0xD4)   # Acentos / subtítulos
GRIS_TEXTO   = RGBColor(0x1F, 0x2D, 0x3D)   # Cuerpo
GRIS_TABLA   = RGBColor(0xF0, 0xF4, 0xF8)   # Fondo tabla par
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
VERDE_OK     = RGBColor(0x1D, 0x7A, 0x45)
ROJO_ALERTA  = RGBColor(0xC4, 0x2B, 0x2B)


# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)

def set_row_height(row, height_cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(height_cm * 567)))
    trPr.append(trH)

def cell_text(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1, color=AZUL_OSCURO, size=None, space_before=12):
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes.get(level, 11))
    run.font.color.rgb = color
    if level == 1:
        # Línea bajo el h1
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),  "single")
        bottom.set(qn("w:sz"),   "6")
        bottom.set(qn("w:space"),"4")
        bottom.set(qn("w:color"),"1E63D4")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def body(doc, text, size=10.5, space_after=6, italic=False, color=GRIS_TEXTO):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    return p

def bullet(doc, text, size=10.5, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ": ")
        r1.bold = True
        r1.font.size = Pt(size)
        r1.font.color.rgb = AZUL_MEDIO
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    r2.font.color.rgb = GRIS_TEXTO

def add_page_break(doc):
    doc.add_page_break()

def simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Encabezado
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell_text(hrow.cells[i], h, bold=True, size=9.5, color=BLANCO,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(hrow.cells[i], "1E63D4")
    set_row_height(hrow, 0.8)
    # Filas
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            bold = (ci == 0)
            cell_text(row.cells[ci], str(val), bold=bold, size=9.5,
                      color=GRIS_TEXTO)
            set_cell_bg(row.cells[ci], bg)
        set_row_height(row, 0.7)
    # Anchos de columna
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ─── DOCUMENTO ───────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Márgenes
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    # ── PORTADA ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("DeepGuard AI")
    r.font.size  = Pt(36)
    r.font.bold  = True
    r.font.color.rgb = AZUL_OSCURO

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Plataforma Forense de Detección de Deepfakes")
    r2.font.size  = Pt(18)
    r2.font.color.rgb = AZUL_MEDIO

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Informe Técnico de Arquitectura y Documentación Completa")
    rs.font.size  = Pt(13)
    rs.font.color.rgb = GRIS_TEXTO
    rs.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    meta = [
        ("Versión del sistema:", "6.0 — Producción"),
        ("Fecha de emisión:",    TODAY),
        ("Clasificación:",       "Confidencial — Proyecto de Título"),
        ("Autor:",               "Gabriel Toro"),
        ("Stack principal:",     "Python 3.13 · FastAPI · Next.js 14 · PyTorch · CUDA 12.4"),
        ("Hardware objetivo:",   "NVIDIA RTX 4070 SUPER (12 GB VRAM)"),
    ]
    for label, val in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "  ")
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = AZUL_OSCURO
        r2 = p.add_run(val)
        r2.font.size = Pt(10)
        r2.font.color.rgb = GRIS_TEXTO

    add_page_break(doc)

    # ── 1. RESUMEN EJECUTIVO ──────────────────────────────────────────────────
    heading(doc, "1. Resumen Ejecutivo", 1)
    body(doc, (
        "DeepGuard AI es una plataforma de análisis forense digital diseñada para detectar "
        "imágenes y videos generados o manipulados mediante inteligencia artificial. El sistema "
        "implementa un pipeline de cinco modelos de aprendizaje profundo organizados en un "
        "meta-ensemble XGBoost con calibración de temperatura, complementado por un Modelo de "
        "Lenguaje y Visión (VLM) de 7 mil millones de parámetros para análisis semántico forense."
    ))
    body(doc, (
        "La plataforma está orientada a casos de uso profesionales: verificación periodística, "
        "peritos forenses digitales, equipos de moderación de contenido y organismos legales "
        "que requieren un resultado probabilístico trazable, con cadena de custodia criptográfica "
        "(SHA-256 + HMAC-SHA256) y explicabilidad completa del proceso de decisión."
    ))
    body(doc, (
        "El sistema alcanza F1=94.7% en el conjunto de validación dorado de 512 imágenes "
        "(ECE=0.084, FPR=10%) manteniendo latencia de análisis de imagen inferior a 5 segundos "
        "en GPU dedicada."
    ))

    heading(doc, "Métricas de rendimiento clave", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Métrica", "Valor", "Condición"],
        [
            ["F1-Score", "94.7%", "Golden set 512 imágenes, split independiente"],
            ["Error de Calibración (ECE)", "0.084", "Temperatura T=0.581"],
            ["Tasa de Falsos Positivos (FPR)", "10%", "Umbral 50%"],
            ["Latencia imagen (GPU)", "< 5s", "RTX 4070 SUPER, incluye LLaVA"],
            ["VRAM total utilizada", "~6.8 GB", "5 modelos + LLaVA 4-bit NF4"],
            ["VRAM disponible restante", "~5.1 GB", "Buffer para análisis concurrente"],
        ],
        [5, 3.5, 7]
    )

    # ── 2. ARQUITECTURA DEL SISTEMA ───────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "2. Arquitectura del Sistema", 1)
    body(doc, (
        "DeepGuard AI implementa una arquitectura de microservicios desacoplada con separación "
        "estricta entre la capa de presentación (Next.js), la capa de orquestación (FastAPI) y "
        "la capa de inferencia (Celery Workers con acceso exclusivo a GPU)."
    ))

    heading(doc, "2.1 Diagrama de componentes", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Capa", "Tecnología", "Responsabilidad"],
        [
            ["Presentación", "Next.js 14 · React · Tailwind CSS · Framer Motion",
             "UI forense, visualización de resultados, polling de tareas"],
            ["API Gateway", "FastAPI · Uvicorn · SlowAPI",
             "Validación, rate limiting, dispatch a Celery, respuestas JSON"],
            ["Cola de tareas", "Celery 5.3 · Redis 3.0",
             "Aislamiento GPU, retry automático, persistencia de resultados"],
            ["Inferencia", "PyTorch · transformers · bitsandbytes",
             "Ensemble de 5 modelos + LLaVA 4-bit en GPU exclusiva"],
            ["Almacenamiento", "Sistema de archivos + Redis",
             "Caché de resultados, cadena de custodia, reportes JSON"],
        ],
        [3.5, 6, 6.5]
    )

    heading(doc, "2.2 Flujo de una solicitud de análisis", 2, AZUL_MEDIO, 12)
    steps = [
        ("1. Recepción", "El usuario carga una imagen o video en el frontend. El archivo se "
         "valida (tipo MIME, tamaño máx. 500 MB) y se transfiere al backend via POST /api/v1/analyze."),
        ("2. SHA-256", "El backend calcula la huella digital SHA-256 del archivo original antes "
         "de cualquier procesamiento, garantizando la inmutabilidad de la evidencia."),
        ("3. Dispatch", "Si Redis está disponible, la tarea se despacha a un Celery Worker GPU "
         "con ID único y se responde 202 Accepted inmediatamente. Sin Redis, el fallback síncrono "
         "procesa la tarea en un hilo de fondo."),
        ("4. Metadatos EXIF/XMP", "El worker extrae metadatos forenses del archivo: cámara, "
         "software de edición, coordenadas GPS, firmas de generadores IA (Midjourney, DALL-E, "
         "SDXL, FLUX, Firefly, Runway, GPT Image, Bing Image Creator)."),
        ("5. OOD Detection", "Se evalúan 5 señales heurísticas para detectar imágenes "
         "fuera de dominio (diseño gráfico, afiches, texto denso) que distorsionan el ensemble."),
        ("6. Ensemble de 5 modelos", "Inferencia en paralelo de los 5 modelos especializados "
         "con pesos adaptativos según presencia de rostro."),
        ("7. Meta-Ensemble XGBoost", "El XGBoost combina los scores individuales con "
         "calibración de temperatura (T=0.581) para producir la probabilidad final."),
        ("8. Correcciones forenses", "Se aplican reglas post-hoc calibradas empíricamente: "
         "OOD Bypass (afiches con rostro sintético) y Compression Veto (fotos reales comprimidas)."),
        ("9. LLaVA Semántico", "LLaVA-1.5-7b-hf analiza la coherencia anatómica y física "
         "de la imagen y puede corregir el score del ensemble en casos límite."),
        ("10. Custodia y persistencia", "Se genera el sello HMAC-SHA256 v2 vinculando "
         "SHA-256 del archivo, score final, score semántico y timestamp UTC. El resultado "
         "JSON se persiste en disco y el frontend lo recibe vía polling."),
    ]
    for label, text in steps:
        bullet(doc, text, bold_prefix=label)

    # ── 3. MODELOS DE IA ──────────────────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "3. Modelos de Inteligencia Artificial", 1)
    body(doc, (
        "El pipeline de detección utiliza cinco modelos especializados, cada uno entrenado "
        "para detectar patrones de manipulación distintos. Se combinan mediante un meta-ensemble "
        "XGBoost con pesos adaptativos según la presencia de rostro humano en la imagen."
    ))

    heading(doc, "3.1 Ensemble principal — 5 modelos", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["ID", "Modelo", "Especialización", "Peso (rostro)", "Peso (sin rostro)"],
        [
            ["A", "prithivMLmods/Deep-Fake-Detector-v2-Model\n(ViT Face-Deepfake)",
             "Deepfakes faciales mediante intercambio de identidad (face-swap, face-reenactment)",
             "30%", "20%"],
            ["B", "Organika/sdxl-detector\n(SDXL Detector)",
             "Imágenes fotorrealistas generadas por Stable Diffusion XL y variantes",
             "20%", "25%"],
            ["C", "CLIP ViT-L/14 + Linear Probe\n(EfficientNet FF++)",
             "Detección generalizada entrenada sobre FaceForensics++ (300 GB)",
             "20%", "20%"],
            ["D", "haywoodsloan/ai-image-detector-deploy\n(AI Art Detector / Swin-v2)",
             "Arte digital y fotografía generada por IA (Midjourney, DALL-E, Firefly)",
             "15%", "20%"],
            ["E", "prithivMLmods/Deepfake-Detect-Siglip2\n(SigLIP Deepfake)",
             "Deepfakes audiovisuales y manipulaciones compuestas",
             "15%", "15%"],
        ],
        [0.8, 5.5, 5.5, 2.2, 2.5]
    )

    heading(doc, "3.2 Meta-Ensemble XGBoost", 2, AZUL_MEDIO, 12)
    body(doc, (
        "El meta-ensemble recibe los scores de los 5 modelos como features y produce una "
        "probabilidad calibrada. Parámetros de entrenamiento:"
    ))
    bullets_meta = [
        ("Algoritmo", "XGBoost con profundidad máxima=2 y regularización L2"),
        ("Features de entrada", "sdxl_detector, ai_art_detector, efficientnet_ffpp, "
         "sdxl_x_aiart (interacción), robust_mean, reliable_std "
         "(ViT y SigLIP excluidos del meta por alta FPR)"),
        ("Calibración de temperatura", "Temperature Scaling con T=0.581 post-entrenamiento"),
        ("Veto por consenso", "Si modelos especializados discrepan fuertemente "
         "(VETO_THRESHOLD=0.25), se aplica un suavizado hacia 50%"),
        ("Conjunto de entrenamiento", "512 imágenes sintéticas balanceadas (16 categorías)"),
    ]
    for label, text in bullets_meta:
        bullet(doc, text, bold_prefix=label)

    heading(doc, "3.3 LLaVA-1.5-7b-hf — Análisis Semántico Forense", 2, AZUL_MEDIO, 12)
    body(doc, (
        "El Modelo de Lenguaje y Visión (VLM) LLaVA-1.5-7b-hf añade una capa de comprensión "
        "semántica que los modelos estadísticos no pueden proporcionar: detección de incoherencias "
        "anatómicas, inconsistencias físicas, artefactos de generación y texto deformado."
    ))
    simple_table(doc,
        ["Parámetro", "Valor"],
        [
            ["Modelo base", "llava-hf/llava-1.5-7b-hf"],
            ["Parámetros", "7 mil millones"],
            ["Cuantización", "4-bit NF4 + bfloat16 (bitsandbytes)"],
            ["Double quantization", "Sí — reduce ~0.4 bits adicionales"],
            ["CPU offload FP32", "Habilitado (llm_int8_enable_fp32_cpu_offload=True)"],
            ["VRAM utilizada", "~5.0 GB (GPU) + overflow CPU RAM si necesario"],
            ["Max memory GPU", "Dinámico: VRAM_total − 7.0 GB reservados para ensemble"],
            ["Max memory CPU", "12 GB RAM del sistema (overflow sin errores)"],
            ["Temperatura de inferencia", "0.05 (respuestas reproducibles y estructuradas)"],
            ["Tokens máximos", "300 (suficiente para JSON + descripción forense)"],
            ["Formato de salida", 'JSON: {"risk_score": 0-100, "semantic_observations": "..."}'],
        ],
        [6, 10]
    )

    heading(doc, "3.4 Fusión Semántica — Reglas de Corrección Calibradas", 2, AZUL_MEDIO, 12)
    body(doc, (
        "La fusión entre el score del ensemble y el análisis LLaVA sigue reglas calibradas "
        "empíricamente (sin reentrenar XGBoost) para máxima transparencia y controlabilidad:"
    ))
    simple_table(doc,
        ["Caso", "Condición", "Acción", "Resultado esperado"],
        [
            ["Corrección ascendente",
             "Ensemble < 35% Y LLaVA > 65%",
             "Delta amplificado × 1.40, mínimo garantizado 40%",
             "Capta manipulaciones sutiles que el ensemble no detectó"],
            ["Corrección descendente",
             "Ensemble > 65% Y LLaVA < 20%",
             "Delta reducido × 0.90, mínimo 15%",
             "Evita falsos positivos por compresión JPEG agresiva"],
            ["Zona de compresión",
             "LLaVA < 10% Y Ensemble ∈ [35%, 55%]",
             "Reducción proporcional, mínimo 12%",
             "Fotografías reales comprimidas en redes sociales"],
            ["Blend ponderado",
             "Resto de casos",
             "70% ensemble + 30% LLaVA (con rostro)",
             "Fusión suave sin corrección"],
        ],
        [3.5, 4, 4, 4.5]
    )
    body(doc, "Resultados calibrados en casos críticos (6/6 tests pasan):", italic=True)
    simple_table(doc,
        ["Caso de prueba", "Ensemble", "LLaVA", "Score final", "Tipo de corrección"],
        [
            ["Foto real comprimida (Messi)", "40%", "8/100",  "19.0%", "semantic_compression_zone"],
            ["Afiche IA híbrido",            "25%", "78/100", "47.3%", "semantic_correction_up"],
            ["Deepfake facial claro",        "85%", "88/100", "85.9%", "semantic_blend"],
            ["LLaVA no disponible",          "42%", "N/A",    "42.0%", "sin cambio (degradado)"],
            ["Imagen real sin cara",         "12%", "10/100", "11.6%", "semantic_blend"],
            ["Corrección descendente",       "72%", "10/100", "55.3%", "semantic_correction_down"],
        ],
        [5, 2.5, 2.5, 2.8, 5]
    )

    # ── 4. DETECTOR OOD ───────────────────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "4. Detector Out-of-Distribution (OOD)", 1)
    body(doc, (
        "El detector OOD identifica imágenes fuera del dominio forense para el que fueron "
        "entrenados los modelos: diseño gráfico, afiches, infografías, texto denso. Estas "
        "imágenes no son deepfakes pero el ensemble podría asignarles scores erróneamente altos."
    ))
    simple_table(doc,
        ["Señal heurística", "Peso", "Descripción"],
        [
            ["edge_density",      "25%", "Densidad de bordes Canny — imágenes vectoriales tienen bordes perfectos"],
            ["extreme_pixels",    "25%", "Fracción de píxeles en valores extremos (0-5 o 250-255)"],
            ["uniform_regions",   "20%", "Regiones de color plano sin textura — típico de diseño digital"],
            ["palette_compact",   "15%", "Paleta de colores reducida — imágenes gráficas usan <50 colores"],
            ["text_regularity",   "15%", "Regularidad espacial de regiones horizontales (texto)"],
        ],
        [4, 2, 10]
    )
    body(doc, (
        "Regla dura: si extreme_pixels ≥ 85%, el score OOD se fuerza a mínimo 0.60 "
        "independientemente de las otras señales. Umbral OOD: 0.43. "
        "Alpha de penalización máximo: 0.35 (mezcla hacia 50%)."
    ), italic=True)

    heading(doc, "4.1 Correcciones Forenses Post-OOD", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Regla", "Condición de activación", "Corrección aplicada"],
        [
            ["OOD Bypass (Regla 1)",
             "Imagen OOD Y (ViT > 75% O AI_Art > 55%)",
             "Score corregido a rango [35%-50%]. Patrón: afiche con rostro sintético real."],
            ["Compression Veto (Regla 2)",
             "AI_Art < 5% Y SDXL < 50%",
             "Score corregido a 10% + AI_Art×40% [clampado 8%-22%]. Patrón: foto real comprimida por JPEG."],
        ],
        [3.5, 6, 6.5]
    )

    # ── 5. PIPELINE DE VIDEO ──────────────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "5. Pipeline de Análisis de Video", 1)
    body(doc, (
        "El análisis de video combina análisis frame-by-frame del ensemble de modelos con "
        "un análisis de consistencia temporal basado en embeddings ViT y flujo óptico Farneback, "
        "diseñado específicamente para detectar deepfakes generados frame a frame."
    ))

    heading(doc, "5.1 Análisis Frame-by-Frame", 2, AZUL_MEDIO, 12)
    bullets_video = [
        ("Extracción", "Hasta 50 frames uniformemente espaciados del video completo"),
        ("Detección de cara", "MTCNN por cada frame — si hay cara, se usa el crop facial como input al ensemble"),
        ("Ensemble adaptativo", "Los 5 modelos procesan cada frame con pesos ajustados según presencia de rostro"),
        ("Frame timeline", "Array de {frame_index, timestamp, manipulation_probability, face_detected} para visualización"),
        ("Score final", "Promedio ponderado de todos los frames con penalización temporal si aplica"),
    ]
    for label, text in bullets_video:
        bullet(doc, text, bold_prefix=label)

    heading(doc, "5.2 Análisis Temporal (ViT-768D + Farneback)", 2, AZUL_MEDIO, 12)
    body(doc, (
        "Los deepfakes generados frame a frame presentan inconsistencias temporales invisibles "
        "al ojo humano pero detectables mediante dos señales complementarias:"
    ))
    simple_table(doc,
        ["Señal", "Método", "Umbral de anomalía"],
        [
            ["Embedding ViT-768D", "Distancia coseno entre embeddings CLS consecutivos. "
             "Videos reales: distancia 0.02-0.18. Deepfake: saltos 0.25-0.60.",
             "Varianza > 0.040 O spike > 0.300"],
            ["Flujo óptico Farneback", "Magnitud del flujo entre regiones faciales consecutivas. "
             "Deepfakes generan irregularidades de textura → varianza anómala.",
             "Varianza > 12.0 O media > 8.0"],
        ],
        [3.5, 9, 4]
    )
    body(doc,
         "Multiplicador de riesgo temporal máximo: +35% al score base del ensemble. "
         "Keyframes a 3 FPS, máximo 90 keyframes.",
         italic=True)

    # ── 6. API REST ───────────────────────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "6. API REST — Endpoints v1", 1)
    body(doc, (
        "Todos los endpoints de análisis se encuentran bajo el prefijo /api/v1/. "
        "La API legacy /api/ se mantiene por compatibilidad pero los clientes deben "
        "usar exclusivamente la versión v1."
    ))

    simple_table(doc,
        ["Método", "Endpoint", "Descripción", "Respuesta"],
        [
            ["POST", "/api/v1/analyze",
             "Subir imagen o video para análisis. Acepta multipart/form-data. "
             "Despacha a Celery si Redis disponible, sino usa sync fallback.",
             "202 Accepted + {task_id, poll_url, estimated_seconds}"],
            ["GET", "/api/v1/tasks/{task_id}",
             "Consultar estado y resultado del análisis. Devuelve status en minúsculas "
             "(pending/processing/completed/failed) compatible con el frontend.",
             "202 si en proceso, 200 con resultado completo si terminó"],
            ["GET", "/api/v1/tasks/{task_id}/custody",
             "Verificar integridad del sello de cadena de custodia. "
             "Recomputa el HMAC y lo compara con el token almacenado.",
             "200 con integrity_valid: true/false"],
            ["GET", "/api/v1/history",
             "Historial de análisis completados (lee JSONs en disco). "
             "Máximo 100 resultados, ordenados por fecha desc.",
             "200 con array de resultados"],
            ["DELETE", "/api/v1/tasks/{task_id}",
             "Eliminar resultado persistido en disco.",
             "200 con mensaje de confirmación"],
            ["GET", "/api/v1/health",
             "Estado completo del sistema: API, Redis, workers Celery, GPU.",
             "200 con estado de todos los componentes"],
            ["GET", "/api/health",
             "Health check simplificado (legacy). Estado de modelos y CUDA.",
             "200 con model_loaded, cuda_available, device"],
        ],
        [1.8, 5, 6.5, 4.5]
    )

    heading(doc, "6.1 Estructura del resultado de análisis (imagen)", 2, AZUL_MEDIO, 12)
    body(doc, "El payload JSON devuelto por GET /api/v1/tasks/{id} para una imagen incluye:")
    simple_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["status", "string", "completed | failed | pending | processing"],
            ["manipulation_probability", "float", "Probabilidad de manipulación en escala 0-100"],
            ["evidence_level", "string", "Very Low / Low / Inconclusive / Moderate / Strong Evidence"],
            ["model_agreement", "string", "High / Medium / Low Consensus"],
            ["ensemble", "object", "Desglose por modelo: score, peso, contribución de cada uno de los 5 modelos"],
            ["semantic_analysis", "object", "risk_score (0-100), semantic_observations, fusion_type, fusion_note"],
            ["chain_of_custody", "object", "seal_version, file_sha256, custody_token (HMAC-SHA256), timestamp_utc"],
            ["forensic_metadata", "object", "ai_generator_detected, editing_software, missing_exif, inconsistencies"],
            ["heatmap", "string|null", "Imagen Grad-CAM++ en base64 (solo si score > 38%)"],
            ["is_ood", "bool", "True si la imagen es out-of-distribution (diseño gráfico, afiche)"],
            ["explanation", "string", "Explicación forense en lenguaje natural en español"],
        ],
        [4, 2.5, 10]
    )

    # ── 7. CADENA DE CUSTODIA ─────────────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "7. Cadena de Custodia Criptográfica", 1)
    body(doc, (
        "Cada análisis genera un Sello de Cadena de Custodia v2 que vincula "
        "criptográficamente la evidencia digital con el resultado forense, haciéndolo "
        "imposible de alterar retroactivamente sin invalidar el token."
    ))

    heading(doc, "7.1 Componentes del sello", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Componente", "Algoritmo", "Propósito"],
        [
            ["Hash del archivo", "SHA-256", "Inmutabilidad de la evidencia original"],
            ["Task UUID", "UUID v4", "Trazabilidad única de la sesión de análisis"],
            ["Score final", "float64", "Resultado auditado del ensemble"],
            ["Score semántico", "int (0-100)", "Score LLaVA integrado en el sello (v2)"],
            ["Timestamp", "Unix epoch ms", "Momento exacto del análisis (UTC)"],
            ["Token HMAC", "HMAC-SHA256", "Firma criptográfica del bundle completo"],
        ],
        [4, 3, 9]
    )

    heading(doc, "7.2 String canónico para verificación", 2, AZUL_MEDIO, 12)
    body(doc, "Formato del string que se firma con HMAC-SHA256:")
    p = doc.add_paragraph()
    run = p.add_run(
        "DG-CUSTODY-v2:{task_id}:{file_sha256}:{final_score:.8f}:{timestamp_unix:.3f}:sem={sem_score}"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = AZUL_MEDIO
    p.paragraph_format.left_indent = Cm(1)
    doc.add_paragraph()

    body(doc, (
        "La verificación es totalmente independiente: cualquier auditor con la clave de firma "
        "puede recomputar el HMAC y compararlo contra el custody_token del reporte, "
        "sin acceso al sistema DeepGuard AI."
    ))

    # ── 8. FRONTEND ───────────────────────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "8. Frontend — Interfaz de Usuario", 1)
    body(doc, (
        "La interfaz está construida con Next.js 14 App Router, Tailwind CSS y Framer Motion. "
        "El diseño sigue una paleta forense profesional inspirada en CrowdStrike y VirusTotal: "
        "fondo oscuro (#080C15), tipografía monospace para valores numéricos y código de colores "
        "semáforo para niveles de evidencia."
    ))

    heading(doc, "8.1 Componentes principales", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Componente", "Archivo", "Función"],
        [
            ["UploadZone", "UploadZone.tsx",
             "Zona de drag-and-drop. Acepta JPG, PNG, WEBP, MP4, MOV, MKV, WEBM (máx. 500 MB). "
             "Inicia el upload y el polling automático."],
            ["ResultCard", "ResultCard.tsx",
             "Tarjeta principal de resultado. Muestra gauge circular animado con el score, "
             "nivel de evidencia, y el panel de desglose de modelos siempre visible."],
            ["ModelBreakdownPanel", "ResultCard.tsx (inline)",
             "Panel siempre visible con barra animada para cada uno de los 5 modelos + LLaVA. "
             "Muestra nombre completo, porcentaje y peso en el ensemble."],
            ["ForensicPanel", "ForensicPanel.tsx",
             "Tab 'Ensemble': tabla completa con Score / Peso / Contribución por modelo, "
             "consenso de modelos, panel LLaVA semántico con observaciones forenses."],
            ["AnalysisProgress", "AnalysisProgress.tsx",
             "Indicador de progreso durante el análisis con etapas descriptivas."],
            ["MetadataPanel", "MetadataPanel.tsx",
             "Visualización de metadatos EXIF/XMP: cámara, software, GPS, firmas de IA."],
            ["HistorySection", "HistorySection.tsx",
             "Historial de los últimos 50 análisis almacenado en localStorage."],
        ],
        [3.5, 4, 9]
    )

    heading(doc, "8.2 Jerarquía visual del resultado", 2, AZUL_MEDIO, 12)
    body(doc, "La pantalla de resultados sigue esta jerarquía de información:")
    hierarchy = [
        ("Score consolidado", "Gauge circular grande con el porcentaje final y nivel de evidencia"),
        ("Desglose por modelo", "Barras animadas de los 5 modelos + LLaVA, siempre visible sin pestañas"),
        ("Fórmula del ensemble", "∑ = peso₁·score₁ + ... = probabilidad_final en tiempo real"),
        ("Tabs de detalle", "Resumen / Ensemble (tabla completa) / Metadatos / Verificación / Heatmap"),
        ("Gráfico temporal", "Solo en videos: AreaChart de probabilidad de manipulación por segundo"),
        ("Cadena de custodia", "Token HMAC y SHA-256 del archivo en el tab Resumen"),
    ]
    for label, text in hierarchy:
        bullet(doc, text, bold_prefix=label)

    heading(doc, "8.3 Comunicación Frontend ↔ Backend", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Función", "Endpoint", "Comportamiento"],
        [
            ["uploadFile()", "POST /api/v1/analyze",
             "Subida multipart con progreso. Retorna task_id para polling."],
            ["getTask()", "GET /api/v1/tasks/{id}",
             "Normaliza status de Celery (SUCCESS→completed, FAILED→failed). "
             "Acepta HTTP 202 para estados intermedios."],
            ["pollTask()", "Intervalo 1.5s sobre getTask()",
             "Polling automático hasta status completed o failed. Cancela con la función retornada."],
            ["getHistory()", "GET /api/v1/history",
             "Lee JSONs en disco del servidor, normaliza la respuesta."],
            ["getHealth()", "GET /api/v1/health",
             "Estado completo (Redis + Celery + GPU) para el indicador 'EN LÍNEA'."],
        ],
        [3, 4.5, 9]
    )

    # ── 9. INFRAESTRUCTURA ────────────────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "9. Infraestructura y Despliegue", 1)

    heading(doc, "9.1 Stack tecnológico completo", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Capa", "Tecnología", "Versión"],
        [
            ["Lenguaje backend",    "Python",                    "3.13"],
            ["Framework API",       "FastAPI + Uvicorn",         "0.104.1 / 0.24.0"],
            ["Cola de tareas",      "Celery + Redis",            "5.3+ / 3.0.504"],
            ["ML Framework",        "PyTorch + CUDA",            "2.x + cu124"],
            ["Modelos transformers","HuggingFace transformers",  "4.40+"],
            ["Cuantización VLM",    "bitsandbytes + accelerate", "0.43+ / 0.29+"],
            ["Meta-ensemble",       "XGBoost",                   "2.x"],
            ["Detección de caras",  "facenet-pytorch (MTCNN)",   "2.5.3+"],
            ["Visión computacional","OpenCV + Pillow",           "4.9+ / 10.3+"],
            ["Validación datos",    "Pydantic v2",               "2.7+"],
            ["Logging",             "Loguru",                    "0.7.2+"],
            ["Rate limiting",       "SlowAPI",                   "0.1.9+"],
            ["Frontend framework",  "Next.js 14 (App Router)",   "14.0.4"],
            ["UI components",       "React + Tailwind CSS",      "18 / 3.x"],
            ["Animaciones",         "Framer Motion",             "10.x"],
            ["Gráficos",           "Recharts",                  "2.x"],
        ],
        [4, 6, 3]
    )

    heading(doc, "9.2 Docker Compose (producción)", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Servicio", "Imagen", "Recursos", "Función"],
        [
            ["redis",   "redis:7-alpine",  "RAM: 2 GB límite",     "Broker y backend de Celery"],
            ["api",     "deepguard-api",   "4 workers Uvicorn, CPU only", "FastAPI sin modelos — solo routing"],
            ["worker",  "deepguard-worker","GPU NVIDIA passthrough", "Celery + 5 modelos + LLaVA en GPU"],
            ["nginx",   "nginx:alpine",    "Proxy inverso",         "SSL termination, static files, rate limit"],
        ],
        [2.5, 4, 4.5, 5.5]
    )

    heading(doc, "9.3 Variables de entorno requeridas", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Variable", "Valor por defecto", "Descripción"],
        [
            ["REDIS_URL", "redis://localhost:6379/0", "URL del broker Celery"],
            ["REDIS_BACKEND", "redis://localhost:6379/1", "URL del backend de resultados"],
            ["DEEPGUARD_SIGNING_KEY", "(cambiar en producción)", "Clave HMAC para cadena de custodia"],
            ["TRANSFORMERS_CACHE", "./models", "Directorio caché de modelos HuggingFace"],
            ["HF_HOME", "./models", "Directorio HOME de HuggingFace Hub"],
            ["MAX_FILE_SIZE_MB", "500", "Tamaño máximo de archivo en MB"],
            ["ALLOWED_ORIGINS_CSV", "http://localhost:3000", "CORS orígenes permitidos (CSV)"],
        ],
        [4.5, 4.5, 7.5]
    )

    # ── 10. SEGURIDAD ─────────────────────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "10. Seguridad y Auditoría", 1)

    heading(doc, "10.1 Medidas de seguridad implementadas", 2, AZUL_MEDIO, 12)
    sec_items = [
        ("Rate limiting", "100 req/minuto por IP con SlowAPI. Configurable por variable de entorno."),
        ("Validación MIME", "Verificación del tipo MIME real del archivo, no solo la extensión. "
         "Rechazo inmediato de tipos no soportados."),
        ("Sanitización de nombres", "Nombres de archivo sanitizados con UUID prefix antes de "
         "guardar en disco. Previene path traversal."),
        ("CORS estricto", "Solo orígenes en ALLOWED_ORIGINS_CSV tienen acceso. "
         "Configurable por entorno."),
        ("Firma HMAC", "Cada resultado está firmado con HMAC-SHA256. Cualquier "
         "alteración posterior del resultado invalida el token."),
        ("Sin SQL", "No hay base de datos SQL — elimina superficie de ataque SQLi. "
         "Persistencia en archivos JSON firmados."),
        ("Cleanup automático", "Archivos subidos eliminados automáticamente tras 1 hora (3600s)."),
        ("Trusted Host", "Middleware TrustedHostMiddleware para prevenir HTTP host header attacks."),
    ]
    for label, text in sec_items:
        bullet(doc, text, bold_prefix=label)

    heading(doc, "10.2 Auditoría del código — resultados", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Check de auditoría", "Resultado"],
        [
            ["Archivos Python auditados",          "30 archivos — 0 errores de sintaxis"],
            ["Reimportaciones locales de logger",  "0 (causa de UnboundLocalError eliminada)"],
            ["BOM U+FEFF en __init__.py",          "Corregido en 2 archivos"],
            ["Rutas API legacy sin prefijo v1",    "0 en código de producción"],
            ["TypeScript tsc --noEmit",            "0 errores"],
            ["Next.js npm run build",              "✓ Compiled successfully (4 páginas estáticas)"],
            ["Tests de fusión semántica",          "6/6 PASS"],
            ["Tests de robustez del pipeline",     "7/7 PASS"],
            ["Tests de corrección de enrutamiento","8/8 PASS"],
        ],
        [8, 8]
    )

    # ── 11. LIMITACIONES ─────────────────────────────────────────────────────
    heading(doc, "11. Limitaciones y Consideraciones Éticas", 1)
    body(doc, (
        "DeepGuard AI es una herramienta de análisis probabilístico. Sus resultados NO "
        "constituyen evidencia forense definitiva ni reemplazan el juicio de un perito "
        "humano certificado. Las siguientes limitaciones deben tenerse en cuenta:"
    ))
    limits = [
        ("Naturaleza probabilística", "El sistema produce probabilidades, no certezas. "
         "Un score de 85% indica alta evidencia pero no prueba absoluta de manipulación."),
        ("Sesgos del conjunto de entrenamiento", "Los modelos fueron entrenados principalmente "
         "con datos de habla inglesa y sujetos de origen occidental. Pueden tener mayor "
         "tasa de error en otras etnias o contextos culturales."),
        ("Nuevas técnicas de generación", "Modelos de IA publicados después de la fecha de "
         "entrenamiento de los detectores pueden evadir la detección parcial o totalmente."),
        ("Compresión extrema", "Imágenes con compresión JPEG muy agresiva (factor > 70%) "
         "pueden confundir los modelos, aunque la Corrección Veto mitiga este efecto."),
        ("Videos de alta compresión", "Streams de video con bitrate muy bajo (< 500 kbps) "
         "degradan la calidad del análisis temporal y de embeddings."),
        ("LLaVA opcional", "Si LLaVA no está disponible (VRAM insuficiente), el sistema "
         "funciona en modo degradado sin análisis semántico, con menor capacidad de "
         "corrección de falsos positivos/negativos."),
    ]
    for label, text in limits:
        bullet(doc, text, bold_prefix=label)

    # ── 12. INSTRUCCIONES DE ARRANQUE ────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "12. Instrucciones de Arranque y Operación", 1)

    heading(doc, "12.1 Requisitos previos", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["Requisito", "Versión mínima", "Notas"],
        [
            ["NVIDIA GPU", "8 GB VRAM", "12 GB recomendado para LLaVA. RTX 3060 / RTX 4070 o superior."],
            ["CUDA", "12.1+", "Instalar con: pip install torch --index-url https://download.pytorch.org/whl/cu124"],
            ["Python", "3.11–3.13", "Python 3.13 recomendado. Crear venv antes de instalar."],
            ["Node.js", "18+", "Para el frontend Next.js. npm 9+ incluido."],
            ["Redis", "3.0+", "En Windows: instalado automáticamente con winget install Redis.Redis"],
            ["RAM sistema", "16 GB+", "Para overflow de LLaVA si VRAM es limitada (< 10 GB)."],
        ],
        [3, 3, 10.5]
    )

    heading(doc, "12.2 Comandos de arranque", 2, AZUL_MEDIO, 12)
    cmds = [
        ("Terminal 1 — Backend FastAPI",
         'cd backend\nvenv\\Scripts\\uvicorn.exe app.main:app --host 0.0.0.0 --port 8000 --reload\n→ Esperar: "Application startup complete."'),
        ("Terminal 2 — Frontend Next.js",
         'cd frontend\nnpm run dev\n→ Esperar: "✓ Ready in X.Xs" — luego abrir http://localhost:3000'),
        ("Terminal 3 — Celery Worker (GPU)",
         'cd backend\nvenv\\Scripts\\celery.exe -A app.celery_app worker --loglevel=info -P threads -Q images,videos,default -c 2\n→ Esperar: "celery@HOSTNAME ready." (~60s por carga de LLaVA)'),
    ]
    for title, cmd in cmds:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = AZUL_OSCURO
        p2 = doc.add_paragraph()
        r2 = p2.add_run(cmd)
        r2.font.name = "Courier New"
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = AZUL_MEDIO
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(6)

    heading(doc, "12.3 URLs del sistema en operación", 2, AZUL_MEDIO, 12)
    simple_table(doc,
        ["URL", "Descripción"],
        [
            ["http://localhost:3000",                "Aplicación web — interfaz principal"],
            ["http://localhost:8000/docs",           "Swagger UI — documentación interactiva de la API"],
            ["http://localhost:8000/api/health",     "Health check simplificado (modelo + CUDA)"],
            ["http://localhost:8000/api/v1/health",  "Health check enterprise (Redis + Celery + GPU)"],
        ],
        [7, 9.5]
    )

    # ── ÍNDICE DE ARCHIVOS CLAVE ──────────────────────────────────────────────
    add_page_break(doc)
    heading(doc, "Anexo A — Índice de Archivos Clave del Proyecto", 1)
    simple_table(doc,
        ["Archivo", "Descripción"],
        [
            ["backend/app/main.py",              "Entry point FastAPI, lifespan, CORS, routers"],
            ["backend/app/config.py",            "Configuración global (Pydantic Settings)"],
            ["backend/app/celery_app.py",        "Configuración del broker Celery + colas"],
            ["backend/app/api/v1/routes.py",     "Endpoints enterprise v1, sync_fallback, serialización"],
            ["backend/app/api/schemas.py",       "Modelos Pydantic: AnalysisResult, SemanticAnalysis, TemporalAnalysis..."],
            ["backend/app/models/deepfake_detector.py", "Carga y gestión de los 5 modelos de ensemble"],
            ["backend/app/models/meta_ensemble.py",     "XGBoost + Temperature Scaling"],
            ["backend/app/models/ood_detector.py",      "Detector OOD 5 señales heurísticas"],
            ["backend/app/services/image_service.py",   "Pipeline completo de análisis de imagen"],
            ["backend/app/services/video_service.py",   "Pipeline frame-by-frame + análisis temporal"],
            ["backend/app/services/video_temporal_service.py", "ViT embeddings + flujo óptico Farneback"],
            ["backend/app/services/semantic_inspection_service.py", "LLaVA-1.5-7b-hf + fusión semántica"],
            ["backend/app/services/analysis_service.py","Orquestador async (ruta legacy con LLaVA completo)"],
            ["backend/app/services/custody_service.py", "SHA-256 + HMAC-SHA256 cadena de custodia v2"],
            ["backend/app/services/forensic_metadata_service.py","EXIF/XMP/IPTC + firmas de generadores IA"],
            ["backend/app/tasks/analysis_tasks.py",     "Tareas Celery: analyze_image_task, analyze_video_task"],
            ["backend/app/utils/forensic_corrections.py","Reglas OOD Bypass y Compression Veto"],
            ["backend/app/utils/helpers.py",             "get_evidence_level, generate_forensic_explanation"],
            ["frontend/src/lib/api.ts",                  "Cliente Axios, normalización v1, polling"],
            ["frontend/src/types/index.ts",              "Tipos TypeScript: AnalysisResult, SemanticAnalysis..."],
            ["frontend/src/components/ResultCard.tsx",   "Card principal: gauge, desglose de modelos, tabs"],
            ["frontend/src/components/ForensicPanel.tsx","Tab Ensemble: tabla detallada + LLaVA panel"],
            ["frontend/src/app/globals.css",             "Paleta forense, variables CSS, clases utilitarias"],
        ],
        [7.5, 9]
    )

    # ── PIE DE PÁGINA ─────────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        run_f = fp.add_run("DeepGuard AI — Informe Técnico Confidencial — " + TODAY)
        run_f.font.size = Pt(8)
        run_f.font.color.rgb = RGBColor(0x7A, 0x88, 0x99)

    # Guardar
    out = r"c:\Users\gabot\OneDrive\Desktop\PROYECTO TITULO FINAL\DeepGuard_AI_Informe_Tecnico.docx"
    doc.save(out)
    print(f"Informe generado: {out}")
    return out


if __name__ == "__main__":
    build()
