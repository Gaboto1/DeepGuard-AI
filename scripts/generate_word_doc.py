"""
Genera documentación completa del proyecto DeepGuard AI en formato Word (.docx)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

ROOT = Path(__file__).parent.parent
OUT  = ROOT / "DeepGuard_AI_Documentacion.docx"


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x1E, 0x63, 0xD4)  # accent blue
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x4A, 0x6E)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
             size: int = 11, color: RGBColor = None) -> None:
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1.2 * (level + 1))


def add_table(doc: Document, headers: list, rows: list,
              col_widths: list = None, header_color: str = "1E63D4") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if r_idx % 2 == 0:
                set_cell_bg(cell, "F0F4FA")

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing


def add_code(doc: Document, code: str) -> None:
    p   = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ─── Build Document ───────────────────────────────────────────────────────────

def build() -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ── PORTADA ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("DeepGuard AI")
    t_run.font.size  = Pt(32)
    t_run.font.bold  = True
    t_run.font.color.rgb = RGBColor(0x1E, 0x63, 0xD4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("Plataforma de Análisis Forense Digital\npara Detección de Deepfakes e Imágenes Generadas por IA")
    s_run.font.size   = Pt(16)
    s_run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x6E)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Documentación Técnica Completa\n").font.size = Pt(12)
    meta.add_run(f"Versión 5.0  ·  {datetime.date.today().strftime('%d de %B de %Y')}").font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tech badges
    tech_line = doc.add_paragraph()
    tech_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tech_run = tech_line.add_run("Python  ·  FastAPI  ·  PyTorch  ·  Next.js  ·  CUDA  ·  RTX 4070 SUPER")
    tech_run.font.size   = Pt(10)
    tech_run.font.italic = True
    tech_run.font.color.rgb = RGBColor(0x7A, 0x88, 0x99)

    doc.add_page_break()

    # ── ÍNDICE ────────────────────────────────────────────────────────────────
    add_heading(doc, "Índice de Contenidos", level=1)
    indice = [
        ("1", "Resumen Ejecutivo"),
        ("2", "Descripción del Proyecto"),
        ("3", "Tecnologías Aplicadas"),
        ("4", "Arquitectura del Sistema"),
        ("5", "Modelos de Inteligencia Artificial"),
        ("6", "Meta-Ensemble y Calibración"),
        ("7", "Análisis Forense y Funcionalidades"),
        ("8", "Pipeline de Análisis"),
        ("9", "Resultados y Métricas"),
        ("10", "Benchmark y Evaluación"),
        ("11", "API y Endpoints"),
        ("12", "Interfaz de Usuario"),
        ("13", "Instalación y Configuración"),
        ("14", "Limitaciones y Trabajo Futuro"),
    ]
    for num, titulo in indice:
        p   = doc.add_paragraph()
        run = p.add_run(f"  {num}. {titulo}")
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 1. RESUMEN EJECUTIVO
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Resumen Ejecutivo", level=1)
    add_para(doc,
        "DeepGuard AI es una plataforma profesional de análisis forense digital que utiliza "
        "un ensemble de cinco modelos de inteligencia artificial para determinar la probabilidad "
        "de que una imagen o video haya sido manipulado o generado artificialmente mediante "
        "técnicas como deepfake, Stable Diffusion, MidJourney, FLUX, DALL-E o cualquier otra "
        "tecnología generativa moderna.")

    add_para(doc,
        "A diferencia de los clasificadores binarios convencionales que emiten veredictos "
        "absolutos (REAL / FAKE), DeepGuard AI opera como una herramienta forense basada en "
        "evidencia probabilística, mostrando al usuario la probabilidad de manipulación, "
        "el nivel de evidencia, el consenso entre modelos y la incertidumbre del análisis.")

    doc.add_paragraph()
    add_para(doc, "Casos de uso principales:", bold=True)
    for caso in [
        "Periodistas y equipos de verificación de hechos (fact-checking)",
        "Analistas OSINT e investigadores digitales",
        "Equipos de Centros de Operaciones de Seguridad (SOC)",
        "Peritos digitales y forenses",
        "Investigadores académicos en autenticidad de medios",
    ]:
        add_bullet(doc, caso)

    doc.add_paragraph()
    add_heading(doc, "Indicadores Clave de Rendimiento", level=2)
    add_table(doc,
        ["Métrica", "Valor", "Descripción"],
        [
            ["F1 Score", "84.2%", "Golden set (20 imágenes independientes)"],
            ["ROC-AUC", "0.910", "Poder discriminativo del ensemble"],
            ["F1 (meta, CV)", "95.1%", "LightGBM en cross-validation 5-fold, 512 imágenes"],
            ["FPR", "10%", "Falsos positivos en fotografías reales"],
            ["ECE", "0.202", "Error de calibración esperado"],
            ["Brier Score", "0.128", "Calibración probabilística"],
            ["VRAM usada", "1.93 GB", "De 12.9 GB disponibles (15%)"],
            ["Latencia imagen", "< 2s", "En RTX 4070 SUPER con CUDA"],
        ],
        col_widths=[5, 3, 8],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. DESCRIPCIÓN DEL PROYECTO
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Descripción del Proyecto", level=1)

    add_heading(doc, "2.1 Problema que Resuelve", level=2)
    add_para(doc,
        "La proliferación de herramientas de inteligencia artificial generativa ha democratizado "
        "la creación de imágenes y videos sintéticos de alta calidad. Tecnologías como Stable "
        "Diffusion, MidJourney, FLUX, DALL-E y los modelos de intercambio facial (deepfakes) "
        "permiten crear contenido visual indistinguible de la realidad para el ojo humano.")
    add_para(doc,
        "Esto representa un riesgo significativo para la integridad informativa, la confianza "
        "en medios visuales y los procesos judiciales que dependen de evidencia fotográfica o "
        "videográfica. DeepGuard AI proporciona una capa de análisis automatizado que ayuda a "
        "evaluar la autenticidad de contenido visual mediante métodos basados en evidencia.")

    add_heading(doc, "2.2 Filosofía de Diseño", level=2)
    add_para(doc, "El sistema se diseñó bajo dos principios fundamentales:", bold=True)
    add_bullet(doc,
        "No emitir veredictos absolutos: El sistema nunca dice REAL o FAKE. "
        "Muestra probabilidades, niveles de evidencia y consenso entre modelos.")
    add_bullet(doc,
        "Transparencia total: Cada resultado incluye los scores individuales de cada modelo, "
        "los pesos aplicados, la contribución de cada modelo al resultado final, los metadatos "
        "EXIF extraídos y los links de verificación externa.")

    add_heading(doc, "2.3 Evolución del Proyecto", level=2)
    add_table(doc,
        ["Versión", "Característica Principal"],
        [
            ["v1.0", "Sistema básico con modelo ViT single"],
            ["v2.0", "Triple ensemble (ViT + SDXL + EfficientNet-B0)"],
            ["v3.0", "Frontend forense en español, eliminación de etiquetas binarias"],
            ["v4.0", "5 modelos: +haywoodsloan/ai-art-detector + SigLIP"],
            ["v5.0", "Meta-ensemble LightGBM, incertidumbre, detección OOD"],
        ],
        col_widths=[3, 13],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. TECNOLOGÍAS APLICADAS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Tecnologías Aplicadas", level=1)

    add_heading(doc, "3.1 Hardware y Entorno de Ejecución", level=2)
    add_table(doc,
        ["Componente", "Especificación"],
        [
            ["GPU", "NVIDIA GeForce RTX 4070 SUPER"],
            ["VRAM", "12 GB GDDR6X"],
            ["CUDA", "12.4"],
            ["Sistema Operativo", "Windows 11 Home"],
            ["RAM", "16 GB+"],
            ["Almacenamiento", "SSD NVMe (para modelos y dataset)"],
        ],
        col_widths=[5, 11],
    )

    add_heading(doc, "3.2 Stack de Backend", level=2)
    add_table(doc,
        ["Tecnología", "Versión", "Rol"],
        [
            ["Python", "3.13.0", "Lenguaje principal del backend"],
            ["PyTorch", "2.6.0+cu124", "Framework de deep learning con soporte CUDA"],
            ["FastAPI", "0.104.1", "Framework web asíncrono para la API REST"],
            ["Uvicorn", "0.24.0", "Servidor ASGI de alto rendimiento"],
            ["Transformers (HuggingFace)", "4.40+", "Carga de modelos ViT, Swin, SigLIP"],
            ["OpenCV", "4.13.0", "Procesamiento de video y extracción de frames"],
            ["facenet-pytorch", "2.6.0", "MTCNN para detección facial"],
            ["pytorch-grad-cam", "1.4.8", "Generación de mapas de atención Grad-CAM++"],
            ["LightGBM", "4.x", "Meta-ensemble para combinación de scores"],
            ["scikit-learn", "1.x", "Métricas, calibración, Logistic Regression"],
            ["piexif / Pillow", "12.2.0", "Extracción de metadatos EXIF"],
            ["imagehash", "4.x", "Hash perceptual para verificación OSINT"],
            ["joblib", "1.x", "Serialización del meta-modelo"],
            ["loguru", "0.7.2", "Sistema de logging estructurado"],
        ],
        col_widths=[5.5, 3, 7.5],
    )

    add_heading(doc, "3.3 Stack de Frontend", level=2)
    add_table(doc,
        ["Tecnología", "Versión", "Rol"],
        [
            ["Next.js", "14.0.4", "Framework React con App Router"],
            ["React", "18.2.0", "Librería de interfaz de usuario"],
            ["TypeScript", "5.3.3", "Tipado estático para mayor robustez"],
            ["TailwindCSS", "3.4.0", "Sistema de diseño utilitario"],
            ["Framer Motion", "10.16.16", "Animaciones de transición"],
            ["Recharts", "2.10.3", "Gráficas de timeline de video"],
            ["react-dropzone", "14.2.3", "Zona de carga con drag & drop"],
            ["lucide-react", "0.303.0", "Iconografía profesional"],
            ["axios", "1.6.2", "Cliente HTTP para llamadas a la API"],
        ],
        col_widths=[5.5, 3, 7.5],
    )

    add_heading(doc, "3.4 Técnicas de Machine Learning Aplicadas", level=2)
    for tecnica in [
        "Vision Transformer (ViT) — para clasificación de imágenes completas",
        "Swin Transformer (Swin-B, Swin-v2) — para detección de imágenes SDXL/AI-art",
        "EfficientNet-B0/B4 — especializado en deepfakes faciales (FaceForensics++)",
        "SigLIP — clasificador deepfake basado en arquitectura Google SigLIP",
        "MTCNN — detección facial multi-escala para localización de caras",
        "Grad-CAM++ — mapas de atención explicables sobre arquitecturas ViT",
        "Meta-ensemble LightGBM — combinación óptima de scores mediante gradient boosting",
        "Temperature Scaling — calibración probabilística de salidas del ensemble",
        "Detección OOD (Out-of-Distribution) — identificación de imágenes fuera del dominio",
        "Perceptual Hashing — generación de fingerprint visual para deduplicación",
        "Cross-validation estratificada (5-fold) — evaluación no sesgada del meta-modelo",
        "Grid Search de pesos — optimización automática de ponderaciones del ensemble",
    ]:
        add_bullet(doc, tecnica)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 4. ARQUITECTURA DEL SISTEMA
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Arquitectura del Sistema", level=1)

    add_heading(doc, "4.1 Visión General", level=2)
    add_para(doc,
        "DeepGuard AI sigue una arquitectura de microservicios desacoplada con un frontend "
        "React y un backend FastAPI que expone una API REST asíncrona. El análisis se ejecuta "
        "en background mediante asyncio + ThreadPoolExecutor para no bloquear el event loop "
        "durante la inferencia GPU.")

    add_code(doc,
"""Usuario (Navegador)
      ↓  HTTP Multipart POST
  [Next.js 14 — :3000]
      ↓  HTTP REST
  [FastAPI — :8000]
      ↓
  [Task Queue (asyncio)]
      ↓
  [5-Model Ensemble + LightGBM Meta]
      ↓  GPU (CUDA 12.4)
  [RTX 4070 SUPER]
      ↓
  [Resultado JSON → Frontend]""")

    add_heading(doc, "4.2 Estructura de Carpetas", level=2)
    add_code(doc,
"""PROYECTO TITULO FINAL/
├── backend/
│   ├── app/
│   │   ├── main.py                ← FastAPI app, lifespan con precarga de modelos
│   │   ├── config.py              ← Configuración (DEVICE, MODELS_DIR, etc.)
│   │   ├── api/
│   │   │   ├── routes.py          ← Endpoints REST
│   │   │   └── schemas.py         ← Modelos Pydantic (EvidenceLevel, UncertaintyLevel...)
│   │   ├── models/
│   │   │   ├── deepfake_detector.py ← 5-model ensemble + integración meta
│   │   │   ├── face_detector.py   ← MTCNN
│   │   │   └── meta_ensemble.py   ← LightGBM meta-classifier
│   │   └── services/
│   │       ├── analysis_service.py  ← Pipeline principal
│   │       ├── image_service.py     ← Análisis imagen (dual: cara + completa)
│   │       ├── video_service.py     ← Análisis video (frame extraction)
│   │       ├── metadata_service.py  ← EXIF
│   │       └── osint_service.py     ← Verificación externa
│   └── requirements.txt
├── frontend/                      ← Next.js 14 App Router
│   └── src/
│       ├── app/page.tsx           ← Página principal
│       └── components/
│           ├── ResultCard.tsx     ← Vista de resultados con pestañas
│           ├── ForensicPanel.tsx  ← Desglose del ensemble
│           ├── MetadataPanel.tsx  ← EXIF
│           └── OsintPanel.tsx     ← Verificación externa
├── models/
│   ├── meta_ensemble/             ← LightGBM entrenado
│   └── trained/                   ← EfficientNet-B4 (si entrenado)
├── tests/
│   ├── golden_set/                ← 25 imágenes de referencia
│   ├── benchmark_extended/        ← 45 imágenes, 10 categorías
│   └── benchmark_massive/         ← 512 imágenes, 16 categorías
└── scripts/                       ← Pipeline de entrenamiento y evaluación""")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. MODELOS DE IA
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Modelos de Inteligencia Artificial", level=1)

    add_heading(doc, "5.1 Los 5 Modelos del Ensemble", level=2)
    add_para(doc,
        "El sistema utiliza cinco modelos especializados que se ejecutan en paralelo "
        "sobre la misma imagen. Cada modelo aporta una perspectiva diferente al análisis.")

    add_table(doc,
        ["ID", "Modelo", "Arquitectura", "Especialización", "VRAM"],
        [
            ["A", "prithivMLmods/Deep-Fake-Detector-v2", "ViT-base (85.8M params)", "Deepfakes faciales (face-swap, reenactment)", "350 MB"],
            ["B", "Organika/sdxl-detector", "Swin Transformer (86.8M)", "SDXL vs fotografías reales — el más fuerte", "350 MB"],
            ["C", "Xicor9/efficientnet-b0-ffpp-c23", "EfficientNet-B0 (5.3M)", "FaceForensics++ c23 — face-swaps 2019-2022", "50 MB"],
            ["D", "haywoodsloan/ai-image-detector-deploy", "Swin v2 (195.2M)", "Arte IA moderno (MJ, FLUX, SDXL, DALL-E, Ideogram)", "800 MB"],
            ["E", "prithivMLmods/Deepfake-Detect-Siglip2", "SigLIP (92.9M)", "Clasificación deepfake general", "381 MB"],
        ],
        col_widths=[0.8, 5.5, 4, 5, 2],
    )

    add_heading(doc, "5.2 Pesos del Ensemble", level=2)
    add_para(doc,
        "Los pesos se ajustan dinámicamente según si MTCNN detecta o no un rostro en la imagen. "
        "El meta-modelo LightGBM aprende la combinación óptima a partir de datos.")

    add_table(doc,
        ["Modelo", "Peso sin cara", "Peso con cara", "Justificación"],
        [
            ["A — ViT face-deepfake", "15%", "35%", "Efectivo con caras, alta FPR sin cara"],
            ["B — SDXL Detector", "70%", "25%", "Mejor modelo standalone, dominante sin cara"],
            ["C — EfficientNet-B0", "5%", "5%", "Limitado a face-swaps 2019, peso mínimo"],
            ["D — AI Art Detector", "5%", "25%", "Especialista en imágenes reales de plataformas IA"],
            ["E — SigLIP", "5%", "10%", "FPR alto solo, útil en ensemble con LightGBM"],
        ],
        col_widths=[5, 3, 3, 5],
    )

    add_heading(doc, "5.3 Detección Facial (MTCNN)", level=2)
    add_para(doc,
        "Antes de pasar la imagen al ensemble, el sistema ejecuta MTCNN (Multi-task "
        "Cascaded Convolutional Networks) para detectar rostros. Si se detecta al menos "
        "un rostro, el análisis se realiza en modo dual:")
    add_bullet(doc, "60% del score proviene del análisis del recorte facial (face crop)")
    add_bullet(doc, "40% del score proviene del análisis de la imagen completa")
    add_para(doc,
        "Esto permite capturar tanto las manipulaciones faciales (face-swap, reenactment) "
        "como los artefactos globales de generación AI (colores imposibles, texturas sintéticas).")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 6. META-ENSEMBLE Y CALIBRACIÓN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Meta-Ensemble y Calibración", level=1)

    add_heading(doc, "6.1 ¿Qué es el Meta-Ensemble?", level=2)
    add_para(doc,
        "En lugar de usar pesos fijos para combinar los 5 modelos (lo cual requiere suposiciones "
        "manuales sobre qué modelo es más importante), DeepGuard v5.0 entrena un meta-clasificador "
        "LightGBM que aprende automáticamente la combinación óptima a partir de datos reales.")
    add_para(doc,
        "El meta-modelo recibe como entrada los 5 scores individuales y produce una "
        "probabilidad calibrada final. Fue entrenado con validación cruzada estratificada "
        "(5-fold) sobre 512 imágenes etiquetadas (256 reales, 256 generadas por IA).")

    add_heading(doc, "6.2 Comparativa de Métodos de Combinación", level=2)
    add_table(doc,
        ["Método", "F1 (CV 5-fold)", "ECE", "Observación"],
        [
            ["Pesos fijos actuales (0.15/0.70/...)", "58.6%", "0.193", "Baseline, manual"],
            ["Grid-search de pesos", "69.1%", "0.209", "Mejor que fijos, pero subóptimo"],
            ["Logistic Regression", "81.9%", "0.077", "Buen balance calidad/velocidad"],
            ["XGBoost", "94.3%", "0.038", "Excelente — casi tan bueno como LightGBM"],
            ["LightGBM (seleccionado)", "95.1%", "0.039", "Mejor F1 y calibración"],
        ],
        col_widths=[6, 3, 2, 5],
        header_color="1A4080",
    )

    add_heading(doc, "6.3 Cuantificación de Incertidumbre", level=2)
    add_para(doc,
        "El sistema calcula automáticamente tres indicadores de incertidumbre para cada análisis:")
    add_table(doc,
        ["Campo", "Descripción", "Valores"],
        [
            ["uncertainty", "Nivel de incertidumbre del resultado", "Baja / Moderada / Alta"],
            ["uncertainty_score", "Entropía normalizada de la predicción (0-1)", "0.0 = certeza, 1.0 = máxima duda"],
            ["ood_signal", "Imagen posiblemente fuera del dominio de entrenamiento", "True / False"],
            ["risk_of_error", "Descripción del riesgo en lenguaje natural", "Texto explicativo"],
        ],
        col_widths=[4, 7, 5],
    )

    add_heading(doc, "6.4 Detección OOD (Out-of-Distribution)", level=2)
    add_para(doc,
        "Si todos los modelos devuelven scores cercanos a 0.5 simultáneamente con baja "
        "desviación estándar, el sistema activa la señal OOD (ood_signal=True). Esto indica "
        "que la imagen podría ser de un tipo no visto durante el entrenamiento, como:")
    for ejemplo in [
        "Memes con texto superpuesto",
        "Renders 3D de videojuegos",
        "Ilustraciones médicas o técnicas",
        "Fotografías históricas extremadamente antiguas",
        "Arte abstracto no fotorrealista",
    ]:
        add_bullet(doc, ejemplo)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 7. ANÁLISIS FORENSE Y FUNCIONALIDADES
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Análisis Forense y Funcionalidades", level=1)

    add_heading(doc, "7.1 Niveles de Evidencia", level=2)
    add_para(doc,
        "El sistema reemplaza las etiquetas binarias (REAL/FAKE) por niveles de evidencia "
        "que reflejan la magnitud de la probabilidad sin emitir un juicio definitivo:")
    add_table(doc,
        ["Probabilidad", "Nivel de Evidencia", "Interpretación"],
        [
            ["0% – 20%", "Evidencia Muy Baja", "No se detectaron patrones de manipulación"],
            ["20% – 40%", "Evidencia Baja", "Señales débiles dentro de la variación normal"],
            ["40% – 60%", "Inconclusa", "Los modelos no llegan a consenso — revisar manualmente"],
            ["60% – 80%", "Evidencia Moderada", "Múltiples modelos detectan patrones sospechosos"],
            ["80% – 100%", "Evidencia Fuerte", "Alta convergencia entre modelos en manipulación"],
        ],
        col_widths=[3.5, 4.5, 8],
    )

    add_heading(doc, "7.2 Análisis de Metadatos EXIF", level=2)
    add_para(doc,
        "El sistema extrae automáticamente todos los metadatos disponibles en el archivo "
        "y genera notas forenses. Los metadatos no modifican el score — son evidencia "
        "complementaria para el analista.")
    add_table(doc,
        ["Campo EXIF", "Relevancia Forense"],
        [
            ["Camera Make / Model", "Confirma origen en cámara real vs software"],
            ["Software", "Detecta Photoshop, Lightroom, o software generativo IA"],
            ["DateTime Original", "Permite verificar cronología del contenido"],
            ["GPS Coordinates", "Ubica geográficamente el origen de la imagen"],
            ["ISO / Focal Length / Aperture", "Parámetros físicos de captura real"],
            ["Ausencia de EXIF", "Común en screenshots, redes sociales e imágenes IA"],
        ],
        col_widths=[5, 11],
    )

    add_heading(doc, "7.3 Verificación Externa (OSINT)", level=2)
    add_para(doc,
        "El módulo OSINT proporciona herramientas para verificación externa:")
    add_bullet(doc, "Links directos a TinEye, Google Lens, Bing Visual Search y Yandex Images")
    add_bullet(doc, "Hash perceptual (phash + dhash) para detectar imágenes casi idénticas")
    add_bullet(doc, "Lista de fuentes prioritarias: Reuters, AP, Getty Images, AFP, BBC, Wikimedia")
    add_bullet(doc, "Disclaimer explícito: encontrar la imagen en una fuente confiable no confirma autenticidad")

    add_heading(doc, "7.4 Mapa de Atención (Grad-CAM++)", level=2)
    add_para(doc,
        "Para imágenes con probabilidad de manipulación superior al 42%, el sistema genera "
        "automáticamente un mapa de atención usando Grad-CAM++ sobre el modelo ViT (Modelo A). "
        "El mapa muestra visualmente qué regiones de la imagen influyeron más en la predicción, "
        "sin afirmar que esas regiones estén manipuladas — es la atención del modelo, "
        "no un diagnóstico de artefactos.")

    add_heading(doc, "7.5 Análisis de Video", level=2)
    add_para(doc,
        "Para archivos de video (MP4, MOV, MKV, WEBM), el sistema extrae hasta 50 frames "
        "equidistantes y ejecuta el ensemble sobre cada uno. Los resultados incluyen:")
    add_bullet(doc, "Probabilidad promedio de manipulación sobre todos los frames")
    add_bullet(doc, "Gráfica de timeline mostrando la evolución frame a frame")
    add_bullet(doc, "Puntuación de inconsistencia temporal (σ) — variación entre frames")
    add_bullet(doc, "Cantidad de frames con rostro detectado")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 8. PIPELINE DE ANÁLISIS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. Pipeline de Análisis", level=1)

    add_heading(doc, "8.1 Imagen — Flujo Completo", level=2)
    steps_img = [
        ("1", "Recepción", "POST /api/analyze con multipart/form-data"),
        ("2", "Validación", "MIME type, extensión, tamaño (máx 500MB)"),
        ("3", "Task creation", "UUID único asignado, estado: pending"),
        ("4", "Background task", "asyncio + ThreadPoolExecutor (no bloquea el servidor)"),
        ("5", "Detección facial", "MTCNN en GPU — detecta y recorta rostros"),
        ("6", "Inferencia dual", "Si hay cara: 60% face crop + 40% imagen completa"),
        ("7", "5 modelos GPU", "A+B+C+D+E en paralelo (PyTorch CUDA)"),
        ("8", "Meta-ensemble", "LightGBM combina 5 scores → probabilidad calibrada"),
        ("9", "Post-processing", "EvidenceLevel, ModelAgreement, Uncertainty, OOD"),
        ("10", "Grad-CAM++", "Solo si prob > 42%, sobre modelo ViT"),
        ("11", "EXIF", "piexif + Pillow, notas forenses automáticas"),
        ("12", "OSINT", "Hash perceptual, links de búsqueda"),
        ("13", "Respuesta", "JSON con todos los campos → frontend polling"),
    ]
    add_table(doc,
        ["Paso", "Etapa", "Descripción"],
        steps_img,
        col_widths=[1.5, 4, 10.5],
    )

    add_heading(doc, "8.2 Endpoints de la API", level=2)
    add_table(doc,
        ["Método", "Endpoint", "Descripción"],
        [
            ["POST", "/api/analyze", "Sube imagen o video e inicia análisis"],
            ["GET", "/api/tasks/{id}", "Consulta estado y resultado de un análisis"],
            ["GET", "/api/history", "Lista los últimos 20 análisis"],
            ["DELETE", "/api/tasks/{id}", "Elimina un análisis del historial"],
            ["GET", "/api/health", "Estado del sistema, GPU, modelos cargados"],
            ["GET", "/docs", "Documentación interactiva Swagger UI"],
        ],
        col_widths=[2, 5, 9],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 9. RESULTADOS Y MÉTRICAS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. Resultados y Métricas", level=1)

    add_heading(doc, "9.1 Ranking de Modelos (Golden Set)", level=2)
    add_table(doc,
        ["Posición", "Modelo", "F1", "ROC-AUC", "FPR", "FNR"],
        [
            ["🥇 1°", "SDXL Detector (B)", "85.7%", "0.880", "20%", "10%"],
            ["🥈 2°", "Ensemble LightGBM", "84.2%", "0.910", "10%", "20%"],
            ["🥉 3°", "Face-Deepfake ViT (A)", "64.3%", "0.360", "90%", "10%"],
            ["4°", "AI Art Detector (D)", "~58%", "~0.40", "70%", "30%"],
            ["5°", "SigLIP (E)", "~60%", "~0.45", "60%", "25%"],
            ["Último", "EfficientNet-B0 (C)", "13.3%", "0.335", "40%", "90%"],
        ],
        col_widths=[2, 5.5, 2, 2.5, 2, 2],
    )

    add_heading(doc, "9.2 Impacto del Meta-Ensemble", level=2)
    add_table(doc,
        ["Métrica", "Antes (pesos fijos)", "Después (LightGBM)", "Mejora"],
        [
            ["F1 Score", "85.7%", "84.2%", "-1.5% (tradeoff)"],
            ["ROC-AUC", "0.880", "0.910", "+3.4% ↑"],
            ["FPR", "20%", "10%", "-50% ↓↓"],
            ["Brier Score", "N/A", "0.128", "Mejor calibrado"],
            ["F1 en Massive (CV)", "49%", "95.1%", "+46pp ↑↑↑"],
        ],
        col_widths=[5, 4, 4, 3],
        header_color="1A4080",
    )
    add_para(doc,
        "La reducción del 50% en la tasa de falsos positivos (FPR) es el logro más "
        "importante, directamente alineado con el objetivo de no flaggear fotografías "
        "reales de alta calidad.", italic=True)

    add_heading(doc, "9.3 Limitaciones Conocidas", level=2)
    add_table(doc,
        ["Limitación", "Severidad", "Descripción"],
        [
            ["Arte fantástico / anime", "Alta", "Sin modelo especializado — FN elevado en este estilo"],
            ["Fotografía vintage/sepia", "Media", "SDXL detector genera FP sistemáticos"],
            ["Imágenes FLUX/GPT Image reales", "Media", "Sin datos de evaluación con imágenes reales de plataformas"],
            ["Sin análisis temporal de video", "Media", "Solo frame-by-frame, sin modelo de movimiento"],
            ["ECE = 0.20", "Media", "Calibración mejorable con más datos etiquetados"],
        ],
        col_widths=[5, 2.5, 8.5],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 10. BENCHMARK Y EVALUACIÓN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. Benchmark y Evaluación", level=1)

    add_heading(doc, "10.1 Conjuntos de Datos de Evaluación", level=2)
    add_table(doc,
        ["Set", "Imágenes", "Categorías", "Descripción"],
        [
            ["Golden Set", "25 (20 etiquetadas)", "3", "Referencia principal independiente"],
            ["Extended", "45", "10", "Selfies, deportes, paisajes + estilos IA variados"],
            ["Massive", "512", "16", "8 reales × 32 + 8 IA × 32 imágenes"],
        ],
        col_widths=[3, 5, 2.5, 5.5],
    )
    add_para(doc,
        "Todas las imágenes son generadas programáticamente con ground truth conocido. "
        "Son independientes del dataset FaceForensics++ usado en entrenamiento.")

    add_heading(doc, "10.2 Categorías del Benchmark Masivo", level=2)
    add_table(doc,
        ["Tipo", "Categorías"],
        [
            ["Reales", "Selfies, Deportes, Naturaleza, Paisajes, Arquitectura, Nocturnas, Noticias, Fotografía Profesional"],
            ["IA", "Midjourney-style, FLUX-style, GPT Image-style, SDXL-style, Ideogram-style, Anime IA, Arte IA, Paisajes IA"],
        ],
        col_widths=[3, 13],
    )

    add_heading(doc, "10.3 Scripts de Evaluación Disponibles", level=2)
    add_table(doc,
        ["Script", "Función"],
        [
            ["scripts/build_massive_benchmark.py", "Genera 512 imágenes en 16 categorías"],
            ["scripts/run_full_evaluation.py --set massive", "Evaluación completa 5 modelos + robustez"],
            ["scripts/train_meta_ensemble.py", "Compara LR/XGBoost/LightGBM y entrena el mejor"],
            ["scripts/calibration_analysis.py", "ECE, Reliability Diagram, Temperature Scaling"],
            ["scripts/compare_models.py", "Comparativa B4 vs B0 vs SDXL vs Ensemble"],
            ["scripts/benchmark_golden_set.py", "Benchmark rápido sobre el golden set"],
        ],
        col_widths=[7, 9],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 11. INSTALACIÓN Y CONFIGURACIÓN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. Instalación y Configuración", level=1)

    add_heading(doc, "11.1 Requisitos del Sistema", level=2)
    add_table(doc,
        ["Componente", "Mínimo", "Recomendado"],
        [
            ["GPU NVIDIA", "6 GB VRAM", "12 GB VRAM (RTX 4070 SUPER)"],
            ["CUDA", "12.1", "12.4"],
            ["Python", "3.11", "3.13"],
            ["Node.js", "18", "24"],
            ["RAM", "8 GB", "16 GB+"],
            ["Almacenamiento", "5 GB (modelos)", "10+ GB (con dataset)"],
        ],
        col_widths=[4, 4, 8],
    )

    add_heading(doc, "11.2 Instalación Rápida (Windows)", level=2)
    add_code(doc,
"""# 1. Ejecutar setup automático
powershell -ExecutionPolicy Bypass -File setup.ps1

# 2. Iniciar todos los servicios (doble clic)
START DEEPGUARD.bat

# Acceder a:
#   http://localhost:3000  (Frontend)
#   http://localhost:8000  (API)
#   http://localhost:8000/docs  (Documentación API)""")

    add_heading(doc, "11.3 Instalación Manual", level=2)
    add_code(doc,
"""# Backend
cd backend
python -m venv venv
venv\\Scripts\\activate
pip install torch torchvision --index-url https://download.pytorch.org/whl/cu124
pip install -r requirements.txt

# Frontend
cd frontend
npm install
echo "NEXT_PUBLIC_API_URL=http://localhost:8000" > .env.local
npm run dev""")

    add_heading(doc, "11.4 Variables de Entorno (.env)", level=2)
    add_table(doc,
        ["Variable", "Valor por defecto", "Descripción"],
        [
            ["DEVICE", "cuda", "Dispositivo de inferencia (cuda / cpu)"],
            ["MODEL_NAME", "dima806/...", "Modelo principal HuggingFace"],
            ["MAX_FILE_SIZE_MB", "500", "Tamaño máximo de archivo"],
            ["MAX_FRAMES", "50", "Máximo de frames para análisis de video"],
            ["MODELS_DIR", "path/to/models", "Caché de modelos descargados"],
        ],
        col_widths=[5, 5, 6],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 12. TRABAJO FUTURO
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12. Limitaciones y Trabajo Futuro", level=1)

    add_heading(doc, "12.1 Limitaciones Actuales", level=2)
    for lim in [
        "El benchmark usa imágenes sintéticas programáticas, no imágenes reales de MidJourney o FLUX. "
        "El rendimiento real con imágenes de plataformas puede diferir.",
        "No existe modelo open-source para detectar específicamente imágenes FLUX 1.x, GPT Image o Ideogram 2.x.",
        "El análisis de video es frame-by-frame sin modelo de flujo temporal. No detecta manipulaciones "
        "que solo son visibles en la secuencia de movimiento.",
        "La verificación externa OSINT requiere acción manual. No hay integración automática con "
        "TinEye, Google Lens ni Bing sin claves API.",
        "EfficientNet-B0 tiene F1=13.3% — prácticamente obsoleto para imágenes IA modernas.",
    ]:
        add_bullet(doc, lim)

    add_heading(doc, "12.2 Mejoras Recomendadas (por prioridad)", level=2)
    add_table(doc,
        ["Prioridad", "Mejora", "Impacto Esperado"],
        [
            ["Alta", "Evaluar con imágenes reales de MidJourney/FLUX/SDXL/DALL-E", "Métricas más realistas"],
            ["Alta", "Integrar TinEye API ($200/año)", "OSINT automático"],
            ["Alta", "Reentrenar meta-ensemble con datos reales", "Generalización mejorada"],
            ["Media", "Modelo para anime / arte fantástico", "Cubrir gap de FN"],
            ["Media", "Integrar EfficientNet-B4 (ya entrenado)", "Reemplazar B0 obsoleto"],
            ["Media", "Modelo temporal para video (AltFreezing)", "Detección en secuencias"],
            ["Baja", "Redis para task queue multi-worker", "Escalabilidad"],
            ["Baja", "PostgreSQL para historial persistente", "Persistencia de datos"],
        ],
        col_widths=[2, 6, 8],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # GLOSARIO
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Glosario de Términos", level=1)
    add_table(doc,
        ["Término", "Definición"],
        [
            ["Deepfake", "Técnica que usa IA para reemplazar el rostro de una persona en video/imagen"],
            ["Diffusion Model", "Arquitectura generativa que crea imágenes mediante proceso iterativo de denoising"],
            ["Ensemble", "Combinación de múltiples modelos para obtener una predicción más robusta"],
            ["ECE (Expected Calibration Error)", "Mide qué tan bien una probabilidad predicha refleja la frecuencia real"],
            ["EvidenceLevel", "Clasificación del resultado en 5 niveles (Muy Baja → Fuerte) sin afirmar verdad absoluta"],
            ["FPR (False Positive Rate)", "Proporción de imágenes reales incorrectamente clasificadas como IA"],
            ["FNR (False Negative Rate)", "Proporción de imágenes IA que el sistema no detecta"],
            ["Grad-CAM++", "Técnica de explicabilidad que muestra qué regiones de la imagen influyeron en la decisión"],
            ["LightGBM", "Algoritmo de gradient boosting optimizado para velocidad y eficiencia de memoria"],
            ["MTCNN", "Multi-task Cascaded CNN para detección y alineación facial en tiempo real"],
            ["OOD (Out-of-Distribution)", "Imagen de un tipo no visto durante el entrenamiento del modelo"],
            ["Perceptual Hash", "Fingerprint visual que identifica imágenes similares aunque tengan pequeñas diferencias"],
            ["ROC-AUC", "Area bajo la curva ROC — mide la capacidad discriminativa general del clasificador"],
            ["SigLIP", "Arquitectura de Google que combina visión y lenguaje para clasificación de imágenes"],
            ["Swin Transformer", "Transformers jerárquicos para visión, más eficientes que ViT en imágenes grandes"],
            ["Temperature Scaling", "Técnica post-hoc de calibración que ajusta la confianza de las predicciones"],
            ["ViT (Vision Transformer)", "Transformer aplicado directamente a parches de imagen para clasificación"],
        ],
        col_widths=[5.5, 10.5],
    )

    # ─── Footer en cada página ────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp_p   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp_p.clear()
        fp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp_p.add_run(f"DeepGuard AI — Documentación Técnica v5.0  ·  {datetime.date.today().strftime('%B %Y')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7A, 0x88, 0x99)

    doc.save(str(OUT))
    print(f"Documento generado: {OUT}")
    size_mb = OUT.stat().st_size / 1e6
    print(f"Tamaño: {size_mb:.2f} MB")


if __name__ == "__main__":
    build()
